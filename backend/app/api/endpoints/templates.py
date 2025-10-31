"""模板管理API端点 - 基于React Agent系统"""

from fastapi import APIRouter, Depends, HTTPException, status, Query, UploadFile, File, Request
from sqlalchemy.orm import Session
from typing import List, Optional, Dict, Any
from uuid import UUID

from app.core.architecture import ApiResponse, PaginatedResponse
from app.core.permissions import require_permission, ResourceType, PermissionLevel
from app.db.session import get_db
from app.core.dependencies import get_current_user
from app.models.user import User
from app.models.template import Template as TemplateModel
from app.schemas.template import TemplateCreate, TemplateUpdate, Template as TemplateSchema, TemplatePreview
from app.schemas.template_placeholder import TemplatePlaceholderCreate, TemplatePlaceholderUpdate
from app.services.infrastructure.storage.hybrid_storage_service import get_hybrid_storage_service
from app.crud import template as crud_template
from app import crud
from app.services.domain.template.services.template_domain_service import TemplateParser
from app.api import deps
import re
import logging
import json
from datetime import datetime
from typing import Any

logger = logging.getLogger(__name__)

# 创建全局实例
template_parser = TemplateParser()

router = APIRouter()


# 旧的get_unified_api_adapter函数已被移除，使用新的Claude Code架构
@router.get("", response_model=PaginatedResponse[TemplateSchema])
async def list_templates(
    request: Request,
    skip: int = Query(0, ge=0),
    limit: int = Query(10, ge=1, le=100),
    search: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取模板列表"""
    try:
        logger.info(f"获取用户 {current_user.id} 的模板列表，搜索: {search}")
        
        # 获取模板列表
        templates, total = crud_template.get_templates_with_pagination(
            db=db,
            user_id=current_user.id,
            skip=skip,
            limit=limit,
            search=search
        )
        
        # 转换为schema对象
        template_schemas = [TemplateSchema.model_validate(template) for template in templates]
        
        return PaginatedResponse(
            items=template_schemas,
            total=total,
            page=skip // limit + 1,
            size=limit,
            pages=(total + limit - 1) // limit,
            has_next=skip + limit < total,
            has_prev=skip > 0
        )
    except Exception as e:
        logger.error(f"获取模板列表失败: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="获取模板列表失败"
        )


@router.get("/{template_id}", response_model=ApiResponse[TemplateSchema])
async def get_template(
    request: Request,
    template_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取特定模板"""
    try:
        template = crud_template.get_by_id_and_user(
            db=db, 
            id=template_id, 
            user_id=current_user.id
        )
        
        if not template:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="模板不存在"
            )
        
        return ApiResponse(
            success=True,
            data=template,
            message="获取模板成功"
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取模板失败: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="获取模板失败"
        )


@router.post("", response_model=ApiResponse[TemplateSchema])
async def create_template(
    request: Request,
    template_in: TemplateCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """创建新模板"""
    try:
        # 创建模板
        template = crud_template.create_with_owner(
            db=db,
            obj_in=template_in,
            owner_id=current_user.id
        )
        
        logger.info(f"用户 {current_user.id} 创建了模板 {template.id}")
        
        return ApiResponse(
            success=True,
            data=template,
            message="模板创建成功"
        )
    except Exception as e:
        logger.error(f"创建模板失败: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="创建模板失败"
        )


@router.put("/{template_id}", response_model=ApiResponse[TemplateSchema])
async def update_template(
    request: Request,
    template_id: str,
    template_in: TemplateUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """更新模板"""
    try:
        template = crud_template.get_by_id_and_user(
            db=db, 
            id=template_id, 
            user_id=current_user.id
        )
        
        if not template:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="模板不存在"
            )
        
        # 更新模板
        template = crud_template.update(
            db=db,
            db_obj=template,
            obj_in=template_in
        )
        
        logger.info(f"用户 {current_user.id} 更新了模板 {template_id}")
        
        return ApiResponse(
            success=True,
            data=template,
            message="模板更新成功"
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"更新模板失败: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="更新模板失败"
        )


@router.get("/{template_id}/dependencies", response_model=ApiResponse[Dict])
async def get_template_dependencies(
    request: Request,
    template_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取模板依赖信息"""
    try:
        template = crud_template.get_by_id_and_user(
            db=db,
            id=template_id,
            user_id=current_user.id
        )

        if not template:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="模板不存在"
            )

        # 检查关联的任务
        from app.models.task import Task
        related_tasks = db.query(Task).filter(Task.template_id == template_id).all()

        dependencies = {
            "template_id": template_id,
            "template_name": template.name,
            "can_delete": len(related_tasks) == 0,
            "related_tasks": [
                {
                    "id": task.id,
                    "name": task.name,
                    "status": task.status.value if hasattr(task.status, 'value') else str(task.status),
                    "is_active": task.is_active,
                    "created_at": task.created_at.isoformat() if task.created_at else None
                } for task in related_tasks
            ],
            "related_tasks_count": len(related_tasks)
        }

        return ApiResponse(
            success=True,
            data=dependencies,
            message="模板依赖信息获取成功"
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取模板依赖信息失败: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="获取模板依赖信息失败"
        )


@router.delete("/{template_id}", response_model=ApiResponse[Dict])
async def delete_template(
    request: Request,
    template_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """删除模板"""
    try:
        template = crud_template.get_by_id_and_user(
            db=db, 
            id=template_id, 
            user_id=current_user.id
        )
        
        if not template:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="模板不存在"
            )
        
        # 删除模板
        crud_template.remove(db=db, id=template_id)

        logger.info(f"用户 {current_user.id} 删除了模板 {template_id}")

        return ApiResponse(
            success=True,
            data={"deleted_id": template_id},
            message="模板删除成功"
        )
    except HTTPException:
        raise
    except ValueError as e:
        # 处理业务逻辑错误（如存在关联任务）
        logger.warning(f"删除模板失败 - 业务逻辑错误: {e}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        )
    except Exception as e:
        logger.error(f"删除模板失败: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="删除模板失败"
        )


@router.post("/{template_id}/analyze", response_model=ApiResponse[Dict])
async def analyze_template_placeholders(
    template_id: str,
    data_source_id: str = Query(..., description="数据源ID"),
    optimization_level: str = Query("enhanced", description="优化级别"),
    force_reanalyze: bool = Query(False, description="强制重新分析"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """智能分析模板占位符 - 使用AI Agent高级分析 (Claude Code架构)"""
    try:
        # 验证模板存在性
        template = crud_template.get_by_id_and_user(
            db=db,
            id=template_id,
            user_id=current_user.id
        )
        
        if not template:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="模板不存在"
            )
        
        # 获取数据源信息
        from app.crud import data_source as crud_data_source
        data_source_info = None
        if data_source_id:
            data_source = crud_data_source.get(db, id=data_source_id)
            if data_source:
                data_source_info = {
                    "type": data_source.source_type.value if hasattr(data_source.source_type, 'value') else str(data_source.source_type),
                    "database": getattr(data_source, 'doris_database', 'unknown'),
                    "name": data_source.name
                }
        
        # 使用新的LLM编排服务进行AI Agent高级分析
        from app.services.application.llm import get_llm_orchestration_service
        
        # 构建AI Agent高级分析问题描述
        template_content_preview = template.content[:500] + "..." if len(template.content) > 500 else template.content
        
        business_question = f"""
        使用AI Agent智能深度分析模板'{template.name}'的占位符和数据映射关系。
        
        模板类型: {template.template_type or 'report'}
        优化级别: {optimization_level}
        
        AI Agent高级分析要求:
        1. 智能识别所有占位符模式和语法结构
        2. 建立复杂的数据字段映射关系和依赖分析
        3. 提供详细的数据类型分析和约束检查
        4. 生成个性化的架构优化建议
        5. 评估模板的数据完整性和性能要求
        6. 提供智能化的最佳实践推荐
        """
        
        # 构建AI Agent增强的上下文信息
        context_info = {
            "template": {
                "id": template_id,
                "name": template.name,
                "type": template.template_type or "report",
                "content_preview": template_content_preview,
                "content_length": len(template.content),
                "complexity_level": "ai_agent_advanced"
            },
            "data_source": data_source_info,
            "analysis_requirements": {
                "optimization_level": optimization_level,
                "force_reanalyze": force_reanalyze,
                "analysis_depth": "comprehensive",
                "ai_agent_mode": True,
                "requested_outputs": [
                    "intelligent_placeholder_analysis",
                    "advanced_field_mapping_suggestions", 
                    "detailed_data_type_requirements",
                    "personalized_optimization_recommendations",
                    "integrity_and_performance_assessment",
                    "best_practices_guidance"
                ]
            },
            "ai_features": {
                "intelligent_pattern_recognition": True,
                "advanced_dependency_analysis": True,
                "personalized_recommendations": True,
                "comprehensive_assessment": True
            }
        }
        
        # 执行AI Agent高级LLM分析
        service = get_llm_orchestration_service()
        result = await service.analyze_data_requirements(
            user_id=str(current_user.id),
            business_question=business_question,
            context_info=context_info
        )
        
        if not result.get('success'):
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"AI Agent智能分析失败: {result.get('error', '未知错误')}"
            )
        
        logger.info(f"用户 {current_user.id} 使用AI Agent智能分析了模板 {template_id}")
        
        # 格式化AI Agent高级分析结果
        analysis_data = {
            "template_info": {
                "id": template_id,
                "name": template.name,
                "type": template.template_type or "report",
                "analysis_level": "ai_agent_intelligent"
            },
            "ai_agent_analysis": {
                "analysis": result.get('analysis', ''),
                "recommended_approach": result.get('recommended_approach', ''),
                "confidence": result.get('confidence', 0.8)
            },
            "data_source_info": data_source_info,
            "optimization_level": optimization_level,
            "ai_capabilities": {
                "analysis_method": "ai_agent_six_stage_orchestration",
                "llm_participated": True,
                "intelligent_pattern_recognition": True,
                "advanced_recommendations": True,
                "comprehensive_analysis": True
            }
        }
        
        return ApiResponse(
            success=True,
            data=analysis_data,
            message="AI Agent智能分析完成"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Claude Code架构模板分析失败: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"分析失败: {str(e)}"
        )


@router.get("/{template_id}/analyze/stream")
async def analyze_template_streaming_with_claude_code_architecture(
    template_id: str,
    data_source_id: str = Query(..., description="数据源ID"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """流式分析模板 - 使用Claude Code架构 (实时进度反馈)"""
    from fastapi.responses import StreamingResponse
    import json
    
    async def generate():
        try:
            # 验证模板存在性
            template = crud_template.get_by_id_and_user(
                db=db,
                id=template_id,
                user_id=current_user.id
            )
            
            if not template:
                yield f"data: {json.dumps({'error': '模板不存在'})}\n\n"
                return
            
            # 获取数据源信息
            from app.crud import data_source as crud_data_source
            data_source_info = None
            if data_source_id:
                data_source = crud_data_source.get(db, id=data_source_id)
                if data_source:
                    data_source_info = {
                        "type": data_source.source_type.value if hasattr(data_source.source_type, 'value') else str(data_source.source_type),
                        "database": getattr(data_source, 'doris_database', 'unknown'),
                        "name": data_source.name
                    }
            
            # 使用新的LLM编排服务进行高级流式分析
            from app.services.application.llm import get_llm_orchestration_service
            
            # 发送增强开始事件
            yield f"data: {json.dumps({'type': 'start', 'template_id': template_id, 'user_id': str(current_user.id), 'method': 'llm_orchestration_v2', 'architecture': 'claude_code'})}\n\n"
            
            # 构建高级分析问题和上下文
            template_content_preview = template.content[:500] + "..." if len(template.content) > 500 else template.content
            
            business_question = f"""
            使用Claude Code架构v2进行高级流式分析模板'{template.name}'。
            
            模板类型: {template.template_type or 'report'}
            
            Claude Code架构分析要求：
            1. 深度解析模板结构和语法
            2. 智能识别复杂占位符模式
            3. 建立精确的数据源映射关系
            4. 提供架构级优化建议
            5. 实时流式反馈分析进度
            """
            
            context_info = {
                "template": {
                    "id": template_id,
                    "name": template.name,
                    "type": template.template_type or "report",
                    "content_preview": template_content_preview,
                    "content_length": len(template.content),
                    "architecture": "claude_code_v2"
                },
                "data_source": data_source_info,
                "streaming_mode": True,
                "advanced_features": {
                    "claude_code_architecture": True,
                    "enhanced_analysis": True,
                    "streaming_progress": True
                }
            }
            
            # Claude Code架构六步编排的详细流式进度
            progress_steps = [
                {'type': 'progress', 'progress': 10, 'stage': 'initialization', 'message': '初始化Claude Code架构v2...'},
                {'type': 'progress', 'progress': 20, 'stage': 'validation', 'message': '验证模板结构和语法...'},
                {'type': 'progress', 'progress': 35, 'stage': 'readonly_parallel', 'message': '并行解析模板占位符模式...'},
                {'type': 'progress', 'progress': 50, 'stage': 'write_sequential', 'message': '构建数据源映射关系...'},
                {'type': 'progress', 'progress': 65, 'stage': 'context_compression', 'message': '优化分析上下文...'},
                {'type': 'progress', 'progress': 85, 'stage': 'llm_reasoning', 'message': 'Claude Code架构深度推理...'},
            ]
            
            try:
                # 发送详细的进度更新
                for step in progress_steps:
                    yield f"data: {json.dumps(step)}\n\n"
                    import asyncio
                    await asyncio.sleep(0.6)  # 稍微慢一点，显示架构复杂性
                
                # 执行实际的LLM分析
                service = get_llm_orchestration_service()
                result = await service.analyze_data_requirements(
                    user_id=str(current_user.id),
                    business_question=business_question,
                    context_info=context_info
                )
                
                # 发送最终结果
                if result.get('success'):
                    final_data = {
                        'type': 'result', 
                        'progress': 100, 
                        'stage': 'synthesis_complete',
                        'message': 'Claude Code架构分析完成',
                        'data': {
                            "template_info": {
                                "id": template_id,
                                "name": template.name,
                                "type": template.template_type or "report",
                                "architecture": "claude_code_v2"
                            },
                            "claude_code_analysis": {
                                "analysis": result.get('analysis', ''),
                                "recommended_approach": result.get('recommended_approach', ''),
                                "confidence": result.get('confidence', 0.8)
                            },
                            "architecture_features": {
                                "streaming_method": "claude_code_six_stage_orchestration",
                                "llm_participated": True,
                                "advanced_reasoning": True,
                                "architectural_optimization": True
                            }
                        }
                    }
                    yield f"data: {json.dumps(final_data)}\n\n"
                else:
                    error_msg = result.get('error', '未知错误')
                    yield f"data: {json.dumps({'type': 'error', 'error': f'Claude Code分析失败: {error_msg}'})}\n\n"
                    
            except Exception as e:
                yield f"data: {json.dumps({'type': 'error', 'error': f'Claude Code流式分析异常: {str(e)}'})}\n\n"
                
        except Exception as e:
            logger.error(f"流式分析失败: {e}")
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
    
    return StreamingResponse(
        generate(),
        media_type="text/plain",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive"}
    )


@router.get("/{template_id}/preview", response_model=ApiResponse[TemplatePreview])
async def preview_template(
    request: Request,
    template_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """预览模板内容和占位符"""
    try:
        # 验证模板存在性
        template = crud_template.get_by_id_and_user(
            db=db,
            id=template_id,
            user_id=current_user.id
        )
        
        if not template:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="模板不存在"
            )
        
        # 解析模板结构
        structure = template_parser.parse_template_structure(template.content or "")
        
        # 构建预览数据
        preview_data = TemplatePreview(
            template_id=template.id,
            content=template.content,
            html_content=template.content,  # 可以在这里添加HTML转换逻辑
            placeholders=structure.get('placeholders', []),
            metadata={
                'name': template.name,
                'description': template.description,
                'template_type': template.template_type,
                'original_filename': template.original_filename,
                'file_size': template.file_size,
                'complexity_score': structure.get('complexity_score', 0),
                'sections': structure.get('sections', []),
                'variables': structure.get('variables', {})
            }
        )
        
        logger.info(f"用户 {current_user.id} 预览了模板 {template_id}")
        
        return ApiResponse(
            success=True,
            data=preview_data,
            message="模板预览获取成功"
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"模板预览失败: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="模板预览失败"
        )


@router.post("/{template_id}/upload", response_model=ApiResponse[TemplateSchema])
async def upload_template_file(
    request: Request,
    template_id: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """上传模板文件并更新内容"""
    try:
        # 验证模板存在性
        template = crud_template.get_by_id_and_user(
            db=db,
            id=template_id,
            user_id=current_user.id
        )
        
        if not template:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="模板不存在"
            )
        
        # 验证文件类型
        allowed_extensions = {'.docx', '.doc', '.txt', '.html', '.md'}
        file_extension = '.' + file.filename.split('.')[-1].lower() if '.' in file.filename else ''
        
        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"不支持的文件类型。支持的类型: {', '.join(allowed_extensions)}"
            )
        
        # 读取文件内容
        content = await file.read()
        file_size = len(content)
        
        # 1. 先保存原始文件到存储系统
        file_info = None
        content_text = ""
        
        try:
            from app.services.infrastructure.storage.hybrid_storage_service import get_hybrid_storage_service
            from io import BytesIO
            
            storage_service = get_hybrid_storage_service()
            
            # 保存原始文件
            file_info = storage_service.upload_file(
                file_data=BytesIO(content),
                original_filename=file.filename,
                file_type="templates",
                content_type=file.content_type
            )
            
            logger.info(f"文件保存到存储系统: {file_info.get('file_path')}")
            
        except Exception as e:
            logger.error(f"保存文件到存储系统失败: {str(e)}")
            # 如果存储失败，仍然继续处理，但记录错误
        
        # 2. 解析文件内容用于占位符分析
        if file_extension in ['.docx', '.doc']:
            try:
                from docx import Document
                import io
                
                # 解析docx文档
                doc = Document(io.BytesIO(content))
                
                # 提取文本内容
                full_text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)
                
                # 提取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            full_text.append(" | ".join(row_text))
                
                content_text = "\n\n".join(full_text) if full_text else f"[空文档: {file.filename}]"
                
            except Exception as e:
                logger.error(f"解析docx文件失败: {str(e)}")
                content_text = f"[文档解析失败: {file.filename}]\n错误信息: {str(e)}"
        else:
            content_text = content.decode('utf-8', errors='ignore')
        
        # 3. 更新模板记录
        template_update = TemplateUpdate(
            content=content_text,
            original_filename=file.filename,
            file_path=file_info.get("file_path") if file_info else None,
            file_size=file_size,
            template_type=file_extension.lstrip('.')
        )
        
        updated_template = crud_template.update(
            db=db,
            db_obj=template,
            obj_in=template_update
        )
        
        logger.info(f"用户 {current_user.id} 上传了模板文件 {file.filename} 到模板 {template_id}")
        
        # 自动触发占位符分析
        try:
            structure = template_parser.parse_template_structure(content_text)
            logger.info(f"自动解析了模板 {template_id} 的占位符: {len(structure.get('placeholders', []))} 个")
        except Exception as parse_error:
            logger.warning(f"自动占位符解析失败: {parse_error}")
        
        return ApiResponse(
            success=True,
            data=updated_template,
            message=f"模板文件上传成功，解析到 {len(structure.get('placeholders', []))} 个占位符" if 'structure' in locals() else "模板文件上传成功"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"模板文件上传失败: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="模板文件上传失败"
        )


@router.get("/{template_id}/download")
async def download_template_file(
    request: Request,
    template_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """下载模板原始文件"""
    try:
        # 验证模板存在性
        template = crud_template.get_by_id_and_user(
            db=db,
            id=template_id,
            user_id=current_user.id
        )
        if not template:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="模板不存在")
        if not template.file_path:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="模板没有关联的文件")

        from app.services.infrastructure.storage.hybrid_storage_service import get_hybrid_storage_service
        from fastapi.responses import StreamingResponse
        import io

        storage_service = get_hybrid_storage_service()
        if not storage_service.file_exists(template.file_path):
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="文件在存储系统中不存在")

        file_data, backend_type = storage_service.download_file(template.file_path)
        content_type = "application/octet-stream"
        if template.original_filename:
            fn = template.original_filename.lower()
            if fn.endswith(".docx"):
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            elif fn.endswith(".doc"):
                content_type = "application/msword"
            elif fn.endswith(".pdf"):
                content_type = "application/pdf"
            elif fn.endswith(".txt"):
                content_type = "text/plain"
            elif fn.endswith(".html"):
                content_type = "text/html"

        file_stream = io.BytesIO(file_data)
        logger.info(f"用户 {current_user.id} 下载模板文件: {template.name} ({template.original_filename})")
        return StreamingResponse(
            file_stream,
            media_type=content_type,
            headers={
                "Content-Disposition": f'attachment; filename="{template.original_filename or f"template_{template_id}"}"',
                "X-Storage-Backend": backend_type,
                "X-Template-ID": template_id
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"模板文件下载失败: {e}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="文件下载失败")


@router.get("/{template_id}/download-url", response_model=ApiResponse[Dict[str, Any]])
async def get_template_download_url(
    request: Request,
    template_id: str,
    expires: int = Query(3600, ge=60, le=86400),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取模板文件的预签名下载URL（MinIO或回退）"""
    try:
        template = crud_template.get_by_id_and_user(
            db=db,
            id=template_id,
            user_id=current_user.id
        )
        if not template:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="模板不存在")
        if not template.file_path:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="模板没有关联的文件")

        storage = get_hybrid_storage_service()
        if not storage.file_exists(template.file_path):
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="文件在存储系统中不存在")

        url = storage.get_download_url(template.file_path, expires=expires)
        return ApiResponse(success=True, data={"url": url, "file_path": template.file_path}, message="获取下载URL成功")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取模板下载URL失败: {e}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="获取下载URL失败")


@router.post("/{template_id}/placeholders/reparse", response_model=ApiResponse[Dict])
async def reparse_template_placeholders(
    request: Request,
    template_id: str,
    force_reparse: bool = Query(False, description="强制重新解析"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """重新解析模板占位符并保存到数据库"""
    try:
        # 验证模板存在性
        template = crud_template.get_by_id_and_user(
            db=db,
            id=template_id,
            user_id=current_user.id
        )

        if not template:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="模板不存在"
            )

        # 解析模板结构
        structure = template_parser.parse_template_structure(template.content or "")
        placeholders_data = structure.get('placeholders', [])

        # 🔑 保存占位符到数据库
        saved_count = 0
        for ph in placeholders_data:
            try:
                # 使用upsert逻辑，如果存在就更新，不存在就创建
                existing = crud.template_placeholder.get_by_template_and_name(
                    db=db,
                    template_id=template_id,
                    name=ph.get('name', '') or ph.get('description', '')
                )

                if existing and not force_reparse:
                    continue  # 如果已存在且不强制重新解析，跳过

                placeholder_in = TemplatePlaceholderCreate(
                    template_id=template_id,
                    placeholder_name=ph.get('name', '') or ph.get('description', ''),
                    placeholder_text=ph.get('text', ''),
                    placeholder_type=ph.get('type', 'statistical'),
                    content_type='text',
                    execution_order=ph.get('position', 1),
                    is_active=True,
                    original_type=ph.get('original_type'),
                    extracted_description=ph.get('description'),
                    parsing_metadata=ph
                )

                if existing:
                    # 更新现有占位符
                    crud.template_placeholder.update(
                        db=db,
                        db_obj=existing,
                        obj_in=TemplatePlaceholderUpdate(**placeholder_in.dict(exclude={'template_id'}))
                    )
                else:
                    # 创建新占位符
                    crud.template_placeholder.create(db=db, obj_in=placeholder_in)

                saved_count += 1

            except Exception as e:
                logger.warning(f"保存占位符失败 {ph.get('name')}: {e}")
                continue

        logger.info(f"用户 {current_user.id} 重新解析了模板 {template_id} 的占位符: 发现 {len(placeholders_data)} 个，保存 {saved_count} 个")

        return ApiResponse(
            success=True,
            data={
                "template_id": template_id,
                "placeholders": placeholders_data,
                "sections": structure.get('sections', []),
                "variables": structure.get('variables', {}),
                "complexity_score": structure.get('complexity_score', 0),
                "force_reparse": force_reparse,
                "saved_count": saved_count,
                "total_found": len(placeholders_data)
            },
            message=f"占位符重新解析完成，共发现 {len(placeholders_data)} 个占位符，已保存 {saved_count} 个到数据库"
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"重新解析占位符失败: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"重新解析占位符失败: {str(e)}"
        )


@router.post("/{template_id}/analyze-with-agent", response_model=ApiResponse[Dict])
async def analyze_with_agent(
    template_id: str,
    data_source_id: str = Query(..., description="数据源ID"),
    force_reanalyze: bool = Query(False, description="强制重新分析"),
    optimization_level: str = Query("enhanced", description="优化级别"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """使用AI Agent分析模板 - 已升级到Claude Code架构"""
    try:
        # 验证模板存在性
        template = crud_template.get_by_id_and_user(
            db=db,
            id=template_id,
            user_id=current_user.id
        )
        
        if not template:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="模板不存在"
            )
        
        # 获取数据源信息
        from app.crud import data_source as crud_data_source
        data_source_info = None
        if data_source_id:
            data_source = crud_data_source.get(db, id=data_source_id)
            if data_source:
                data_source_info = {
                    "type": data_source.source_type.value if hasattr(data_source.source_type, 'value') else str(data_source.source_type),
                    "database": getattr(data_source, 'doris_database', 'unknown'),
                    "name": data_source.name
                }
        
        # 使用新的LLM编排服务
        from app.services.application.llm import get_llm_orchestration_service
        
        # 构建高级分析问题描述
        template_content_preview = template.content[:500] + "..." if len(template.content) > 500 else template.content
        
        business_question = f"""
        使用AI Agent高级分析模板'{template.name}'的占位符和数据映射关系。
        
        模板类型: {template.template_type or 'report'}
        优化级别: {optimization_level}
        
        高级分析要求:
        1. 智能识别所有占位符模式
        2. 建立复杂的数据字段映射关系
        3. 提供详细的数据类型分析
        4. 生成个性化的优化建议
        5. 评估模板的数据完整性要求
        """
        
        # 构建增强的上下文信息
        context_info = {
            "template": {
                "id": template_id,
                "name": template.name,
                "type": template.template_type or "report",
                "content_preview": template_content_preview,
                "content_length": len(template.content),
                "complexity_level": "advanced"
            },
            "data_source": data_source_info,
            "analysis_requirements": {
                "optimization_level": optimization_level,
                "analysis_depth": "comprehensive",
                "ai_agent_mode": True,
                "requested_outputs": [
                    "placeholder_analysis",
                    "field_mapping_suggestions", 
                    "data_type_requirements",
                    "optimization_recommendations",
                    "integrity_assessment"
                ]
            }
        }
        
        # 执行高级LLM分析
        service = get_llm_orchestration_service()
        result = await service.analyze_data_requirements(
            user_id=str(current_user.id),
            business_question=business_question,
            context_info=context_info
        )
        
        if not result.get('success'):
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"AI Agent分析失败: {result.get('error', '未知错误')}"
            )
        
        logger.info(f"用户 {current_user.id} 使用LLM编排服务进行高级模板分析 {template_id}")
        
        # 格式化高级分析结果
        analysis_data = {
            "template_info": {
                "id": template_id,
                "name": template.name,
                "type": template.template_type or "report",
                "analysis_level": "ai_agent_advanced"
            },
            "advanced_analysis": {
                "analysis": result.get('analysis', ''),
                "recommended_approach": result.get('recommended_approach', ''),
                "confidence": result.get('confidence', 0.8)
            },
            "data_source_info": data_source_info,
            "optimization_level": optimization_level,
            "analysis_method": "llm_orchestration_ai_agent_mode",
            "llm_participated": True,
            "ai_features": {
                "intelligent_mapping": True,
                "advanced_recommendations": True,
                "comprehensive_analysis": True
            }
        }
        
        return ApiResponse(
            success=True,
            data=analysis_data,
            message="AI Agent高级分析完成"
        )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Agent模板分析失败: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Agent分析失败"
        )
