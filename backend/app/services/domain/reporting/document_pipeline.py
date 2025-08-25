"""
Word文档生成流水线系统

基于模板的稳定文档生成流水线，支持：
1. 模板解析和验证
2. 占位符替换
3. 格式化和样式应用
4. Word文档生成和导出
5. 文档质量检查
"""

import asyncio
import json
import logging
import os
import re
import time
from dataclasses import dataclass, field
from enum import Enum
from typing import Any, Dict, List, Optional, Union
from datetime import datetime
import uuid
import base64

# 尝试导入python-docx，如果没有则使用模拟
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)


class DocumentFormat(Enum):
    """文档格式枚举"""
    DOCX = "docx"
    PDF = "pdf"
    HTML = "html"
    MARKDOWN = "md"


class ContentType(Enum):
    """内容类型枚举"""
    TEXT = "text"
    TABLE = "table"
    CHART = "chart"
    IMAGE = "image"
    LIST = "list"
    PARAGRAPH = "paragraph"


@dataclass
class PlaceholderInfo:
    """占位符信息"""
    name: str
    type: str
    description: str
    content_type: ContentType
    required: bool = True
    default_value: str = ""
    format_rules: Dict[str, Any] = field(default_factory=dict)


@dataclass
class DocumentTemplate:
    """文档模板"""
    template_id: str
    name: str
    content: str
    placeholders: List[PlaceholderInfo]
    styles: Dict[str, Any] = field(default_factory=dict)
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ProcessedContent:
    """处理后的内容"""
    placeholder_name: str
    content_type: ContentType
    raw_content: Any
    formatted_content: str
    style_info: Dict[str, Any] = field(default_factory=dict)
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class DocumentGenerationResult:
    """文档生成结果"""
    success: bool
    document_path: str
    document_format: DocumentFormat
    generation_time: float
    placeholder_count: int
    processed_placeholders: int
    failed_placeholders: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    error_message: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


class TemplateParser:
    """模板解析器"""
    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
        # 支持多种占位符格式：
        # 1. {{type:description}} - 标准格式
        # 2. {{variable}} - 简单变量
        # 3. {variable} - 单花括号格式
        # 4. {{@variable=value}} - 带默认值的变量
        # 5. 【中文占位符】- 中文格式
        # 6. [占位符] - 方括号格式
        self.placeholder_patterns = [
            r'\{\{([^:}]+):([^}]+)\}\}',  # {{type:description}}
            r'\{\{@(\w+)\s*=\s*([^}]+)\}\}',  # {{@variable=value}}
            r'\{\{([^}]+)\}\}',  # {{variable}}
            r'\{([^}]+)\}',  # {variable}
            r'【([^】]+)】',  # 【中文占位符】
            r'\[([^\]]+)\]'  # [占位符]
        ]
    
    def parse_template(self, template_content: str) -> DocumentTemplate:
        """解析模板内容"""
        template_id = str(uuid.uuid4())
        
        # 提取占位符
        placeholders = self._extract_placeholders(template_content)
        
        # 分析模板结构
        styles = self._analyze_template_styles(template_content)
        
        return DocumentTemplate(
            template_id=template_id,
            name="Parsed Template",
            content=template_content,
            placeholders=placeholders,
            styles=styles,
            metadata={
                "parsed_at": datetime.now().isoformat(),
                "placeholder_count": len(placeholders)
            }
        )
    
    def extract_placeholders(self, content: str) -> List[Dict[str, Any]]:
        """提取占位符（公共方法）"""
        self.logger.info(f"开始解析模板占位符，内容长度: {len(content)}")
        
        # 首先尝试检测是否为二进制Word文档
        text_content = self._extract_text_from_content(content)
        self.logger.info(f"提取文本内容，长度: {len(text_content)}")
        
        placeholders_info = self._extract_placeholders_improved(text_content)
        self.logger.info(f"提取到 {len(placeholders_info)} 个占位符")
        
        # 转换为字典格式以兼容现有代码
        result = []
        for placeholder in placeholders_info:
            result.append({
                "name": placeholder.name,
                "type": placeholder.type,
                "description": placeholder.description,
                "content_type": placeholder.content_type.value if hasattr(placeholder.content_type, 'value') else str(placeholder.content_type),
                "required": placeholder.required
            })
        
        return result
    
    def _extract_text_from_content(self, content: str) -> str:
        """从内容中提取文本（支持二进制Word文档）"""
        # 检查是否为二进制数据（Word文档的十六进制表示）
        if self._is_binary_content(content):
            self.logger.info("检测到二进制Word文档，尝试解析...")
            try:
                # 清理十六进制内容（去除空格和换行符）
                clean_hex = content.replace(' ', '').replace('\n', '').replace('\r', '')
                
                # 验证十六进制长度
                if len(clean_hex) % 2 != 0:
                    self.logger.warning("十六进制内容长度不正确")
                    return content
                
                # 尝试将十六进制字符串转换为二进制数据
                binary_data = bytes.fromhex(clean_hex)
                
                # 检查是否为Word文档（ZIP格式）
                if binary_data.startswith(b'PK'):
                    # 进一步验证是否为DOCX文件
                    if b'word/' in binary_data or b'[Content_Types].xml' in binary_data:
                        return self._extract_text_from_word_doc(binary_data)
                    else:
                        self.logger.warning("ZIP文件但非DOCX格式")
                        return ""
                else:
                    self.logger.warning("无法识别的二进制格式")
                    return content
            except ValueError as e:
                self.logger.error(f"十六进制解码失败: {e}")
                return content
            except Exception as e:
                self.logger.error(f"解析二进制文档失败: {e}")
                return content
        else:
            # 纯文本内容，直接返回
            return content
    
    def _is_binary_content(self, content: str) -> bool:
        """检查内容是否为二进制数据的十六进制表示"""
        if not content or len(content) < 20:
            return False
        
        # 清理内容（去除空格和换行符）
        clean_content = content.replace(' ', '').replace('\n', '').replace('\r', '')
        
        # 检查长度是否为偶数（十六进制必须是偶数长度）
        if len(clean_content) % 2 != 0:
            return False
        
        # 检查是否只包含十六进制字符
        if not all(c in '0123456789ABCDEFabcdef' for c in clean_content):
            return False
        
        # 检查是否以常见的文档格式魔术字节开头
        hex_prefix = clean_content[:20].lower()
        # PK（ZIP/Word文档）= 504b
        # PDF = 255044462d
        # Microsoft Office documents often start with 504b (ZIP format)
        if hex_prefix.startswith('504b') or hex_prefix.startswith('255044462d'):
            return True
        
        # 如果内容很长且全是十六进制字符，也可能是二进制文件
        if len(clean_content) > 1000:
            return True
            
        return False
    
    def _extract_text_from_word_doc(self, binary_data: bytes) -> str:
        """从Word文档二进制数据中提取文本"""
        text_parts = []
        temp_file_path = None
        
        try:
            import tempfile
            import os
            
            # 保存二进制数据到临时文件
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(binary_data)
                temp_file_path = temp_file.name
            
            try:
                if HAS_DOCX:
                    # 使用python-docx解析
                    from docx import Document
                    doc = Document(temp_file_path)
                    
                    # 提取所有段落文本
                    for paragraph in doc.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            text_parts.append(text)
                    
                    # 提取表格文本
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text = cell.text.strip()
                                if text:
                                    text_parts.append(text)
                    
                    # 提取页眉页脚（如果可能）
                    try:
                        for section in doc.sections:
                            # 页眉
                            header = section.header
                            for paragraph in header.paragraphs:
                                text = paragraph.text.strip()
                                if text:
                                    text_parts.append(text)
                            
                            # 页脚
                            footer = section.footer
                            for paragraph in footer.paragraphs:
                                text = paragraph.text.strip()
                                if text:
                                    text_parts.append(text)
                    except Exception as header_footer_error:
                        self.logger.debug(f"提取页眉页脚失败: {header_footer_error}")
                    
                    extracted_text = '\n'.join(text_parts)
                    self.logger.info(f"成功从Word文档提取文本，长度: {len(extracted_text)}")
                    return extracted_text
                
                else:
                    self.logger.warning("python-docx不可用，尝试备用方法")
                    # 备用方法：作为ZIP文件读取
                    return self._extract_text_fallback(temp_file_path)
            
            finally:
                # 清理临时文件
                if temp_file_path and os.path.exists(temp_file_path):
                    try:
                        os.unlink(temp_file_path)
                    except Exception as cleanup_error:
                        self.logger.warning(f"删除临时文件失败: {cleanup_error}")
                    
        except Exception as e:
            self.logger.error(f"解析Word文档失败: {e}")
            return ""
    
    def _extract_text_fallback(self, file_path: str) -> str:
        """备用文本提取方法（当python-docx不可用时）"""
        try:
            import zipfile
            text_parts = []
            
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                # 查找document.xml文件
                if 'word/document.xml' in zip_file.namelist():
                    document_xml = zip_file.read('word/document.xml')
                    # 简单的XML文本提取
                    xml_text = document_xml.decode('utf-8', errors='ignore')
                    # 使用正则表达式提取文本内容
                    text_matches = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml_text)
                    text_parts.extend(text_matches)
            
            extracted_text = ' '.join(text_parts)
            self.logger.info(f"备用方法提取文本，长度: {len(extracted_text)}")
            return extracted_text
        
        except Exception as e:
            self.logger.error(f"备用文本提取方法失败: {e}")
            return ""
    
    def _extract_placeholders(self, content: str) -> List[PlaceholderInfo]:
        """提取占位符（私有方法）"""
        placeholders = []
        seen_placeholders = set()  # 用于去重
        
        self.logger.info(f"从内容中提取占位符，内容长度: {len(content)}")
        
        # 使用多个正则表达式模式
        for i, pattern in enumerate(self.placeholder_patterns):
            try:
                matches = re.finditer(pattern, content, re.IGNORECASE | re.MULTILINE)
                
                for match in matches:
                    placeholder_info = self._parse_placeholder_match(match, i)
                    
                    if placeholder_info and placeholder_info.name:
                        # 去重处理
                        placeholder_key = f"{placeholder_info.name}_{placeholder_info.type}"
                        if placeholder_key not in seen_placeholders:
                            seen_placeholders.add(placeholder_key)
                            placeholders.append(placeholder_info)
                            self.logger.debug(f"提取占位符: {placeholder_info.name} - {placeholder_info.description}")
            
            except Exception as e:
                self.logger.warning(f"模式 {i} 匹配失败: {e}")
        
        self.logger.info(f"提取到 {len(placeholders)} 个唯一占位符")
        return placeholders
    
    def _parse_placeholder_match(self, match, pattern_index: int) -> Optional[PlaceholderInfo]:
        """解析正则匹配结果为占位符信息"""
        groups = match.groups()
        full_text = match.group(0)
        
        placeholder_name = ""
        placeholder_description = ""
        placeholder_type = "text"
        
        try:
            if pattern_index == 0:  # {{type:description}} 格式
                if len(groups) >= 2 and groups[0] and groups[1]:
                    placeholder_type = groups[0].strip()
                    placeholder_description = groups[1].strip()
                    placeholder_name = f"{placeholder_type}_{hash(placeholder_description) % 1000}"
            
            elif pattern_index == 1:  # {{@variable=value}} 格式
                if len(groups) >= 2 and groups[0] and groups[1]:
                    placeholder_name = groups[0].strip()
                    placeholder_description = f"默认值: {groups[1].strip()}"
                    placeholder_type = "variable"
            
            elif pattern_index == 2:  # {{variable}} 格式
                if len(groups) >= 1 and groups[0]:
                    placeholder_name = groups[0].strip()
                    placeholder_description = f"变量: {placeholder_name}"
            
            elif pattern_index == 3:  # {variable} 格式
                if len(groups) >= 1 and groups[0]:
                    placeholder_name = groups[0].strip()
                    placeholder_description = f"简单变量: {placeholder_name}"
            
            elif pattern_index == 4:  # 【中文占位符】格式
                if len(groups) >= 1 and groups[0]:
                    placeholder_name = groups[0].strip()
                    placeholder_description = f"中文占位符: {placeholder_name}"
                    placeholder_type = "chinese"
            
            elif pattern_index == 5:  # [占位符] 格式
                if len(groups) >= 1 and groups[0]:
                    placeholder_name = groups[0].strip()
                    placeholder_description = f"方括号占位符: {placeholder_name}"
            
            # 验证占位符名称
            if not placeholder_name or len(placeholder_name) < 1:
                return None
            
            # 推断更精确的类型
            inferred_type = self._infer_placeholder_type(placeholder_name)
            if inferred_type != "text":
                placeholder_type = inferred_type
            
            # 推断内容类型
            content_type = self._infer_content_type(placeholder_name)
            
            return PlaceholderInfo(
                name=placeholder_name,
                type=placeholder_type,
                description=placeholder_description,
                content_type=content_type,
                required=True
            )
        
        except Exception as e:
            self.logger.warning(f"解析占位符匹配失败: {e}")
            return None
    
    def _infer_content_type(self, placeholder_name: str) -> ContentType:
        """推断内容类型"""
        name_lower = placeholder_name.lower()
        
        if any(keyword in name_lower for keyword in ['图表', 'chart', '图', '可视化']):
            return ContentType.CHART
        elif any(keyword in name_lower for keyword in ['表格', 'table', '列表', 'list']):
            return ContentType.TABLE
        elif any(keyword in name_lower for keyword in ['图片', 'image', '照片']):
            return ContentType.IMAGE
        elif any(keyword in name_lower for keyword in ['清单', '项目', 'item']):
            return ContentType.LIST
        else:
            return ContentType.TEXT
    
    def _infer_placeholder_type(self, placeholder_name: str) -> str:
        """推断占位符类型"""
        name_lower = placeholder_name.lower()
        
        if any(keyword in name_lower for keyword in ['统计', '数量', '总数', '计算']):
            return "statistic"
        elif any(keyword in name_lower for keyword in ['分析', '洞察', '趋势']):
            return "analysis"
        elif any(keyword in name_lower for keyword in ['图表', 'chart']):
            return "chart"
        else:
            return "text"
    
    def _extract_placeholders_improved(self, content: str) -> List[PlaceholderInfo]:
        """改进的占位符提取方法 - 解决重复和过度匹配问题"""
        if not content:
            return []
        
        placeholders = []
        seen_placeholders = set()  # 用于去重
        
        # 定义优先级排序的正则表达式
        patterns = [
            # 高优先级：具体格式的占位符
            (r'\{\{([^:}]+):([^}]+)\}\}', 'typed_placeholder'),  # {{type:description}}
            (r'\{\{@(\w+)\s*=\s*([^}]+)\}\}', 'default_value'),  # {{@var=value}}
            
            # 中优先级：简单双括号占位符  
            (r'\{\{([^{}]+)\}\}', 'simple_double'),  # {{variable}}
            
            # 低优先级：单括号和其他格式（避免过度匹配）
            (r'\{([^{}\[\]【】]+)\}', 'simple_single'),  # {variable}
            (r'【([^】]+)】', 'chinese_bracket'),  # 【中文】
            (r'\[([^\[\]{}【】]+)\]', 'square_bracket'),  # [variable]
        ]
        
        # 记录已匹配的位置，避免重复匹配
        matched_positions = set()
        
        for pattern, pattern_type in patterns:
            try:
                matches = list(re.finditer(pattern, content, re.IGNORECASE | re.MULTILINE))
                
                for match in matches:
                    start, end = match.span()
                    
                    # 检查是否与已匹配的位置重叠
                    if any(start < pos[1] and end > pos[0] for pos in matched_positions):
                        continue
                    
                    placeholder_info = self._parse_improved_match(match, pattern_type)
                    
                    if placeholder_info and self._is_valid_placeholder(placeholder_info):
                        # 更强的去重逻辑
                        placeholder_key = self._get_placeholder_key(placeholder_info)
                        
                        if placeholder_key not in seen_placeholders:
                            seen_placeholders.add(placeholder_key)
                            matched_positions.add((start, end))
                            placeholders.append(placeholder_info)
                        else:
                            # 如果已存在，选择更好的版本（优先带描述的格式）
                            existing_idx = next((i for i, p in enumerate(placeholders) 
                                               if self._get_placeholder_key(p) == placeholder_key), None)
                            
                            if existing_idx is not None:
                                existing = placeholders[existing_idx]
                                if self._is_better_placeholder(placeholder_info, existing):
                                    placeholders[existing_idx] = placeholder_info
                            
            except Exception as e:
                self.logger.warning(f"模式匹配失败 {pattern_type}: {e}")
        
        self.logger.info(f"改进提取器提取到 {len(placeholders)} 个唯一占位符")
        return placeholders
    
    def _parse_improved_match(self, match, pattern_type: str) -> Optional[PlaceholderInfo]:
        """解析改进的匹配结果"""
        groups = match.groups()
        full_text = match.group(0)
        
        if pattern_type == 'typed_placeholder':
            # {{type:description}}
            if len(groups) >= 2:
                placeholder_type = groups[0].strip()
                description = groups[1].strip()
                
                # 验证内容质量
                if self._is_meaningful_content(placeholder_type) and self._is_meaningful_content(description):
                    return PlaceholderInfo(
                        name=f"{placeholder_type}:{description}",
                        type=self._normalize_type(placeholder_type),
                        description=description,
                        content_type=self._infer_content_type(placeholder_type)
                    )
                    
        elif pattern_type == 'default_value':
            # {{@var=value}}
            if len(groups) >= 2:
                var_name = groups[0].strip()
                default_val = groups[1].strip()
                
                if self._is_meaningful_content(var_name):
                    return PlaceholderInfo(
                        name=var_name,
                        type="variable",
                        description=f"变量，默认值: {default_val}",
                        content_type=ContentType.TEXT,
                        default_value=default_val
                    )
                    
        elif pattern_type == 'simple_double':
            # {{variable}}
            if len(groups) >= 1:
                var_name = groups[0].strip()
                
                if self._is_meaningful_content(var_name) and ':' in var_name:
                    # 尝试解析为 type:description 格式
                    parts = var_name.split(':', 1)
                    if len(parts) == 2 and self._is_meaningful_content(parts[0]) and self._is_meaningful_content(parts[1]):
                        return PlaceholderInfo(
                            name=var_name,
                            type=self._normalize_type(parts[0]),
                            description=parts[1].strip(),
                            content_type=self._infer_content_type(parts[0])
                        )
                elif self._is_meaningful_content(var_name):
                    return PlaceholderInfo(
                        name=var_name,
                        type=self._infer_placeholder_type(var_name),
                        description=var_name,
                        content_type=ContentType.TEXT
                    )
                    
        elif pattern_type in ['simple_single', 'chinese_bracket', 'square_bracket']:
            # 单括号、中文括号、方括号格式
            if len(groups) >= 1:
                var_name = groups[0].strip()
                
                # 只有内容有意义且不太长才提取
                if (self._is_meaningful_content(var_name) and 
                    len(var_name) <= 50 and  # 限制长度
                    not self._contains_common_text_patterns(var_name)):
                    
                    return PlaceholderInfo(
                        name=var_name,
                        type=self._infer_placeholder_type(var_name),
                        description=f"简单占位符: {var_name}",
                        content_type=ContentType.TEXT
                    )
        
        return None
    
    def _is_valid_placeholder(self, placeholder: PlaceholderInfo) -> bool:
        """验证占位符是否有效"""
        if not placeholder or not placeholder.name:
            return False
        
        # 排除过短或过长的占位符
        if len(placeholder.name) < 2 or len(placeholder.name) > 100:
            return False
        
        # 排除纯数字或特殊字符
        if placeholder.name.isdigit() or not any(c.isalnum() or c in '：:_-' for c in placeholder.name):
            return False
        
        # 排除明显的文档结构内容
        exclude_patterns = [
            r'^\s*第[一二三四五六七八九十\d]+[章节部分]\s*$',  # 章节标题
            r'^\s*\d+\.\d+\s*$',  # 纯数字编号
            r'^\s*[第\d]+[页条项]\s*$',  # 页码等
        ]
        
        for pattern in exclude_patterns:
            if re.match(pattern, placeholder.name):
                return False
        
        return True
    
    def _is_meaningful_content(self, text: str) -> bool:
        """检查文本是否有意义"""
        if not text or len(text.strip()) < 2:
            return False
        
        # 排除纯数字、纯符号
        if text.isdigit() or not any(c.isalnum() for c in text):
            return False
        
        return True
    
    def _contains_common_text_patterns(self, text: str) -> bool:
        """检查是否包含常见文本模式（非占位符内容）"""
        text_patterns = [
            r'二、数据图表',  # 文档标题
            r'第[一二三四五六七八九十\d]+[章节]',  # 章节标题
            r'[\d]+\.\d+',  # 编号
        ]
        
        for pattern in text_patterns:
            if re.search(pattern, text):
                return True
        
        return False
    
    def _normalize_type(self, type_str: str) -> str:
        """标准化类型名称"""
        type_mapping = {
            '统计': 'statistic', '数量': 'statistic', '计算': 'statistic',
            '图表': 'chart', 'chart': 'chart',
            '表格': 'table', 'table': 'table',
            '分析': 'analysis', 'analysis': 'analysis',
            '时间': 'datetime', '日期': 'datetime', '周期': 'datetime',
            '区域': 'text', '地区': 'text', '标题': 'text',
            '变量': 'variable', 'variable': 'variable'
        }
        
        return type_mapping.get(type_str.lower(), 'text')
    
    def _infer_content_type(self, type_str: str) -> ContentType:
        """推断内容类型"""
        if any(keyword in type_str.lower() for keyword in ['图表', 'chart']):
            return ContentType.CHART
        elif any(keyword in type_str.lower() for keyword in ['表格', 'table']):
            return ContentType.TABLE
        else:
            return ContentType.TEXT
    
    def _get_placeholder_key(self, placeholder: PlaceholderInfo) -> str:
        """生成占位符的唯一标识键 - 用于去重"""
        normalized_name = placeholder.name.lower()
        
        # 移除数字后缀 (如 _955, _87)
        normalized_name = re.sub(r'[_\d]+$', '', normalized_name)
        
        # 处理类型:描述格式，提取核心语义
        if ':' in normalized_name:
            parts = normalized_name.split(':', 1)
            type_part = parts[0].strip()
            desc_part = parts[1].strip()
        else:
            # 对于没有冒号的占位符，尝试从名称推断语义
            type_part = ""
            desc_part = normalized_name
            
            # 从名称中提取类型信息
            if desc_part.startswith('周期') or '年份' in desc_part:
                type_part = '周期'
                # 像 "周期_955" 这种应与 "周期:报告年份" 归并
                if re.fullmatch(r'周期[_\d]*', placeholder.name):
                    desc_part = '年份'
                else:
                    desc_part = '年份' if '年份' in desc_part else desc_part.replace('周期', '').strip('_')
            elif desc_part.startswith('区域') or '地区' in desc_part:
                type_part = '区域'  
                # 像 "区域_87" 这种应与 "区域:地区名称" 归并
                if re.fullmatch(r'区域[_\d]*', placeholder.name):
                    desc_part = '地区'
                else:
                    desc_part = '地区' if '地区' in desc_part else desc_part.replace('区域', '').strip('_')
            elif desc_part.startswith('统计') or '总数' in desc_part or '件数' in desc_part:
                type_part = '统计'
                # 像 "统计_627" 这种应与 "统计:总投诉件数" 归并
                if re.fullmatch(r'统计[_\d]*', placeholder.name):
                    desc_part = '总数'
                else:
                    desc_part = '总数' if ('总数' in desc_part or '件数' in desc_part) else desc_part.replace('统计', '').strip('_')
        
        # 标准化描述部分
        desc_mapping = {
            '报告年份': '年份',
            '地区名称': '地区', 
            '统计开始日期': '开始日期',
            '统计结束日期': '结束日期',
            '总投诉件数': '总数',
            '投诉趋势折线图': '趋势图'
        }
        
        desc_part = desc_mapping.get(desc_part, desc_part)
        
        # 生成标准化的键
        if type_part:
            semantic_key = f"{type_part}:{desc_part}"
        else:
            semantic_key = desc_part
            
        return semantic_key
    
    def _is_better_placeholder(self, new_placeholder: PlaceholderInfo, existing_placeholder: PlaceholderInfo) -> bool:
        """判断新占位符是否比现有占位符更好"""
        # 1. 优先选择类型:描述格式的占位符
        if ':' in new_placeholder.name and ':' not in existing_placeholder.name:
            return True
        elif ':' not in new_placeholder.name and ':' in existing_placeholder.name:
            return False
        
        # 2. 优先选择描述更具体的占位符
        if len(new_placeholder.description) > len(existing_placeholder.description):
            return True
        elif len(new_placeholder.description) < len(existing_placeholder.description):
            return False
        
        # 3. 优先选择不带数字后缀的占位符
        if not re.search(r'_\d+$', new_placeholder.name) and re.search(r'_\d+$', existing_placeholder.name):
            return True
        elif re.search(r'_\d+$', new_placeholder.name) and not re.search(r'_\d+$', existing_placeholder.name):
            return False
        
        # 4. 优先选择内容类型更精确的占位符
        content_type_priority = {
            'chart': 3,
            'table': 2, 
            'text': 1
        }
        
        new_priority = content_type_priority.get(new_placeholder.content_type.value if hasattr(new_placeholder.content_type, 'value') else str(new_placeholder.content_type), 0)
        existing_priority = content_type_priority.get(existing_placeholder.content_type.value if hasattr(existing_placeholder.content_type, 'value') else str(existing_placeholder.content_type), 0)
        
        return new_priority > existing_priority
    
    def _analyze_template_styles(self, content: str) -> Dict[str, Any]:
        """分析模板样式"""
        styles = {
            "default_font": "Arial",
            "default_size": 12,
            "heading_styles": {},
            "paragraph_styles": {},
            "table_styles": {}
        }
        
        # 分析标题样式
        heading_pattern = r'^(#{1,6})\s+(.+)$'
        headings = re.findall(heading_pattern, content, re.MULTILINE)
        
        for heading_level, heading_text in headings:
            level = len(heading_level)
            styles["heading_styles"][f"heading_{level}"] = {
                "font_size": max(16 - level, 10),
                "bold": True,
                "color": "black"
            }
        
        return styles
    
    def parse_doc_placeholders(self, doc_path: str) -> Dict[str, Any]:
        """解析DOC文档中的占位符（API兼容方法）"""
        try:
            # 读取文档内容
            if HAS_DOCX:
                try:
                    doc = Document(doc_path)
                    text_content = ""
                    for paragraph in doc.paragraphs:
                        text_content += paragraph.text + "\n"
                except Exception as docx_error:
                    self.logger.warning(f"python-docx解析失败，使用二进制方式: {docx_error}")
                    # 降级到二进制读取
                    with open(doc_path, 'rb') as f:
                        binary_data = f.read()
                    text_content = self._extract_text_from_word_doc(binary_data)
            else:
                # 如果没有python-docx，尝试二进制读取
                with open(doc_path, 'rb') as f:
                    binary_data = f.read()
                text_content = self._extract_text_from_word_doc(binary_data)
            
            # 提取占位符
            placeholders_info = self._extract_placeholders(text_content)
            
            # 分类占位符
            stats_placeholders = []
            chart_placeholders = []
            
            for placeholder in placeholders_info:
                placeholder_dict = {
                    "description": placeholder.description,
                    "placeholder_text": f"{{{{{placeholder.name}}}}}",
                    "placeholder_name": placeholder.name,
                    "placeholder_type": placeholder.type,
                    "content_type": placeholder.content_type.value if hasattr(placeholder.content_type, 'value') else str(placeholder.content_type)
                }
                
                # 根据类型分类
                if placeholder.type in ["统计", "statistic", "count", "sum", "average"]:
                    placeholder_dict["analysis_requirements"] = {
                        "data_operation": self._infer_data_operation(placeholder.name),
                        "aggregation_type": self._infer_aggregation_type(placeholder.name),
                        "time_dimension": "时间" in placeholder.name or "年" in placeholder.name or "月" in placeholder.name,
                        "geographic_dimension": "地区" in placeholder.name or "区域" in placeholder.name,
                        "requires_grouping": "分组" in placeholder.name or "按" in placeholder.name
                    }
                    stats_placeholders.append(placeholder_dict)
                elif placeholder.type in ["图表", "chart", "graph"]:
                    placeholder_dict["chart_requirements"] = {
                        "chart_type": self._infer_chart_type(placeholder.name),
                        "data_series": 1,
                        "x_axis": "category",
                        "y_axis": "value",
                        "show_legend": True
                    }
                    chart_placeholders.append(placeholder_dict)
                else:
                    # 默认归类为统计
                    placeholder_dict["analysis_requirements"] = {
                        "data_operation": "unknown",
                        "aggregation_type": None,
                        "time_dimension": False,
                        "geographic_dimension": False,
                        "requires_grouping": False
                    }
                    stats_placeholders.append(placeholder_dict)
            
            return {
                "success": True,
                "stats_placeholders": stats_placeholders,
                "chart_placeholders": chart_placeholders,
                "total_placeholders": len(placeholders_info),
                "document_path": doc_path
            }
            
        except Exception as e:
            self.logger.error(f"解析DOC占位符失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "stats_placeholders": [],
                "chart_placeholders": [],
                "total_placeholders": 0,
                "document_path": doc_path
            }
    
    def _infer_data_operation(self, placeholder_name: str) -> str:
        """推断数据操作类型"""
        name_lower = placeholder_name.lower()
        if any(keyword in name_lower for keyword in ['总数', '数量', 'count', '个数']):
            return 'count'
        elif any(keyword in name_lower for keyword in ['总和', '合计', 'sum', '总计']):
            return 'sum'
        elif any(keyword in name_lower for keyword in ['平均', 'avg', 'average', '均值']):
            return 'average'
        elif any(keyword in name_lower for keyword in ['最大', 'max', '最高']):
            return 'max'
        elif any(keyword in name_lower for keyword in ['最小', 'min', '最低']):
            return 'min'
        else:
            return 'unknown'
    
    def _infer_aggregation_type(self, placeholder_name: str) -> Optional[str]:
        """推断聚合类型"""
        name_lower = placeholder_name.lower()
        if any(keyword in name_lower for keyword in ['按月', '月度', 'monthly']):
            return 'monthly'
        elif any(keyword in name_lower for keyword in ['按年', '年度', 'yearly']):
            return 'yearly'
        elif any(keyword in name_lower for keyword in ['按日', '日度', 'daily']):
            return 'daily'
        elif any(keyword in name_lower for keyword in ['按区域', '地区', 'region']):
            return 'regional'
        else:
            return None
    
    def _infer_chart_type(self, placeholder_name: str) -> str:
        """推断图表类型"""
        name_lower = placeholder_name.lower()
        if any(keyword in name_lower for keyword in ['柱状图', 'bar', '条形图']):
            return 'bar'
        elif any(keyword in name_lower for keyword in ['折线图', 'line', '趋势图']):
            return 'line'
        elif any(keyword in name_lower for keyword in ['饼图', 'pie', '饼状图']):
            return 'pie'
        elif any(keyword in name_lower for keyword in ['散点图', 'scatter']):
            return 'scatter'
        else:
            return 'bar'  # 默认柱状图


class ContentFormatter:
    """内容格式化器"""
    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
    
    def format_content(self, raw_content: Any, content_type: ContentType, 
                      format_rules: Dict[str, Any] = None) -> ProcessedContent:
        """格式化内容"""
        format_rules = format_rules or {}
        
        if content_type == ContentType.TEXT:
            formatted = self._format_text(raw_content, format_rules)
        elif content_type == ContentType.TABLE:
            formatted = self._format_table(raw_content, format_rules)
        elif content_type == ContentType.CHART:
            formatted = self._format_chart(raw_content, format_rules)
        elif content_type == ContentType.LIST:
            formatted = self._format_list(raw_content, format_rules)
        elif content_type == ContentType.IMAGE:
            formatted = self._format_image(raw_content, format_rules)
        else:
            formatted = str(raw_content)
        
        return ProcessedContent(
            placeholder_name="",
            content_type=content_type,
            raw_content=raw_content,
            formatted_content=formatted,
            metadata={"formatted_at": datetime.now().isoformat()}
        )
    
    def _format_text(self, content: Any, format_rules: Dict) -> str:
        """格式化文本内容"""
        if isinstance(content, (dict, list)):
            # 如果是结构化数据，转换为文本
            if isinstance(content, dict):
                if "value" in content:
                    return str(content["value"])
                elif "results" in content:
                    results = content["results"]
                    if isinstance(results, dict):
                        # 提取主要信息
                        if "numeric_statistics" in results:
                            stats = results["numeric_statistics"]
                            text_parts = []
                            for field, field_stats in stats.items():
                                mean = field_stats.get("mean", 0)
                                count = field_stats.get("count", 0)
                                text_parts.append(f"{field}: 平均值 {mean:.2f}，共 {count} 项")
                            return "; ".join(text_parts)
                        else:
                            return json.dumps(results, ensure_ascii=False, indent=2)
                    else:
                        return str(results)
                else:
                    # 选择关键字段
                    key_fields = ["total", "count", "sum", "average", "result", "value"]
                    for field in key_fields:
                        if field in content:
                            return str(content[field])
                    return json.dumps(content, ensure_ascii=False)
            elif isinstance(content, list) and content:
                if isinstance(content[0], dict):
                    # 如果是单值结果
                    if len(content) == 1:
                        return self._format_text(content[0], format_rules)
                    # 如果是多行数据，格式化为摘要
                    return f"共 {len(content)} 项数据"
                else:
                    return ", ".join(str(item) for item in content[:5])
        
        return str(content)
    
    def _format_table(self, content: Any, format_rules: Dict) -> str:
        """格式化表格内容"""
        if not isinstance(content, list) or not content:
            return "无表格数据"
        
        # 如果是查询结果格式
        if isinstance(content, dict) and "data" in content:
            content = content["data"]
        
        if not isinstance(content, list) or not content:
            return "无表格数据"
        
        # 获取列名
        if isinstance(content[0], dict):
            columns = list(content[0].keys())
        else:
            return "无效表格数据格式"
        
        # 限制显示行数
        max_rows = format_rules.get("max_rows", 10)
        display_data = content[:max_rows]
        
        # 生成表格字符串
        table_lines = []
        
        # 表头
        header = " | ".join(columns)
        table_lines.append(header)
        table_lines.append(" | ".join(["---"] * len(columns)))
        
        # 数据行
        for row in display_data:
            if isinstance(row, dict):
                values = []
                for col in columns:
                    value = row.get(col, "")
                    # 格式化数值
                    if isinstance(value, float):
                        values.append(f"{value:.2f}")
                    else:
                        values.append(str(value))
                table_lines.append(" | ".join(values))
        
        # 如果有更多数据，添加省略号
        if len(content) > max_rows:
            table_lines.append(f"... (共 {len(content)} 行)")
        
        return "\n".join(table_lines)
    
    def _format_chart(self, content: Any, format_rules: Dict) -> str:
        """格式化图表内容"""
        if isinstance(content, dict):
            chart_type = content.get("chart_type", "unknown")
            data_points = 0
            
            if "chart_data" in content:
                data_points = len(content["chart_data"])
            elif "data" in content:
                data_points = len(content["data"]) if isinstance(content["data"], list) else 0
            
            return f"[{chart_type}图表，包含{data_points}个数据点]"
        
        return "[图表内容]"
    
    def _format_list(self, content: Any, format_rules: Dict) -> str:
        """格式化列表内容"""
        if isinstance(content, list):
            list_items = []
            for i, item in enumerate(content[:10]):  # 限制10项
                if isinstance(item, dict):
                    # 提取主要信息
                    if "name" in item and "value" in item:
                        list_items.append(f"{i+1}. {item['name']}: {item['value']}")
                    else:
                        list_items.append(f"{i+1}. {json.dumps(item, ensure_ascii=False)}")
                else:
                    list_items.append(f"{i+1}. {item}")
            
            if len(content) > 10:
                list_items.append(f"... (共 {len(content)} 项)")
            
            return "\n".join(list_items)
        
        return str(content)
    
    def _format_image(self, content: Any, format_rules: Dict) -> str:
        """格式化图片内容"""
        if isinstance(content, dict) and "image_path" in content:
            return f"[图片: {content['image_path']}]"
        elif isinstance(content, str) and content.startswith("data:image"):
            return "[Base64图片内容]"
        else:
            return "[图片内容]"


class WordDocumentGenerator:
    """Word文档生成器"""
    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
    
    def generate_document(self, template: DocumentTemplate, 
                         processed_contents: List[ProcessedContent],
                         output_path: str = None) -> DocumentGenerationResult:
        """生成Word文档"""
        start_time = time.time()
        
        try:
            # 生成输出路径
            if not output_path:
                timestamp = int(time.time())
                output_path = f"/tmp/report_{timestamp}.docx"
            
            if HAS_DOCX:
                # 使用python-docx生成真实文档
                result = self._generate_docx_document(template, processed_contents, output_path)
            else:
                # 使用模拟生成
                result = self._generate_mock_document(template, processed_contents, output_path)
            
            generation_time = time.time() - start_time
            result.generation_time = generation_time
            
            return result
            
        except Exception as e:
            generation_time = time.time() - start_time
            self.logger.error(f"文档生成失败: {e}")
            
            return DocumentGenerationResult(
                success=False,
                document_path="",
                document_format=DocumentFormat.DOCX,
                generation_time=generation_time,
                placeholder_count=len(template.placeholders),
                processed_placeholders=0,
                error_message=str(e)
            )
    
    def _generate_docx_document(self, template: DocumentTemplate,
                               processed_contents: List[ProcessedContent],
                               output_path: str) -> DocumentGenerationResult:
        """使用python-docx生成真实文档"""
        doc = Document()
        
        # 设置文档标题
        title = doc.add_heading(template.name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 处理模板内容
        content_lines = template.content.split('\n')
        content_map = {pc.placeholder_name: pc for pc in processed_contents}
        
        failed_placeholders = []
        processed_count = 0
        warnings = []
        
        for line in content_lines:
            line = line.strip()
            if not line:
                # 添加空行
                doc.add_paragraph()
                continue
            
            # 检查是否包含占位符
            placeholders_in_line = re.findall(r'\{([^}]+)\}', line)
            
            if placeholders_in_line:
                # 替换占位符
                processed_line = line
                
                for placeholder_name in placeholders_in_line:
                    if placeholder_name in content_map:
                        processed_content = content_map[placeholder_name]
                        replacement = processed_content.formatted_content
                        processed_line = processed_line.replace(f"{{{placeholder_name}}}", replacement)
                        processed_count += 1
                    else:
                        processed_line = processed_line.replace(f"{{{placeholder_name}}}", f"[未处理: {placeholder_name}]")
                        failed_placeholders.append(placeholder_name)
                        warnings.append(f"占位符 {placeholder_name} 未找到对应内容")
                
                # 根据内容类型添加到文档
                if any(pc.content_type == ContentType.TABLE for pc in content_map.values()
                       if pc.placeholder_name in placeholders_in_line):
                    # 如果包含表格内容，添加为表格
                    self._add_table_to_doc(doc, processed_line, content_map, placeholders_in_line)
                else:
                    # 添加为段落
                    self._add_paragraph_to_doc(doc, processed_line, line)
            else:
                # 普通文本行
                self._add_paragraph_to_doc(doc, line)
        
        # 保存文档
        doc.save(output_path)
        
        return DocumentGenerationResult(
            success=True,
            document_path=output_path,
            document_format=DocumentFormat.DOCX,
            generation_time=0,  # 会在上级函数中设置
            placeholder_count=len(template.placeholders),
            processed_placeholders=processed_count,
            failed_placeholders=failed_placeholders,
            warnings=warnings
        )
    
    def _generate_mock_document(self, template: DocumentTemplate,
                               processed_contents: List[ProcessedContent],
                               output_path: str) -> DocumentGenerationResult:
        """生成模拟文档（当没有python-docx时）"""
        content_map = {pc.placeholder_name: pc for pc in processed_contents}
        
        # 替换模板中的占位符
        final_content = template.content
        processed_count = 0
        failed_placeholders = []
        
        for placeholder in template.placeholders:
            placeholder_pattern = f"{{{placeholder.name}}}"
            
            if placeholder.name in content_map:
                replacement = content_map[placeholder.name].formatted_content
                final_content = final_content.replace(placeholder_pattern, replacement)
                processed_count += 1
            else:
                final_content = final_content.replace(placeholder_pattern, f"[未处理: {placeholder.name}]")
                failed_placeholders.append(placeholder.name)
        
        # 保存为文本文件（模拟）
        text_output_path = output_path.replace('.docx', '.txt')
        
        try:
            with open(text_output_path, 'w', encoding='utf-8') as f:
                f.write(f"# {template.name}\n\n")
                f.write(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"占位符总数: {len(template.placeholders)}\n")
                f.write(f"已处理: {processed_count}\n")
                f.write(f"未处理: {len(failed_placeholders)}\n\n")
                f.write("=" * 50 + "\n\n")
                f.write(final_content)
            
            self.logger.info(f"模拟文档已保存: {text_output_path}")
            
        except Exception as e:
            self.logger.warning(f"保存模拟文档失败: {e}")
            text_output_path = output_path
        
        return DocumentGenerationResult(
            success=True,
            document_path=text_output_path,
            document_format=DocumentFormat.DOCX,
            generation_time=0,
            placeholder_count=len(template.placeholders),
            processed_placeholders=processed_count,
            failed_placeholders=failed_placeholders,
            warnings=["使用模拟文档生成器（未安装python-docx）"] if not HAS_DOCX else [],
            metadata={"mock_generation": not HAS_DOCX}
        )
    
    def _add_paragraph_to_doc(self, doc, content: str, original_line: str = None):
        """添加段落到文档"""
        if content.startswith('#'):
            # 标题
            level = min(len(content) - len(content.lstrip('#')), 6)
            title_text = content.lstrip('#').strip()
            doc.add_heading(title_text, level)
        else:
            # 普通段落
            para = doc.add_paragraph(content)
            
            # 如果内容包含数字，设置为右对齐
            if re.search(r'\d+\.?\d*', content):
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _add_table_to_doc(self, doc, content: str, content_map: Dict,
                          placeholders_in_line: List[str]):
        """添加表格到文档"""
        # 简化实现：将表格内容作为格式化文本添加
        doc.add_paragraph(content)
        
        # 这里可以扩展为真正的表格插入逻辑
        for placeholder_name in placeholders_in_line:
            if (placeholder_name in content_map and 
                content_map[placeholder_name].content_type == ContentType.TABLE):
                
                table_content = content_map[placeholder_name].formatted_content
                # 简单添加为代码块样式
                para = doc.add_paragraph(table_content)
                para.style = 'Normal'


class DocumentPipeline:
    """文档生成流水线"""
    
    def __init__(self):
        self.pipeline_id = str(uuid.uuid4())
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
        
        # 初始化组件
        self.template_parser = TemplateParser()
        self.content_formatter = ContentFormatter()
        self.document_generator = WordDocumentGenerator()
        
        # 流水线统计
        self.stats = {
            "templates_processed": 0,
            "documents_generated": 0,
            "total_processing_time": 0,
            "success_rate": 0
        }
    
    async def process_template(self, template_content: str, placeholder_data: Dict[str, Any],
                              output_path: str = None, template_name: str = "Generated Report") -> DocumentGenerationResult:
        """处理模板并生成文档"""
        start_time = time.time()
        
        try:
            self.logger.info(f"开始处理模板: {template_name}")
            
            # 1. 解析模板
            template = self.template_parser.parse_template(template_content)
            template.name = template_name
            
            # 2. 处理占位符数据
            processed_contents = []
            
            for placeholder in template.placeholders:
                if placeholder.name in placeholder_data:
                    raw_content = placeholder_data[placeholder.name]
                    
                    # 格式化内容
                    processed_content = self.content_formatter.format_content(
                        raw_content=raw_content,
                        content_type=placeholder.content_type,
                        format_rules=placeholder.format_rules
                    )
                    processed_content.placeholder_name = placeholder.name
                    
                    processed_contents.append(processed_content)
            
            # 3. 生成文档
            result = self.document_generator.generate_document(
                template=template,
                processed_contents=processed_contents,
                output_path=output_path
            )
            
            # 4. 更新统计信息
            self.stats["templates_processed"] += 1
            if result.success:
                self.stats["documents_generated"] += 1
            
            total_time = time.time() - start_time
            self.stats["total_processing_time"] += total_time
            self.stats["success_rate"] = self.stats["documents_generated"] / self.stats["templates_processed"]
            
            self.logger.info(f"模板处理完成: {template_name}, 耗时: {total_time:.2f}秒")
            
            return result
            
        except Exception as e:
            self.logger.error(f"模板处理失败: {e}")
            return DocumentGenerationResult(
                success=False,
                document_path="",
                document_format=DocumentFormat.DOCX,
                generation_time=time.time() - start_time,
                placeholder_count=0,
                processed_placeholders=0,
                error_message=str(e)
            )
    
    def get_pipeline_stats(self) -> Dict[str, Any]:
        """获取流水线统计信息"""
        avg_processing_time = (
            self.stats["total_processing_time"] / self.stats["templates_processed"]
            if self.stats["templates_processed"] > 0 else 0
        )
        
        return {
            **self.stats,
            "pipeline_id": self.pipeline_id,
            "average_processing_time": avg_processing_time,
            "uptime": datetime.now().isoformat()
        }
    
    async def batch_process_templates(self, template_requests: List[Dict[str, Any]]) -> List[DocumentGenerationResult]:
        """批量处理模板"""
        self.logger.info(f"开始批量处理 {len(template_requests)} 个模板")
        
        results = []
        
        for i, request in enumerate(template_requests):
            self.logger.info(f"处理模板 {i+1}/{len(template_requests)}")
            
            result = await self.process_template(
                template_content=request.get("template_content", ""),
                placeholder_data=request.get("placeholder_data", {}),
                output_path=request.get("output_path"),
                template_name=request.get("template_name", f"Template_{i+1}")
            )
            
            results.append(result)
        
        self.logger.info(f"批量处理完成，成功: {sum(1 for r in results if r.success)}/{len(results)}")
        
        return results


# 全局流水线实例
document_pipeline = DocumentPipeline()


# 便捷函数
async def generate_document_from_template(template_content: str, placeholder_data: Dict[str, Any],
                                        output_path: str = None, template_name: str = "Report") -> DocumentGenerationResult:
    """从模板生成文档的便捷函数"""
    return await document_pipeline.process_template(
        template_content=template_content,
        placeholder_data=placeholder_data,
        output_path=output_path,
        template_name=template_name
    )


def create_sample_template() -> str:
    """创建示例模板"""
    return """
# {报告标题}

## 执行摘要
{执行摘要}

## 数据概览
在本报告期间，我们收集和分析了以下数据：

### 基础统计
- 总记录数：{总记录数}
- 平均值：{平均值}
- 数据完整性：{数据完整性}

## 详细分析

### 分类统计
{分类统计表格}

### 趋势分析
{趋势分析图表}

### 关键洞察
{关键洞察列表}

## 结论和建议
基于以上分析，我们得出以下结论：

{结论和建议}

## 附录
- 数据来源：{数据来源}
- 生成时间：{生成时间}
- 报告版本：{报告版本}
"""


if __name__ == "__main__":
    # 测试用例
    async def test_document_pipeline():
        """测试文档生成流水线"""
        print("🧪 测试文档生成流水线...")
        
        # 创建示例模板
        template_content = create_sample_template()
        
        # 准备占位符数据
        placeholder_data = {
            "报告标题": "月度数据分析报告",
            "执行摘要": "本报告分析了本月的业务数据，发现了多个重要趋势和机会点。",
            "总记录数": 1234,
            "平均值": 89.5,
            "数据完整性": "95.2%",
            "分类统计表格": [
                {"类别": "产品A", "数量": 150, "占比": "30%"},
                {"类别": "产品B", "数量": 120, "占比": "24%"},
                {"类别": "产品C", "数量": 230, "占比": "46%"}
            ],
            "趋势分析图表": {
                "chart_type": "line",
                "data": [{"month": "1月", "value": 100}, {"month": "2月", "value": 120}]
            },
            "关键洞察列表": [
                "产品C表现最佳，占比达到46%",
                "整体趋势呈上升态势",
                "数据质量良好，完整性超过95%"
            ],
            "结论和建议": "建议继续关注产品C的发展，同时改进产品A和B的市场策略。",
            "数据来源": "内部业务系统",
            "生成时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "报告版本": "v1.0"
        }
        
        # 生成文档
        result = await generate_document_from_template(
            template_content=template_content,
            placeholder_data=placeholder_data,
            template_name="测试报告"
        )
        
        print(f"✅ 文档生成完成!")
        print(f"成功: {result.success}")
        print(f"文档路径: {result.document_path}")
        print(f"生成时间: {result.generation_time:.2f}秒")
        print(f"占位符处理: {result.processed_placeholders}/{result.placeholder_count}")
        
        if result.warnings:
            print("⚠️ 警告:")
            for warning in result.warnings:
                print(f"  - {warning}")
        
        if result.failed_placeholders:
            print("❌ 失败的占位符:")
            for failed in result.failed_placeholders:
                print(f"  - {failed}")
        
        # 显示流水线统计
        stats = document_pipeline.get_pipeline_stats()
        print(f"\n📊 流水线统计:")
        print(f"  - 处理模板数: {stats['templates_processed']}")
        print(f"  - 生成文档数: {stats['documents_generated']}")
        print(f"  - 成功率: {stats['success_rate']:.2%}")
        print(f"  - 平均处理时间: {stats['average_processing_time']:.2f}秒")
    
    # 运行测试
    asyncio.run(test_document_pipeline())