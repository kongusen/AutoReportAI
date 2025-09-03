import base64
import io
import logging
import re
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, Optional, List
from pathlib import Path

import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from PIL import Image

from app.services.infrastructure.storage.file_storage_service import file_storage_service

logger = logging.getLogger(__name__)


class WordGeneratorService:
    # Regex patterns for different content types
    IMG_REGEX = re.compile(r'<img src="data:image/png;base64,([^"]+)">')
    CHART_PLACEHOLDER_REGEX = re.compile(r'\{\{chart:(\w+)(?::([^}]+))?\}\}')  # {{chart:bar}} or {{chart:bar:title}}
    
    # Chart configuration
    DEFAULT_CHART_WIDTH = Inches(6.0)
    DEFAULT_CHART_HEIGHT = Inches(4.0)
    CHART_MAX_WIDTH = Inches(6.5)
    CHART_MAX_HEIGHT = Inches(5.0)

    def generate_report_from_template(
        self,
        template_content: str,
        placeholder_values: Dict[str, Any],
        title: str = "自动生成报告",
        format: str = "docx",
        chart_results: Optional[List[Dict[str, Any]]] = None
    ) -> str:
        """
        基于模板生成报告文件
        
        Args:
            template_content: 模板内容 (hex编码的二进制数据或文本)
            placeholder_values: 占位符值字典
            title: 报告标题
            format: 文件格式 (docx)
            chart_results: 图表生成结果列表
            
        Returns:
            存储后的文件路径
        """
        try:
            # 1. 处理模板内容
            if self._is_binary_template(template_content):
                # 处理二进制模板 (hex编码)
                return self._generate_from_binary_template(
                    template_content, placeholder_values, title, chart_results
                )
            else:
                # 处理文本模板
                return self._generate_from_text_template(
                    template_content, placeholder_values, title, chart_results
                )
                
        except Exception as e:
            logger.error(f"基于模板的报告生成失败: {str(e)}")
            # 降级到普通生成
            processed_content = self._replace_placeholders_in_text(
                template_content, placeholder_values
            )
            return self.generate_report(processed_content, title, format)

    def _is_binary_template(self, content: str) -> bool:
        """判断是否为二进制模板内容"""
        if not content:
            return False
        # 检查是否为hex编码
        content_clean = content.replace(' ', '').replace('\n', '')
        return len(content_clean) > 100 and all(c in '0123456789ABCDEFabcdef' for c in content_clean)

    def _generate_from_binary_template(
        self,
        template_content: str,
        placeholder_values: Dict[str, Any],
        title: str,
        chart_results: Optional[List[Dict[str, Any]]] = None
    ) -> str:
        """从二进制模板生成报告"""
        try:
            # 解析hex编码的二进制数据
            binary_data = bytes.fromhex(template_content.replace(' ', '').replace('\n', ''))
            
            # 加载为Word文档
            template_buffer = BytesIO(binary_data)
            doc = docx.Document(template_buffer)
            
            # 替换占位符（包括图表占位符）
            self._replace_placeholders_in_doc(doc, placeholder_values)
            
            # 插入图表
            if chart_results:
                self._insert_charts_in_doc(doc, chart_results)
            
            # 保存到内存
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            # 生成文件名并上传 (保持DOCX格式)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{title}_{timestamp}.docx"
            
            file_info = file_storage_service.upload_file(
                file_data=doc_buffer,
                original_filename=filename,
                file_type="report",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            logger.info(f"基于二进制模板的报告生成成功: {file_info['file_path']}")
            return file_info['file_path']
            
        except Exception as e:
            logger.warning(f"二进制模板处理失败，降级到文本处理: {e}")
            raise

    def _generate_from_text_template(
        self,
        template_content: str,
        placeholder_values: Dict[str, Any],
        title: str,
        chart_results: Optional[List[Dict[str, Any]]] = None
    ) -> str:
        """从文本模板生成报告"""
        # 替换占位符
        processed_content = self._replace_placeholders_in_text(template_content, placeholder_values)
        
        # 如果有图表，需要生成Word文档并插入图表
        if chart_results:
            return self._generate_report_with_charts(processed_content, title, chart_results)
        else:
            # 使用普通生成方法
            return self.generate_report(processed_content, title)

    def _replace_placeholders_in_doc(self, doc, placeholder_values: Dict[str, Any]):
        """在Word文档中替换占位符"""
        logger.error("🔥🔥🔥 CLAUDE修复的占位符替换代码正在运行! 🔥🔥🔥")
        logger.info(f"开始占位符替换 - 输入 {len(placeholder_values)} 个占位符")
        
        # 处理占位符值，提取实际值
        processed_values = {}
        logger.info("处理占位符值详情:")
        for key, value_info in placeholder_values.items():
            logger.info(f"处理占位符: {key}, 原始值: {value_info}")
            
            if isinstance(value_info, dict) and "value" in value_info:
                extracted_value = self._extract_value_from_result(value_info["value"])
                processed_values[key] = str(extracted_value) if extracted_value is not None else ""
            else:
                extracted_value = self._extract_value_from_result(value_info)
                processed_values[key] = str(extracted_value) if extracted_value is not None else ""
            
            logger.info(f"提取后的值: {key} = '{processed_values[key]}'")
            
            # 如果值为空或者是 "None"，为常见占位符提供默认值
            if not processed_values[key] or processed_values[key] in ["None", "null", ""]:
                default_value = self._get_default_placeholder_value(key)
                if default_value:
                    processed_values[key] = default_value
                    logger.warning(f"⚠️  使用默认值: {key} = {default_value}")
                else:
                    logger.error(f"❌ 占位符无法处理: {key} (无默认值)")
        
        # 显示处理后的占位符值（仅显示前5个）
        logger.info(f"处理后的占位符值（显示前5个）:")
        count = 0
        for key, value in processed_values.items():
            if count >= 5:
                break
            logger.info(f"  {key} = '{value}' (长度: {len(value)})")
            count += 1
        
        # 替换段落中的占位符
        replacements_made = 0
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            if original_text.strip() and "{{" in original_text:  # 只处理包含占位符的非空段落
                logger.debug(f"检查段落文本: {original_text}")
                paragraph_changed = False
                for key, value in processed_values.items():
                    # 支持多种占位符格式，包括中文冒号分隔符
                    patterns = [
                        f"{{{{{key}}}}}",    # {{key}}
                        f"{{{key}}}",        # {key}
                    ]
                        
                    for pattern in patterns:
                        if pattern in paragraph.text:
                            paragraph.text = paragraph.text.replace(pattern, value)
                            replacements_made += 1
                            logger.info(f"段落替换成功: {pattern} -> {value}")
                            paragraph_changed = True
                            break  # 找到一个匹配就跳出，避免重复替换
                
                if paragraph_changed:
                    logger.debug(f"段落替换后: {paragraph.text}")
        
        logger.info(f"段落中完成 {replacements_made} 个占位符替换")
        
        # 替换表格中的占位符
        table_replacements = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip() and "{{" in cell.text:  # 只处理包含占位符的非空单元格
                        logger.debug(f"检查表格单元格文本: {cell.text}")
                        cell_changed = False
                        for key, value in processed_values.items():
                            # 支持多种占位符格式，包括中文冒号分隔符
                            patterns = [
                                f"{{{{{key}}}}}",    # {{key}}
                                f"{{{key}}}",        # {key}
                            ]
                                
                            for pattern in patterns:
                                if pattern in cell.text:
                                    cell.text = cell.text.replace(pattern, value)
                                    table_replacements += 1
                                    logger.info(f"表格替换成功: {pattern} -> {value}")
                                    cell_changed = True
                                    break  # 找到一个匹配就跳出，避免重复替换
                        
                        if cell_changed:
                            logger.debug(f"表格单元格替换后: {cell.text}")
        
        logger.info(f"表格中完成 {table_replacements} 个占位符替换")
        total_replacements = replacements_made + table_replacements
        logger.info(f"总共完成 {total_replacements} 个占位符替换")

    def _get_default_placeholder_value(self, key: str) -> str:
        """为常见占位符提供默认值"""
        from datetime import datetime
        current_time = datetime.now()
        
        # 时间相关占位符
        if "报告年份" in key:
            return str(current_time.year)
        elif "统计开始日期" in key:
            return f"{current_time.year}-{current_time.month:02d}-01"
        elif "统计结束日期" in key:
            next_month = current_time.month + 1 if current_time.month < 12 else 1
            next_year = current_time.year if current_time.month < 12 else current_time.year + 1
            return f"{next_year}-{next_month:02d}-01"
        elif "地区名称" in key:
            return "云南省"  # 默认地区
        
        # 数量相关占位符
        elif "投诉件数" in key or "件数" in key:
            return "0"
        elif "占比" in key or "百分比" in key:
            return "0.0"
        elif "时长" in key:
            return "0"
        
        return None

    def _extract_value_from_result(self, value: Any) -> Any:
        """从查询结果中提取实际数值"""
        try:
            # 如果是None或简单类型，直接返回
            if value is None or isinstance(value, (str, int, float, bool)):
                return value
            
            # 如果是DorisQueryResult对象或类似结构
            if hasattr(value, 'data') and hasattr(value, 'execution_time'):
                return self._extract_value_from_result(value.data)
            
            # 如果是pandas DataFrame
            if hasattr(value, 'iloc') and hasattr(value, '__len__'):
                try:
                    if len(value) > 0 and hasattr(value, 'empty') and not value.empty:
                        return value.iloc[0, 0]
                    return 0
                except Exception:
                    # 如果DataFrame访问失败，尝试其他方法
                    pass
            
            # 如果是包含data字段的字典
            if isinstance(value, dict):
                # 首先检查是否有嵌套的data结构
                if 'data' in value:
                    nested_result = self._extract_value_from_result(value['data'])
                    if nested_result != value['data']:  # 避免无限递归
                        return nested_result
                
                # 尝试从常见字段中获取数值
                for key in ['count', 'total', 'sum', 'avg', 'value', 'result', 'amount']:
                    if key in value:
                        extracted = self._extract_value_from_result(value[key])
                        if extracted != value[key]:  # 避免无限递归
                            return extracted
                
                # 如果字典只有一个键值对，返回值
                if len(value) == 1:
                    return next(iter(value.values()))
                    
                # 如果没有找到特定字段，返回第一个非None值
                for v in value.values():
                    if v is not None and not isinstance(v, dict):
                        return self._extract_value_from_result(v)
                
            # 如果是列表或元组
            if isinstance(value, (list, tuple)) and value:
                first_item = value[0]
                if isinstance(first_item, dict):
                    # 如果第一个元素是字典，尝试获取其值
                    return self._extract_value_from_result(first_item)
                else:
                    # 如果是简单值，直接返回
                    return first_item
                
            # 如果是其他对象类型，尝试转换为字符串
            if hasattr(value, '__str__') and not str(value).startswith('<'):
                return str(value)
                
            # 最后返回"无数据"而不是对象引用
            return "无数据"
            
        except Exception as e:
            logger.warning(f"提取查询结果数值失败: {e}, 值类型: {type(value)}")
            # 发生异常时返回"无数据"而不是原值
            return "无数据"

    def _replace_placeholders_in_text(self, content: str, placeholder_values: Dict[str, Any]) -> str:
        """在文本中替换占位符"""
        processed_content = content
        
        # 处理占位符值，提取实际值
        for key, value_info in placeholder_values.items():
            if isinstance(value_info, dict) and "value" in value_info:
                extracted_value = self._extract_value_from_result(value_info["value"])
                value = str(extracted_value)
            else:
                extracted_value = self._extract_value_from_result(value_info)
                value = str(extracted_value)
                
            # 支持多种占位符格式
            patterns = [
                f"{{{{{key}}}}}",  # {{key}}
                f"{{{key}}}",      # {key}
            ]
            for pattern in patterns:
                processed_content = processed_content.replace(pattern, value)
        return processed_content

    def _generate_report_with_charts(
        self,
        content: str,
        title: str,
        chart_results: List[Dict[str, Any]]
    ) -> str:
        """
        生成包含图表的报告
        
        Args:
            content: 文本内容
            title: 报告标题
            chart_results: 图表生成结果列表
            
        Returns:
            存储后的文件路径
        """
        try:
            # 创建Word文档
            doc = docx.Document()
            
            # 设置标题
            title_paragraph = doc.add_heading(title, level=1)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 处理内容，识别图表占位符
            self._process_content_with_charts(doc, content, chart_results)
            
            # 保存文档
            return self._save_doc_to_storage(doc, title)
            
        except Exception as e:
            logger.error(f"生成包含图表的报告失败: {e}")
            # 降级到普通报告生成
            return self.generate_report(content, title)

    def _process_content_with_charts(
        self,
        doc: docx.Document,
        content: str,
        chart_results: List[Dict[str, Any]]
    ):
        """
        处理内容并插入图表
        
        Args:
            doc: Word文档对象
            content: 文本内容
            chart_results: 图表结果列表
        """
        # 按行分割内容
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()  # 空行
                continue
            
            # 检查是否包含图表占位符
            chart_matches = self.CHART_PLACEHOLDER_REGEX.findall(line)
            
            if chart_matches:
                # 处理图表占位符行
                self._process_chart_line(doc, line, chart_matches, chart_results)
            else:
                # 普通文本行
                paragraph = doc.add_paragraph(line)
                
    def _process_chart_line(
        self,
        doc: docx.Document,
        line: str,
        chart_matches: List[tuple],
        chart_results: List[Dict[str, Any]]
    ):
        """
        处理包含图表占位符的行
        
        Args:
            doc: Word文档对象
            line: 包含图表占位符的行
            chart_matches: 图表占位符匹配结果
            chart_results: 图表结果列表
        """
        remaining_text = line
        
        for chart_type, chart_title in chart_matches:
            chart_placeholder = f"{{{{chart:{chart_type}"
            if chart_title:
                chart_placeholder += f":{chart_title}"
            chart_placeholder += "}}"
            
            # 查找匹配的图表结果
            matching_chart = self._find_matching_chart(chart_type, chart_title, chart_results)
            
            if matching_chart:
                # 添加图表前的文本
                before_chart = remaining_text.split(chart_placeholder)[0]
                if before_chart.strip():
                    doc.add_paragraph(before_chart.strip())
                
                # 插入图表
                self._insert_single_chart(doc, matching_chart)
                
                # 更新剩余文本
                parts = remaining_text.split(chart_placeholder, 1)
                remaining_text = parts[1] if len(parts) > 1 else ""
            else:
                # 图表未找到，替换为提示文本
                replacement_text = f"[图表未找到: {chart_type}]"
                remaining_text = remaining_text.replace(chart_placeholder, replacement_text)
        
        # 添加剩余文本
        if remaining_text.strip():
            doc.add_paragraph(remaining_text.strip())

    def _find_matching_chart(
        self,
        chart_type: str,
        chart_title: Optional[str],
        chart_results: List[Dict[str, Any]]
    ) -> Optional[Dict[str, Any]]:
        """
        查找匹配的图表结果
        
        Args:
            chart_type: 图表类型 (bar, line, pie)
            chart_title: 图表标题 (可选)
            chart_results: 图表结果列表
            
        Returns:
            匹配的图表结果或None
        """
        for chart in chart_results:
            # 匹配图表类型
            chart_result_type = chart.get('chart_type', '').replace('_chart', '')
            if chart_result_type == chart_type:
                # 如果指定了标题，进一步匹配标题
                if chart_title:
                    if chart_title.lower() in chart.get('title', '').lower():
                        return chart
                else:
                    # 没有指定标题，返回第一个匹配类型的图表
                    return chart
        
        # 如果没有精确匹配，返回第一个图表作为备选
        return chart_results[0] if chart_results else None

    def _insert_charts_in_doc(self, doc: docx.Document, chart_results: List[Dict[str, Any]]):
        """
        在文档中插入图表（替换占位符）
        
        Args:
            doc: Word文档对象
            chart_results: 图表结果列表
        """
        logger.info(f"开始在文档中插入 {len(chart_results)} 个图表")
        
        chart_inserted = 0
        
        # 遍历所有段落，查找图表占位符
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            if not original_text or '{{chart:' not in original_text:
                continue
                
            logger.debug(f"检查段落: {original_text}")
            
            # 查找图表占位符
            chart_matches = self.CHART_PLACEHOLDER_REGEX.findall(original_text)
            
            if chart_matches:
                # 清除段落内容
                paragraph.clear()
                
                # 处理每个图表占位符
                for chart_type, chart_title in chart_matches:
                    matching_chart = self._find_matching_chart(chart_type, chart_title, chart_results)
                    
                    if matching_chart:
                        # 插入图表到段落
                        self._insert_chart_in_paragraph(paragraph, matching_chart)
                        chart_inserted += 1
                        logger.info(f"成功插入图表: {matching_chart.get('title', chart_type)}")
                    else:
                        # 图表未找到，添加提示文本
                        run = paragraph.add_run(f"[图表未找到: {chart_type}]")
                        run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)  # 红色
                        logger.warning(f"图表未找到: {chart_type}")
        
        logger.info(f"文档中成功插入 {chart_inserted} 个图表")

    def _insert_single_chart(self, doc: docx.Document, chart_info: Dict[str, Any]):
        """
        在文档中插入单个图表
        
        Args:
            doc: Word文档对象
            chart_info: 图表信息
        """
        try:
            chart_filepath = chart_info.get('filepath')
            chart_title = chart_info.get('title', '图表')
            
            if not chart_filepath or not Path(chart_filepath).exists():
                logger.warning(f"图表文件不存在: {chart_filepath}")
                error_paragraph = doc.add_paragraph(f"[图表文件不存在: {chart_title}]")
                return
            
            # 获取图片尺寸
            width, height = self._calculate_chart_size(chart_filepath)
            
            # 添加图表标题（如果有）
            if chart_title and chart_title != '图表':
                title_paragraph = doc.add_paragraph()
                title_run = title_paragraph.add_run(chart_title)
                title_run.font.size = Pt(12)
                title_run.font.bold = True
                title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 插入图片
            chart_paragraph = doc.add_paragraph()
            chart_run = chart_paragraph.add_run()
            chart_run.add_picture(chart_filepath, width=width, height=height)
            chart_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加间距
            doc.add_paragraph()
            
            logger.info(f"成功插入图表: {chart_title} ({chart_filepath})")
            
        except Exception as e:
            logger.error(f"插入图表失败: {e}")
            error_paragraph = doc.add_paragraph(f"[图表插入失败: {chart_info.get('title', '未知图表')}]")

    def _insert_chart_in_paragraph(self, paragraph, chart_info: Dict[str, Any]):
        """
        在段落中插入图表
        
        Args:
            paragraph: 段落对象
            chart_info: 图表信息
        """
        try:
            chart_filepath = chart_info.get('filepath')
            chart_title = chart_info.get('title', '图表')
            
            if not chart_filepath or not Path(chart_filepath).exists():
                logger.warning(f"图表文件不存在: {chart_filepath}")
                run = paragraph.add_run(f"[图表文件不存在: {chart_title}]")
                run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
                return
            
            # 获取图片尺寸
            width, height = self._calculate_chart_size(chart_filepath)
            
            # 添加换行和图表
            paragraph.add_run().add_break()
            
            # 插入图片
            run = paragraph.add_run()
            run.add_picture(chart_filepath, width=width, height=height)
            
            # 添加图表标题
            if chart_title and chart_title != '图表':
                paragraph.add_run().add_break()
                title_run = paragraph.add_run(chart_title)
                title_run.font.size = Pt(10)
                title_run.italic = True
            
        except Exception as e:
            logger.error(f"在段落中插入图表失败: {e}")
            run = paragraph.add_run(f"[图表插入失败: {chart_info.get('title', '未知图表')}]")
            run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)

    def _calculate_chart_size(self, chart_filepath: str) -> tuple:
        """
        计算图表在文档中的适当尺寸
        
        Args:
            chart_filepath: 图表文件路径
            
        Returns:
            (width, height) 元组
        """
        try:
            # 使用PIL获取图片尺寸
            with Image.open(chart_filepath) as img:
                img_width, img_height = img.size
                aspect_ratio = img_height / img_width
                
                # 计算适当的文档尺寸
                doc_width = min(self.DEFAULT_CHART_WIDTH, self.CHART_MAX_WIDTH)
                doc_height = doc_width * aspect_ratio
                
                # 限制最大高度
                if doc_height > self.CHART_MAX_HEIGHT:
                    doc_height = self.CHART_MAX_HEIGHT
                    doc_width = doc_height / aspect_ratio
                
                return doc_width, doc_height
                
        except Exception as e:
            logger.warning(f"计算图表尺寸失败，使用默认尺寸: {e}")
            return self.DEFAULT_CHART_WIDTH, self.DEFAULT_CHART_HEIGHT

    def _save_doc_to_storage(self, doc: docx.Document, title: str) -> str:
        """
        保存Word文档到存储系统
        
        Args:
            doc: Word文档对象
            title: 文档标题
            
        Returns:
            文件路径
        """
        try:
            # 保存到内存
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            # 生成文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{title}_{timestamp}.docx"
            
            # 上传到存储系统
            file_info = file_storage_service.upload_file(
                file_data=doc_buffer,
                original_filename=filename,
                file_type="report",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            logger.info(f"Word文档保存成功: {file_info['file_path']}")
            return file_info['file_path']
            
        except Exception as e:
            logger.error(f"保存Word文档失败: {e}")
            raise

    def _get_file_format_info(self, format_type: str) -> Dict[str, str]:
        """获取文件格式信息"""
        format_map = {
            'docx': {
                'extension': '.docx',
                'content_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            },
            'xlsx': {
                'extension': '.xlsx', 
                'content_type': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            },
            'txt': {
                'extension': '.txt',
                'content_type': 'text/plain'
            },
            'html': {
                'extension': '.html',
                'content_type': 'text/html'
            }
        }
        return format_map.get(format_type, format_map['docx'])

    def generate_report(
        self,
        content: str,
        title: str = "自动生成报告",
        format: str = "docx"
    ) -> str:
        """
        生成报告文件并存储到MinIO或本地
        
        Args:
            content: 报告内容
            title: 报告标题  
            format: 文件格式 (docx)
            
        Returns:
            存储后的文件路径
        """
        try:
            # 创建Word文档
            doc = docx.Document()
            
            # 添加标题
            doc.add_heading(title, 0)
            
            # 添加生成时间
            doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("")  # 空行
            
            # 处理内容
            self._process_content(doc, content)
            
            # 将文档保存到内存
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            # 获取格式信息并生成文件名
            format_info = self._get_file_format_info(format)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{title}_{timestamp}{format_info['extension']}"
            
            # 上传到存储服务
            file_info = file_storage_service.upload_file(
                file_data=doc_buffer,
                original_filename=filename,
                file_type="report",
                content_type=format_info['content_type']
            )
            
            logger.info(f"报告生成成功: {file_info['file_path']}")
            return file_info['file_path']
            
        except Exception as e:
            logger.error(f"报告生成失败: {str(e)}")
            raise

    def generate_report_from_content(self, composed_content: str, output_path: str):
        """
        Generates a .docx report from a string containing the fully composed content.
        This content can include text and special <img> tags for base64 images.
        
        保留原有方法以保持向后兼容
        """
        doc = docx.Document()

        # Split the content into paragraphs based on newlines
        paragraphs = composed_content.split("\n")

        for para_text in paragraphs:
            self._process_paragraph(doc, para_text)

        doc.save(output_path)

    def _process_content(self, doc, content: str):
        """处理报告内容，支持文本和图像"""
        # 分段处理
        paragraphs = content.split("\n")
        
        for para_text in paragraphs:
            self._process_paragraph(doc, para_text)

    def _process_paragraph(self, doc, para_text: str):
        """
        Processes a single paragraph of text, adding text and images to the document.
        """
        # Find all image tags in the paragraph
        img_matches = list(self.IMG_REGEX.finditer(para_text))

        if not img_matches:
            # If no images, add the whole paragraph as text
            doc.add_paragraph(para_text)
            return

        current_pos = 0
        for match in img_matches:
            # Add text before the image
            start, end = match.span()
            if start > current_pos:
                doc.add_paragraph(para_text[current_pos:start])

            # Add the image
            base64_data = match.group(1)
            try:
                image_stream = io.BytesIO(base64.b64decode(base64_data))
                doc.add_picture(image_stream, width=Inches(6.0))
            except Exception as e:
                print(f"Error decoding or adding picture: {e}")
                # Add a placeholder text on error
                doc.add_paragraph(f"[Image could not be loaded: {e}]")

            current_pos = end

        # Add any remaining text after the last image
        if current_pos < len(para_text):
            doc.add_paragraph(para_text[current_pos:])


word_generator_service = WordGeneratorService()
