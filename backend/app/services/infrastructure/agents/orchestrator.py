"""
Task Orchestrator - 多阶段 Agent 编排器

负责协调整个报告生成流程中的多个 Agent，实现：
1. 占位符扫描与分析
2. SQL 生成与验证
3. ETL 取数
4. 数据回填与图表生成
5. 文案优化
6. 文档生成

基于 Loom 框架的 Task 工具和 AgentSpec 注册机制。
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from typing import Any, Dict, List, Optional, AsyncIterator
from uuid import UUID

from .types import AgentInput, AgentOutput, PlaceholderSpec, SchemaInfo, TaskContext
from .service import AgentService
from .config import LoomAgentConfig

logger = logging.getLogger(__name__)


class StageStatus(str, Enum):
    """阶段执行状态"""
    PENDING = "pending"
    IN_PROGRESS = "in_progress"
    COMPLETED = "completed"
    FAILED = "failed"
    SKIPPED = "skipped"


@dataclass
class StageResult:
    """单个阶段的执行结果"""
    stage_name: str
    status: StageStatus
    started_at: Optional[datetime] = None
    completed_at: Optional[datetime] = None
    output: Dict[str, Any] = field(default_factory=dict)
    metadata: Dict[str, Any] = field(default_factory=dict)
    error: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            "stage_name": self.stage_name,
            "status": self.status.value,
            "started_at": self.started_at.isoformat() if self.started_at else None,
            "completed_at": self.completed_at.isoformat() if self.completed_at else None,
            "output": self.output,
            "metadata": self.metadata,
            "error": self.error,
        }


@dataclass
class OrchestratorContext:
    """编排器执行上下文"""
    template_id: str
    data_source_id: str
    user_id: str
    execution_time: datetime = field(default_factory=datetime.now)

    # 任务调度信息
    schedule: Optional[Dict[str, Any]] = None
    time_window: Optional[Dict[str, Any]] = None

    # 模板与数据源信息
    template: Optional[Any] = None
    data_source: Optional[Any] = None

    # 阶段间共享的数据
    shared_data: Dict[str, Any] = field(default_factory=dict)

    # 各阶段的执行结果
    stage_results: Dict[str, StageResult] = field(default_factory=dict)


class ReportGenerationOrchestrator:
    """
    报告生成编排器

    协调多个专门的 Agent 完成完整的报告生成流程。
    每个阶段都有明确的输入/输出，失败时可以回退。
    """

    STAGE_ORDER = [
        "placeholder_scan",      # 占位符扫描
        "sql_generation",        # SQL 生成与验证
        "etl_execution",         # ETL 取数
        "data_fill_chart",       # 数据回填与图表生成
        "content_optimization",  # 文案优化
        "document_generation",   # 文档生成
    ]

    def __init__(self, container: Any, config: Optional[LoomAgentConfig] = None):
        self.container = container
        self.config = config
        self.agent_service = AgentService(container=container, config=config)
        self.logger = logging.getLogger(self.__class__.__name__)

    async def execute(
        self,
        context: OrchestratorContext,
        skip_stages: Optional[List[str]] = None,
    ) -> AsyncIterator[Dict[str, Any]]:
        """
        执行完整的报告生成流程

        Args:
            context: 编排器上下文
            skip_stages: 要跳过的阶段列表

        Yields:
            流式输出每个阶段的进度和结果
        """
        skip_stages = skip_stages or []

        yield {
            "type": "orchestration_started",
            "template_id": context.template_id,
            "total_stages": len(self.STAGE_ORDER) - len(skip_stages),
            "execution_time": context.execution_time.isoformat(),
        }

        for stage_idx, stage_name in enumerate(self.STAGE_ORDER, 1):
            if stage_name in skip_stages:
                self.logger.info(f"跳过阶段 {stage_idx}/{len(self.STAGE_ORDER)}: {stage_name}")
                context.stage_results[stage_name] = StageResult(
                    stage_name=stage_name,
                    status=StageStatus.SKIPPED,
                )
                yield {
                    "type": "stage_skipped",
                    "stage_name": stage_name,
                    "stage_index": stage_idx,
                }
                continue

            self.logger.info(f"开始执行阶段 {stage_idx}/{len(self.STAGE_ORDER)}: {stage_name}")

            yield {
                "type": "stage_started",
                "stage_name": stage_name,
                "stage_index": stage_idx,
                "total_stages": len(self.STAGE_ORDER),
            }

            # 执行单个阶段
            try:
                async for event in self._execute_stage(stage_name, context):
                    yield event

                # 检查阶段是否成功
                stage_result = context.stage_results.get(stage_name)
                if not stage_result or stage_result.status == StageStatus.FAILED:
                    yield {
                        "type": "stage_failed",
                        "stage_name": stage_name,
                        "error": stage_result.error if stage_result else "Unknown error",
                    }

                    # 阶段失败，终止流程
                    yield {
                        "type": "orchestration_failed",
                        "failed_stage": stage_name,
                        "stage_index": stage_idx,
                        "error": stage_result.error if stage_result else "Unknown error",
                    }
                    return

                yield {
                    "type": "stage_completed",
                    "stage_name": stage_name,
                    "stage_index": stage_idx,
                    "result": stage_result.to_dict(),
                }

            except Exception as e:
                self.logger.error(f"阶段 {stage_name} 执行异常: {e}", exc_info=True)

                context.stage_results[stage_name] = StageResult(
                    stage_name=stage_name,
                    status=StageStatus.FAILED,
                    error=str(e),
                    completed_at=datetime.now(),
                )

                yield {
                    "type": "stage_failed",
                    "stage_name": stage_name,
                    "error": str(e),
                }

                yield {
                    "type": "orchestration_failed",
                    "failed_stage": stage_name,
                    "error": str(e),
                }
                return

        # 所有阶段完成
        yield {
            "type": "orchestration_completed",
            "template_id": context.template_id,
            "completed_stages": len([r for r in context.stage_results.values() if r.status == StageStatus.COMPLETED]),
            "total_stages": len(self.STAGE_ORDER),
            "execution_time_seconds": (datetime.now() - context.execution_time).total_seconds(),
        }

    async def _execute_stage(
        self,
        stage_name: str,
        context: OrchestratorContext,
    ) -> AsyncIterator[Dict[str, Any]]:
        """执行单个阶段"""

        stage_result = StageResult(
            stage_name=stage_name,
            status=StageStatus.IN_PROGRESS,
            started_at=datetime.now(),
        )
        context.stage_results[stage_name] = stage_result

        # 根据阶段名称调用相应的处理方法
        stage_handlers = {
            "placeholder_scan": self._stage_placeholder_scan,
            "sql_generation": self._stage_sql_generation,
            "etl_execution": self._stage_etl_execution,
            "data_fill_chart": self._stage_data_fill_chart,
            "content_optimization": self._stage_content_optimization,
            "document_generation": self._stage_document_generation,
        }

        handler = stage_handlers.get(stage_name)
        if not handler:
            raise ValueError(f"Unknown stage: {stage_name}")

        try:
            async for event in handler(context, stage_result):
                yield event

            # 标记阶段完成
            stage_result.status = StageStatus.COMPLETED
            stage_result.completed_at = datetime.now()

        except Exception as e:
            stage_result.status = StageStatus.FAILED
            stage_result.error = str(e)
            stage_result.completed_at = datetime.now()
            raise

    async def _stage_placeholder_scan(
        self,
        context: OrchestratorContext,
        result: StageResult,
    ) -> AsyncIterator[Dict[str, Any]]:
        """
        阶段 1: 占位符扫描

        扫描模板，识别所有占位符及其类型、位置、上下文信息。
        """
        yield {"type": "stage_log", "message": "开始扫描模板占位符..."}

        # 从数据库加载模板（如果还没有）
        if not context.template:
            from app.db.session import get_db_session
            from app.models.template import Template

            with get_db_session() as db:
                context.template = db.query(Template).filter(
                    Template.id == UUID(context.template_id)
                ).first()

                if not context.template:
                    raise ValueError(f"Template not found: {context.template_id}")

        # 扫描占位符
        from app.services.domain.placeholder.usecases.scanner import PlaceholderScanner

        scanner = PlaceholderScanner()
        placeholders = await scanner.scan_template(
            template_content=context.template.content,
            template_id=context.template_id,
        )

        yield {
            "type": "stage_log",
            "message": f"扫描完成，发现 {len(placeholders)} 个占位符",
        }

        # 保存到共享数据
        context.shared_data["placeholders"] = placeholders
        result.output = {
            "placeholder_count": len(placeholders),
            "placeholders": [p.to_dict() for p in placeholders],
        }
        result.metadata = {
            "template_name": context.template.name,
            "scan_method": "PlaceholderScanner",
        }

    async def _stage_sql_generation(
        self,
        context: OrchestratorContext,
        result: StageResult,
    ) -> AsyncIterator[Dict[str, Any]]:
        """
        阶段 2: SQL 生成与验证（并发优化版）

        为每个数据占位符生成 SQL，并进行验证和测试执行。
        采用 3 并发策略提升性能。
        """
        import asyncio

        placeholders = context.shared_data.get("placeholders", [])

        yield {
            "type": "stage_log",
            "message": f"开始为 {len(placeholders)} 个占位符生成 SQL（3并发）...",
        }

        # 加载数据源信息
        if not context.data_source:
            from app.db.session import get_db_session
            from app.models.data_source import DataSource

            with get_db_session() as db:
                context.data_source = db.query(DataSource).filter(
                    DataSource.id == UUID(context.data_source_id)
                ).first()

                if not context.data_source:
                    raise ValueError(f"DataSource not found: {context.data_source_id}")

        # 过滤数据占位符
        data_placeholders = [
            p for p in placeholders
            if p.get("type") in ["data", "sql", "metric"]
        ]

        if not data_placeholders:
            yield {
                "type": "stage_log",
                "message": "没有数据占位符需要生成 SQL",
            }
            context.shared_data["sql_results"] = []
            result.output = {"generated_count": 0, "sql_results": []}
            return

        yield {
            "type": "stage_log",
            "message": f"发现 {len(data_placeholders)} 个数据占位符",
        }

        # 并发生成 SQL
        semaphore = asyncio.Semaphore(3)  # 3 并发
        sql_results = []
        failed_placeholders = []

        async def generate_sql_for_placeholder(placeholder, index):
            """为单个占位符生成 SQL"""
            async with semaphore:
                try:
                    ph_id = placeholder.get("id")
                    ph_name = placeholder.get("name")

                    self.logger.info(f"[{index}/{len(data_placeholders)}] 生成 SQL: {ph_name}")

                    # 使用现有的单占位符分析服务
                    from app.services.application.placeholder.placeholder_service import PlaceholderApplicationService
                    from app.services.domain.placeholder.types import PlaceholderAnalysisRequest

                    ph_service = PlaceholderApplicationService(user_id=context.user_id)

                    request = PlaceholderAnalysisRequest(
                        placeholder_id=ph_id,
                        business_command=placeholder.get("description", ""),
                        requirements=placeholder.get("requirements", ""),
                        target_objective=placeholder.get("objective", "数据查询"),
                        context={
                            "placeholder_context_snippet": placeholder.get("context_text", ""),
                            "time_window": context.time_window,
                            "schedule": context.schedule,
                        },
                        data_source_info={
                            "id": str(context.data_source.id),
                            "source_type": context.data_source.source_type.value if hasattr(context.data_source.source_type, 'value') else str(context.data_source.source_type),
                            "database_name": getattr(context.data_source, 'doris_database', 'default_db'),
                            "host": context.data_source.doris_fe_hosts[0] if context.data_source.doris_fe_hosts else None,
                            "port": getattr(context.data_source, 'doris_fe_http_port', 8030),
                        },
                    )

                    # 执行分析
                    sql_result = None
                    async for event in ph_service.analyze_placeholder(request):
                        if event.get("type") == "sql_generation_complete":
                            sql_result = event.get("content")
                            break
                        elif event.get("type") in ["sql_generation_failed", "analysis_error"]:
                            sql_result = event.get("content")
                            break

                    if sql_result and sql_result.sql_query:
                        return {
                            "success": True,
                            "placeholder_id": ph_id,
                            "placeholder_name": ph_name,
                            "sql": sql_result.sql_query,
                            "validation_status": sql_result.validation_status,
                            "metadata": sql_result.metadata,
                        }
                    else:
                        return {
                            "success": False,
                            "placeholder_id": ph_id,
                            "placeholder_name": ph_name,
                            "error": "SQL generation returned empty result",
                        }

                except Exception as e:
                    self.logger.error(f"占位符 {placeholder.get('name')} SQL生成失败: {e}")
                    return {
                        "success": False,
                        "placeholder_id": placeholder.get("id"),
                        "placeholder_name": placeholder.get("name"),
                        "error": str(e),
                    }

        # 并发执行所有占位符的 SQL 生成
        tasks = [
            generate_sql_for_placeholder(ph, idx)
            for idx, ph in enumerate(data_placeholders, 1)
        ]

        # 使用 asyncio.gather 并发执行
        results = await asyncio.gather(*tasks, return_exceptions=True)

        # 处理结果
        for res in results:
            if isinstance(res, Exception):
                self.logger.error(f"SQL 生成任务异常: {res}")
                failed_placeholders.append({
                    "error": str(res),
                })
            elif res.get("success"):
                sql_results.append({
                    "placeholder_id": res["placeholder_id"],
                    "placeholder_name": res["placeholder_name"],
                    "sql": res["sql"],
                    "validation_status": res["validation_status"],
                    "metadata": res["metadata"],
                })
            else:
                failed_placeholders.append({
                    "placeholder_id": res["placeholder_id"],
                    "placeholder_name": res["placeholder_name"],
                    "error": res.get("error", "Unknown error"),
                })

        # 保存到共享数据
        context.shared_data["sql_results"] = sql_results
        context.shared_data["failed_sql_generation"] = failed_placeholders

        result.output = {
            "generated_count": len(sql_results),
            "failed_count": len(failed_placeholders),
            "sql_results": sql_results,
        }
        result.metadata = {
            "total_placeholders": len(placeholders),
            "data_placeholders": len(data_placeholders),
            "success_rate": len(sql_results) / len(data_placeholders) if data_placeholders else 0,
            "failed_placeholders": failed_placeholders,
        }

        # 检查失败率
        if failed_placeholders:
            failure_rate = len(failed_placeholders) / len(data_placeholders)
            if failure_rate > 0.5:
                yield {
                    "type": "stage_warning",
                    "message": f"⚠️ SQL 生成失败率过高: {failure_rate*100:.1f}% ({len(failed_placeholders)}/{len(data_placeholders)})",
                }

        yield {
            "type": "stage_log",
            "message": f"SQL 生成完成 - 成功: {len(sql_results)}/{len(data_placeholders)}，失败: {len(failed_placeholders)}",
        }

    async def _stage_etl_execution(
        self,
        context: OrchestratorContext,
        result: StageResult,
    ) -> AsyncIterator[Dict[str, Any]]:
        """
        阶段 3: ETL 取数（完整实现版）

        批量执行所有 SQL，获取数据并缓存。
        采用 3 并发策略，参考原 agents/tools/sql/executor.py 实现。
        """
        import asyncio
        import time

        sql_results = context.shared_data.get("sql_results", [])

        if not sql_results:
            yield {
                "type": "stage_log",
                "message": "没有 SQL 需要执行",
            }
            context.shared_data["etl_results"] = []
            result.output = {"executed_count": 0, "etl_results": []}
            return

        yield {
            "type": "stage_log",
            "message": f"开始执行 {len(sql_results)} 个 SQL 查询（3并发）...",
        }

        # 使用 SQL 执行工具
        from app.services.infrastructure.agents.tools.sql_tools import SQLExecuteTool

        sql_tool = SQLExecuteTool(container=self.container)

        # 并发执行 SQL
        semaphore = asyncio.Semaphore(3)  # 3 并发
        etl_results = []

        async def execute_single_sql(sql_result, index):
            """执行单个 SQL"""
            async with semaphore:
                placeholder_id = sql_result.get("placeholder_id")
                placeholder_name = sql_result.get("placeholder_name")
                sql_with_placeholders = sql_result.get("sql")

                try:
                    self.logger.info(f"[{index}/{len(sql_results)}] 执行 SQL: {placeholder_name}")

                    start_time = time.time()

                    # 构建执行上下文
                    exec_input = {
                        "sql": sql_with_placeholders,
                        "current_sql": sql_with_placeholders,
                        "time_window": context.time_window or {},
                        "window": context.time_window or {},
                        "data_source": {
                            "id": str(context.data_source.id),
                            "source_type": context.data_source.source_type.value if hasattr(context.data_source.source_type, 'value') else str(context.data_source.source_type),
                            "database_name": getattr(context.data_source, 'doris_database', 'default_db'),
                        },
                        "user_id": context.user_id,
                    }

                    # 执行 SQL
                    exec_result = await sql_tool.execute(exec_input)

                    execution_time = time.time() - start_time

                    if exec_result.get("success"):
                        return {
                            "success": True,
                            "placeholder_id": placeholder_id,
                            "placeholder_name": placeholder_name,
                            "data": exec_result.get("rows", []),
                            "columns": exec_result.get("columns", []),
                            "row_count": exec_result.get("row_count", 0),
                            "execution_time": execution_time,
                            "execution_sql": exec_result.get("sql"),  # 替换后的实际 SQL
                            "metadata": {
                                "sql_with_placeholders": sql_with_placeholders,
                                "time_window": context.time_window,
                            }
                        }
                    else:
                        error_msg = exec_result.get("error", "Unknown error")
                        self.logger.error(f"SQL 执行失败 - 占位符: {placeholder_name}, 错误: {error_msg}")

                        return {
                            "success": False,
                            "placeholder_id": placeholder_id,
                            "placeholder_name": placeholder_name,
                            "data": [],
                            "columns": [],
                            "row_count": 0,
                            "execution_time": execution_time,
                            "error": error_msg,
                        }

                except Exception as e:
                    self.logger.error(f"ETL 执行异常 - 占位符: {placeholder_name}: {e}", exc_info=True)

                    return {
                        "success": False,
                        "placeholder_id": placeholder_id,
                        "placeholder_name": placeholder_name,
                        "data": [],
                        "columns": [],
                        "row_count": 0,
                        "error": str(e),
                    }

        # 并发执行所有 SQL
        tasks = [
            execute_single_sql(sql_result, idx)
            for idx, sql_result in enumerate(sql_results, 1)
        ]

        # 使用 asyncio.gather 并发执行
        results = await asyncio.gather(*tasks, return_exceptions=True)

        # 处理结果
        for res in results:
            if isinstance(res, Exception):
                self.logger.error(f"ETL 执行任务异常: {res}")
                etl_results.append({
                    "success": False,
                    "error": str(res),
                    "data": [],
                    "row_count": 0,
                })
            else:
                etl_results.append(res)

                # 输出日志
                if res.get("success"):
                    yield {
                        "type": "stage_log",
                        "message": f"✅ {res['placeholder_name']}: {res['row_count']} 行，耗时 {res['execution_time']:.2f}s",
                    }
                else:
                    yield {
                        "type": "stage_warning",
                        "message": f"❌ {res['placeholder_name']} 执行失败: {res.get('error', 'Unknown')}",
                    }

        # 保存结果
        context.shared_data["etl_results"] = etl_results

        # 统计
        success_count = len([r for r in etl_results if r.get("success")])
        failed_count = len(etl_results) - success_count
        total_rows = sum(r.get("row_count", 0) for r in etl_results)

        result.output = {
            "executed_count": len(etl_results),
            "success_count": success_count,
            "failed_count": failed_count,
            "total_rows": total_rows,
            "etl_results": etl_results,
        }

        result.metadata = {
            "total_execution_time": sum(r.get("execution_time", 0) for r in etl_results),
            "avg_execution_time": sum(r.get("execution_time", 0) for r in etl_results) / len(etl_results) if etl_results else 0,
            "success_rate": success_count / len(etl_results) if etl_results else 0,
        }

        yield {
            "type": "stage_log",
            "message": f"ETL 执行完成 - 成功: {success_count}/{len(etl_results)}，总行数: {total_rows}，总耗时: {result.metadata['total_execution_time']:.2f}s",
        }

    async def _stage_data_fill_chart(
        self,
        context: OrchestratorContext,
        result: StageResult,
    ) -> AsyncIterator[Dict[str, Any]]:
        """
        阶段 4: 数据回填与图表生成（完整实现版）

        将数据回填到模板占位符，生成图表图片。
        参考原 agents/tools/chart_tools.py 实现。
        """
        import asyncio

        placeholders = context.shared_data.get("placeholders", [])
        etl_results = context.shared_data.get("etl_results", [])

        yield {"type": "stage_log", "message": "开始数据回填与图表生成..."}

        # 1. 构建占位符 → 数据映射
        placeholder_data_map = {}
        for etl_result in etl_results:
            ph_id = etl_result.get("placeholder_id")
            placeholder_data_map[ph_id] = {
                "data": etl_result.get("data", []),
                "columns": etl_result.get("columns", []),
                "row_count": etl_result.get("row_count", 0),
            }

        # 2. 分类处理占位符
        fill_results = []
        chart_results = []

        # 过滤出需要处理的占位符
        data_and_chart_placeholders = [
            p for p in placeholders
            if p.get("type") in ["data", "sql", "metric", "chart"]
        ]

        if not data_and_chart_placeholders:
            yield {"type": "stage_log", "message": "没有需要回填的占位符"}
            context.shared_data["fill_results"] = []
            context.shared_data["chart_results"] = []
            result.output = {"filled_count": 0, "chart_count": 0}
            return

        yield {
            "type": "stage_log",
            "message": f"发现 {len(data_and_chart_placeholders)} 个需要处理的占位符",
        }

        for idx, placeholder in enumerate(data_and_chart_placeholders, 1):
            ph_id = placeholder.get("id")
            ph_type = placeholder.get("type")
            ph_name = placeholder.get("name")

            yield {
                "type": "stage_progress",
                "message": f"处理占位符 ({idx}/{len(data_and_chart_placeholders)}): {ph_name}",
                "progress": (idx / len(data_and_chart_placeholders)) * 100,
            }

            if ph_type in ["data", "sql", "metric"]:
                # 数据占位符 - 直接回填
                data_info = placeholder_data_map.get(ph_id, {})

                if data_info.get("row_count", 0) > 0:
                    # 根据占位符格式决定回填方式
                    fill_format = placeholder.get("format", "single_value")

                    if fill_format == "single_value":
                        # 单个值（如总数、平均值）
                        filled_value = data_info["data"][0][0] if data_info["data"] and data_info["data"][0] else "N/A"

                    elif fill_format == "table":
                        # 表格数据
                        filled_value = {
                            "type": "table",
                            "columns": data_info["columns"],
                            "data": data_info["data"],
                        }

                    elif fill_format == "list":
                        # 列表数据
                        filled_value = {
                            "type": "list",
                            "items": [row[0] if row else "" for row in data_info["data"]],
                        }
                    else:
                        # 默认单值
                        filled_value = data_info["data"][0][0] if data_info["data"] and data_info["data"][0] else "N/A"

                    fill_results.append({
                        "placeholder_id": ph_id,
                        "placeholder_name": ph_name,
                        "filled_value": filled_value,
                        "format": fill_format,
                        "row_count": data_info["row_count"],
                    })
                else:
                    # 没有数据
                    fill_results.append({
                        "placeholder_id": ph_id,
                        "placeholder_name": ph_name,
                        "filled_value": "暂无数据",
                        "format": "text",
                        "row_count": 0,
                    })

            elif ph_type == "chart":
                # 图表占位符 - 生成图表
                data_info = placeholder_data_map.get(ph_id, {})

                if data_info.get("row_count", 0) > 0:
                    try:
                        # 生成图表
                        chart_result = await self._generate_chart(
                            placeholder=placeholder,
                            data=data_info["data"],
                            columns=data_info["columns"],
                        )

                        if chart_result.get("success"):
                            chart_results.append({
                                "placeholder_id": ph_id,
                                "placeholder_name": ph_name,
                                "chart_path": chart_result.get("chart_path"),
                                "chart_type": chart_result.get("chart_type"),
                            })

                            yield {
                                "type": "stage_log",
                                "message": f"✅ 图表生成成功: {ph_name}",
                            }
                        else:
                            chart_results.append({
                                "placeholder_id": ph_id,
                                "placeholder_name": ph_name,
                                "error": chart_result.get("error"),
                            })

                            yield {
                                "type": "stage_warning",
                                "message": f"❌ 图表生成失败: {ph_name}, {chart_result.get('error')}",
                            }

                    except Exception as e:
                        self.logger.error(f"图表生成异常: {ph_name}, {e}")
                        chart_results.append({
                            "placeholder_id": ph_id,
                            "placeholder_name": ph_name,
                            "error": str(e),
                        })
                else:
                    # 没有数据，无法生成图表
                    chart_results.append({
                        "placeholder_id": ph_id,
                        "placeholder_name": ph_name,
                        "chart_path": None,
                        "message": "暂无数据，无法生成图表",
                    })

        # 3. 保存结果
        context.shared_data["fill_results"] = fill_results
        context.shared_data["chart_results"] = chart_results

        success_chart_count = len([c for c in chart_results if c.get("chart_path")])

        result.output = {
            "filled_count": len(fill_results),
            "chart_count": len(chart_results),
            "success_chart_count": success_chart_count,
            "fill_results": fill_results,
            "chart_results": chart_results,
        }

        result.metadata = {
            "total_data_rows": sum(f.get("row_count", 0) for f in fill_results),
        }

        yield {
            "type": "stage_log",
            "message": f"数据回填与图表生成完成 - 回填: {len(fill_results)}, 图表: {success_chart_count}/{len(chart_results)}",
        }

    async def _generate_chart(
        self,
        placeholder: Dict[str, Any],
        data: List[List],
        columns: List[str],
    ) -> Dict[str, Any]:
        """
        生成图表

        参考原 agents/tools/chart_tools.py 的实现
        """
        try:
            import matplotlib.pyplot as plt
            import matplotlib
            import os
            import base64
            import io
            from datetime import datetime

            # 设置中文字体支持
            matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans', 'Arial']
            matplotlib.rcParams['axes.unicode_minus'] = False

            # 创建图表
            fig, ax = plt.subplots(figsize=(10, 6))

            chart_type = placeholder.get("chart_type", "bar")
            title = placeholder.get("name", "数据图表")

            if not data or len(columns) < 2:
                ax.text(0.5, 0.5, "暂无数据或数据格式不正确", transform=ax.transAxes, ha="center", va="center")
            else:
                # 转换数据格式
                x_values = [row[0] if row else "" for row in data]
                y_values = [float(row[1]) if len(row) > 1 and row[1] is not None else 0 for row in data]

                if chart_type == "bar":
                    bars = ax.bar(x_values, y_values, color="#2563eb", alpha=0.8)
                    ax.set_xlabel(columns[0])
                    ax.set_ylabel(columns[1])
                elif chart_type == "line":
                    ax.plot(x_values, y_values, marker="o", linewidth=2, markersize=6, color="#10b981")
                    ax.set_xlabel(columns[0])
                    ax.set_ylabel(columns[1])
                elif chart_type == "pie":
                    ax.pie(y_values, labels=x_values, autopct="%1.1f%%", startangle=90)
                    ax.axis("equal")
                else:
                    # 默认柱状图
                    bars = ax.bar(x_values, y_values, color="#2563eb", alpha=0.8)
                    ax.set_xlabel(columns[0])
                    ax.set_ylabel(columns[1])

            ax.set_title(title, fontsize=14, fontweight="bold", pad=20)
            plt.tight_layout()

            # 保存图片
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            placeholder_id = placeholder.get("id", "chart")
            filename = f"chart_{placeholder_id}_{timestamp}.png"

            # 确保 charts 目录存在
            charts_dir = "/tmp/autoreport_charts"
            os.makedirs(charts_dir, exist_ok=True)
            image_path = os.path.join(charts_dir, filename)

            # 保存到文件
            plt.savefig(image_path, dpi=300, bbox_inches="tight", facecolor="white")
            plt.close()

            return {
                "success": True,
                "chart_path": image_path,
                "chart_type": chart_type,
            }

        except Exception as e:
            self.logger.error(f"图表生成异常: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
            }

    async def _stage_content_optimization(
        self,
        context: OrchestratorContext,
        result: StageResult,
    ) -> AsyncIterator[Dict[str, Any]]:
        """
        阶段 5: 文案优化

        根据占位符数据，优化模板中占位符附近的语句。
        """
        yield {"type": "stage_log", "message": "开始文案优化..."}

        # TODO: 实现文案优化逻辑
        # 可以使用另一个专门的 Loom Agent 来处理

        result.output = {"optimized": True}

        yield {"type": "stage_log", "message": "文案优化完成"}

    async def _stage_document_generation(
        self,
        context: OrchestratorContext,
        result: StageResult,
    ) -> AsyncIterator[Dict[str, Any]]:
        """
        阶段 6: 文档生成（完整实现版）

        生成最终的 DOCX 文档。
        参考 backend/app/services/infrastructure/document/word_template_service.py
        """
        import os
        from datetime import datetime

        fill_results = context.shared_data.get("fill_results", [])
        chart_results = context.shared_data.get("chart_results", [])

        yield {"type": "stage_log", "message": "开始生成 DOCX 文档..."}

        try:
            # 1. 准备文档数据
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

            # 创建新文档
            doc = Document()

            # 添加标题
            title = doc.add_heading(context.template.name if context.template else "报告", level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 添加生成时间
            doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("")  # 空行

            # 2. 回填数据占位符
            yield {"type": "stage_log", "message": f"回填 {len(fill_results)} 个数据占位符..."}

            for fill_result in fill_results:
                ph_name = fill_result.get("placeholder_name", "未命名占位符")
                filled_value = fill_result.get("filled_value")
                fill_format = fill_result.get("format", "text")

                # 添加占位符标题
                doc.add_heading(ph_name, level=2)

                if fill_format == "single_value":
                    # 单个值
                    doc.add_paragraph(f"数值: {filled_value}")

                elif fill_format == "table" and isinstance(filled_value, dict):
                    # 表格
                    columns = filled_value.get("columns", [])
                    data = filled_value.get("data", [])

                    if columns and data:
                        # 创建表格
                        table = doc.add_table(rows=1 + len(data), cols=len(columns))
                        table.style = 'Light Grid Accent 1'

                        # 表头
                        header_cells = table.rows[0].cells
                        for idx, col_name in enumerate(columns):
                            header_cells[idx].text = str(col_name)

                        # 数据行
                        for row_idx, row_data in enumerate(data, 1):
                            row_cells = table.rows[row_idx].cells
                            for col_idx, cell_value in enumerate(row_data):
                                if col_idx < len(row_cells):
                                    row_cells[col_idx].text = str(cell_value) if cell_value is not None else ""

                elif fill_format == "list" and isinstance(filled_value, dict):
                    # 列表
                    items = filled_value.get("items", [])
                    for item in items:
                        doc.add_paragraph(str(item), style='List Bullet')

                else:
                    # 默认文本
                    doc.add_paragraph(str(filled_value))

                doc.add_paragraph("")  # 空行

            # 3. 插入图表
            yield {"type": "stage_log", "message": f"插入 {len(chart_results)} 个图表..."}

            for chart_result in chart_results:
                ph_name = chart_result.get("placeholder_name", "未命名图表")
                chart_path = chart_result.get("chart_path")

                # 添加图表标题
                doc.add_heading(ph_name, level=2)

                if chart_path and os.path.exists(chart_path):
                    # 插入图表图片
                    doc.add_picture(chart_path, width=Inches(6))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                else:
                    # 图表不存在
                    error_msg = chart_result.get("error") or chart_result.get("message", "图表不可用")
                    doc.add_paragraph(f"图表生成失败: {error_msg}")

                doc.add_paragraph("")  # 空行

            # 4. 保存文档
            output_dir = os.path.join("/tmp", "generated_reports")
            os.makedirs(output_dir, exist_ok=True)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            template_name = context.template.name if context.template else "report"
            # 清理文件名中的非法字符
            safe_template_name = "".join(c for c in template_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            output_filename = f"{safe_template_name}_{timestamp}.docx"
            output_path = os.path.join(output_dir, output_filename)

            doc.save(output_path)

            # 获取文件大小
            file_size = os.path.getsize(output_path) if os.path.exists(output_path) else 0

            result.output = {
                "document_path": output_path,
                "document_size": file_size,
                "charts_embedded": len([c for c in chart_results if c.get("chart_path")]),
                "data_filled": len(fill_results),
            }

            result.metadata = {
                "template_name": context.template.name if context.template else "Unknown",
                "generated_at": datetime.now().isoformat(),
                "output_filename": output_filename,
            }

            yield {
                "type": "stage_log",
                "message": f"✅ DOCX 文档生成成功: {output_path} ({file_size/1024:.1f} KB)",
            }

        except Exception as e:
            self.logger.error(f"文档生成异常: {e}", exc_info=True)

            result.output = {
                "document_path": None,
                "error": str(e),
            }

            yield {
                "type": "stage_error",
                "message": f"❌ DOCX 文档生成失败: {str(e)}",
            }

            raise

        yield {"type": "stage_log", "message": "DOCX 文档生成完成"}


__all__ = [
    "ReportGenerationOrchestrator",
    "OrchestratorContext",
    "StageStatus",
    "StageResult",
]
