"""
Word文档导出服务

负责将报告数据导出为Word文档，支持：
1. 模板解析和占位符替换
2. 图表插入和布局
3. 统一字体和格式设置
4. 文档结构化处理
5. 多种导出格式支持
"""

import logging
import os
import asyncio
from typing import Dict, List, Any, Optional, Union, Tuple
from datetime import datetime
from dataclasses import dataclass
from enum import Enum
import json

logger = logging.getLogger(__name__)


class DocumentFormat(Enum):
    """文档格式"""
    DOCX = "docx"
    PDF = "pdf"
    HTML = "html"


@dataclass
class DocumentConfig:
    """文档配置"""
    template_id: str
    output_format: DocumentFormat = DocumentFormat.DOCX
    font_name: str = "宋体"
    font_size: int = 12
    line_spacing: float = 1.5
    margins: Dict[str, float] = None
    header: Optional[str] = None
    footer: Optional[str] = None
    watermark: Optional[str] = None

    # 图表配置
    max_chart_width_inches: float = 6.5  # 最大图表宽度（英寸）
    chart_dpi: int = 150  # 图表分辨率
    chart_alignment: str = "center"  # 图表对齐方式

    def __post_init__(self):
        if self.margins is None:
            self.margins = {"top": 2.54, "bottom": 2.54, "left": 3.17, "right": 3.17}  # cm


@dataclass
class DocumentExportResult:
    """文档导出结果"""
    success: bool
    document_path: Optional[str] = None
    file_size_bytes: int = 0
    page_count: int = 0
    export_time_seconds: float = 0.0
    error: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None


class WordExportService:
    """
    Word文档导出服务
    
    提供完整的Word文档生成功能：
    1. 解析模板文档
    2. 替换占位符内容
    3. 插入图表和图片
    4. 应用统一格式
    5. 导出为多种格式
    """
    
    def __init__(self, user_id: str):
        if not user_id:
            raise ValueError("user_id is required for WordExportService")
        self.user_id = user_id
        self.output_dir = f"/tmp/documents/{user_id}"
        os.makedirs(self.output_dir, exist_ok=True)
        
        # 确保导入必要的库
        self._ensure_document_dependencies()
    
    def _ensure_document_dependencies(self):
        """确保文档处理依赖库可用"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            self.Document = Document
            self.Inches = Inches
            self.Pt = Pt
            self.WD_PARAGRAPH_ALIGNMENT = WD_PARAGRAPH_ALIGNMENT
            logger.info("Word文档处理依赖库加载成功")
        except ImportError as e:
            logger.warning(f"Word文档处理依赖库加载失败: {e}")
            self.Document = None
    
    async def export_report_document(
        self,
        template_id: str,
        placeholder_data: Dict[str, Any],
        etl_data: Dict[str, Any],
        chart_data: Dict[str, Any],
        config: Optional[DocumentConfig] = None
    ) -> DocumentExportResult:
        """
        导出报告文档
        
        Args:
            template_id: 模板ID
            placeholder_data: 占位符数据
            etl_data: ETL处理数据
            chart_data: 图表数据
            config: 文档配置
            
        Returns:
            文档导出结果
        """
        start_time = datetime.now()
        
        logger.info(f"开始导出报告文档: template_id={template_id}")
        
        if not self.Document:
            return DocumentExportResult(
                success=False,
                error="Word文档处理库不可用"
            )
        
        try:
            # 使用默认配置
            if config is None:
                config = DocumentConfig(template_id=template_id)
            
            # 1. 获取模板文档
            template_doc = await self._load_template_document(template_id)
            
            # 2. 处理占位符替换
            await self._replace_placeholders(template_doc, placeholder_data, etl_data)

            # 3. 插入图表
            await self._insert_charts(template_doc, chart_data, config)

            # 4. 应用文档格式
            self._apply_document_formatting(template_doc, config)
            
            # 5. 保存文档
            document_path = await self._save_document(template_doc, template_id, config.output_format)
            
            end_time = datetime.now()
            export_time = (end_time - start_time).total_seconds()
            
            # 获取文档信息
            file_size = os.path.getsize(document_path) if document_path and os.path.exists(document_path) else 0
            # 简单计算段落数作为页数估算基础（避免xpath的namespaces兼容性问题）
            page_count = len(template_doc.paragraphs)
            
            logger.info(f"文档导出完成: {document_path}, 大小: {file_size} bytes, 页数: {page_count}")
            
            return DocumentExportResult(
                success=True,
                document_path=document_path,
                file_size_bytes=file_size,
                page_count=max(page_count // 25, 1),  # 粗略估计页数
                export_time_seconds=export_time,
                metadata={
                    "template_id": template_id,
                    "output_format": config.output_format.value,
                    "font_name": config.font_name,
                    "export_timestamp": end_time.isoformat()
                }
            )
            
        except Exception as e:
            logger.error(f"文档导出失败: {e}")
            return DocumentExportResult(
                success=False,
                error=str(e)
            )
    
    async def _load_template_document(self, template_id: str):
        """加载模板文档"""
        db = None
        try:
            from app.crud import template as crud_template
            from app.db.session import SessionLocal
            
            db = SessionLocal()
            template = crud_template.get(db, id=template_id)
            if not template:
                logger.error(f"模板不存在: {template_id}")
                raise ValueError(f"模板不存在: {template_id}")
            
            # 如果有模板文件路径，加载文档
            if hasattr(template, 'file_path') and template.file_path:
                if os.path.exists(template.file_path):
                    logger.info(f"从文件加载模板: {template.file_path}")
                    return self.Document(template.file_path)
                else:
                    logger.warning(f"模板文件不存在: {template.file_path}，将创建新文档")
            
            # 否则创建新文档并使用模板内容
            doc = self.Document()
            
            # 添加标题
            title = doc.add_heading(template.name or '报告', 0)
            title.alignment = self.WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加模板内容
            if template.content:
                # 按段落分割内容，保留空行作为段落间距
                paragraphs = template.content.split('\n')
                for paragraph_text in paragraphs:
                    if paragraph_text.strip():
                        doc.add_paragraph(paragraph_text.strip())
                    else:
                        # 空行作为段落间距
                        doc.add_paragraph()
            
            logger.info(f"创建新文档模板，段落数: {len(doc.paragraphs)}")
            return doc
                
        except Exception as e:
            logger.error(f"加载模板文档失败: {e}", exc_info=True)
            # 返回空文档作为后备
            return self.Document()
        finally:
            if db is not None:
                try:
                    db.close()
                except Exception as close_error:
                    logger.warning(f"关闭数据库会话失败: {close_error}")
    
    async def _replace_placeholders(
        self,
        doc,
        placeholder_data: Dict[str, Any],
        etl_data: Dict[str, Any]
    ):
        """替换文档中的占位符（智能改写，保持格式）"""
        try:
            # 构建替换映射
            replacement_map = {}
            # 🆕 记录需要智能改写的占位符（值为0或null）
            smart_rewrite_needed = {}

            # 优先使用调用方提供的直接替换值
            direct_values = placeholder_data.get('direct_values') if isinstance(placeholder_data, dict) else None
            if isinstance(direct_values, dict):
                for name, value in direct_values.items():
                    try:
                        # 检查是否需要智能改写
                        if value is None or value == 0 or value == "":
                            smart_rewrite_needed[name] = value
                            logger.info(f"🔧 占位符 {name} 值为空/0，标记为需要智能改写")
                        else:
                            # 数据类型验证和转换
                            text_value = self._validate_and_convert_placeholder_value(name, value)
                            if text_value is not None:
                                replacement_map[f"{{{{{name}}}}}"] = text_value
                                replacement_map[f"{{{name}}}"] = text_value
                    except Exception as e:
                        logger.warning(f"占位符 {name} 数据验证失败: {e}")
                        continue

            # 回退到从ETL数据推断
            validation_results = placeholder_data.get('validation_results', []) if isinstance(placeholder_data, dict) else []
            for result in validation_results:
                placeholder_name = result.get('placeholder_name', '')
                key_braced = f"{{{{{placeholder_name}}}}}"
                key_single = f"{{{placeholder_name}}}"
                if placeholder_name and key_braced not in replacement_map and key_single not in replacement_map:
                    replacement_value = self._get_placeholder_value(result, etl_data)
                    # 检查是否需要智能改写
                    if replacement_value == "[数据为空]" or replacement_value == "0" or not replacement_value:
                        smart_rewrite_needed[placeholder_name] = replacement_value
                        logger.info(f"🔧 占位符 {placeholder_name} 值为空，标记为需要智能改写")
                    else:
                        replacement_map[key_braced] = replacement_value
                        replacement_map[key_single] = replacement_value

            # 🆕 步骤1：智能改写包含空值占位符的句子
            if smart_rewrite_needed:
                logger.info(f"🤖 开始智能改写 {len(smart_rewrite_needed)} 个占位符所在的句子")
                await self._smart_rewrite_empty_placeholders(doc, smart_rewrite_needed, placeholder_data)

            # 步骤2：常规占位符替换
            if not replacement_map:
                logger.warning("没有找到可替换的占位符（可能都已被智能改写）")
                return

            replaced_count = 0

            # 替换段落中的占位符（保持格式，支持多个占位符）
            for paragraph in doc.paragraphs:
                if not paragraph.runs:
                    continue

                for run in paragraph.runs:
                    original_text = run.text
                    if not original_text:
                        continue

                    modified_text = original_text
                    for placeholder, value in replacement_map.items():
                        if placeholder in modified_text:
                            modified_text = modified_text.replace(placeholder, str(value))
                            replaced_count += 1

                    if modified_text != original_text:
                        run.text = modified_text

            # 替换表格中的占位符（保持格式）
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if not cell.paragraphs:
                            continue

                        for paragraph in cell.paragraphs:
                            full_text = "".join(run.text for run in paragraph.runs)
                            if not full_text:
                                continue

                            for placeholder, value in replacement_map.items():
                                if placeholder in full_text:
                                    for run in paragraph.runs:
                                        if placeholder in run.text:
                                            run.text = run.text.replace(placeholder, str(value))
                                            replaced_count += 1
                                            break
                                    break

            logger.info(f"✅ 已替换 {replaced_count} 个占位符实例（智能改写: {len(smart_rewrite_needed)}，常规替换: {len(replacement_map)}）")

        except Exception as e:
            logger.error(f"替换占位符失败: {e}", exc_info=True)

    async def _smart_rewrite_empty_placeholders(
        self,
        doc,
        empty_placeholders: Dict[str, Any],
        placeholder_data: Dict[str, Any]
    ):
        """
        🆕 智能改写包含空值/0值占位符的句子

        当占位符的值为0、null或空时，不是直接显示"0"或"无数据"，
        而是调用LLM根据上下文改写整个句子，使其更自然。

        例如：
        - 原句："本月退货成功数量为{{退货数量}}件"
        - 值为0时：改为"本月暂无退货成功记录"
        - 值为null时：改为"本月退货数据暂未统计"
        """
        try:
            from app.services.infrastructure.agents import create_llm_service
            llm_service = create_llm_service(container=self.container if hasattr(self, 'container') else None)

            rewritten_count = 0

            # 遍历所有段落，找到包含空值占位符的句子
            for paragraph in doc.paragraphs:
                full_text = "".join(run.text for run in paragraph.runs)
                if not full_text:
                    continue

                # 检查段落是否包含任何空值占位符
                needs_rewrite = False
                matched_placeholders = []
                for placeholder_name, value in empty_placeholders.items():
                    if f"{{{{{placeholder_name}}}}}" in full_text or f"{{{placeholder_name}}}" in full_text:
                        needs_rewrite = True
                        matched_placeholders.append((placeholder_name, value))

                if not needs_rewrite:
                    continue

                # 调用LLM改写句子
                logger.info(f"🤖 改写句子: {full_text[:100]}...")

                # 构建改写prompt
                placeholder_info = "\n".join([
                    f"- {name}: 值为{repr(value)}（{'null' if value is None else '0' if value == 0 else '空'}）"
                    for name, value in matched_placeholders
                ])

                rewrite_prompt = f"""请改写以下句子，使其在数据为空/0/null时仍然自然流畅。

原句子：{full_text}

占位符信息：
{placeholder_info}

改写要求：
1. 如果数据为0，改为"暂无XXX"、"尚未有XXX"等表达
2. 如果数据为null/空，改为"数据暂未统计"、"信息暂缺"等表达
3. 保持句子的业务含义和语气
4. 删除数量单位（如"件"、"个"、"元"）
5. 语言简洁自然，符合中文表达习惯

请直接输出改写后的句子，不要包含任何解释或标记。"""

                try:
                    response = await llm_service.generate(
                        prompt=rewrite_prompt,
                        temperature=0.3,
                    )

                    rewritten_text = response.get("content", "").strip()

                    if rewritten_text and rewritten_text != full_text:
                        logger.info(f"✅ 改写成功: {rewritten_text[:100]}...")

                        # 替换段落内容（保持格式）
                        if paragraph.runs:
                            # 清空所有run的文本
                            for run in paragraph.runs[1:]:
                                run.text = ""
                            # 在第一个run中设置新文本（保留格式）
                            paragraph.runs[0].text = rewritten_text
                            rewritten_count += 1
                    else:
                        logger.warning(f"⚠️ LLM改写结果为空或未改变，保持原文")

                except Exception as llm_error:
                    logger.error(f"LLM改写失败: {llm_error}")
                    # 降级处理：使用默认文本替换
                    fallback_text = self._generate_fallback_text(full_text, matched_placeholders)
                    if paragraph.runs:
                        for run in paragraph.runs[1:]:
                            run.text = ""
                        paragraph.runs[0].text = fallback_text
                        logger.info(f"📝 使用降级文本: {fallback_text}")

            logger.info(f"🎉 智能改写完成，共改写 {rewritten_count} 个句子")

        except Exception as e:
            logger.error(f"智能改写过程异常: {e}", exc_info=True)

    def _generate_fallback_text(
        self,
        original_text: str,
        matched_placeholders: List[Tuple[str, Any]]
    ) -> str:
        """生成降级文本（当LLM不可用时）"""
        # 简单的规则替换
        for placeholder_name, value in matched_placeholders:
            # 移除占位符和周围的数字、单位
            import re
            pattern = rf'(为|是|有|共|达到?)\s*{{{{{placeholder_name}}}}}\s*([件个元项条笔])?'
            if value is None:
                original_text = re.sub(pattern, r'\1数据暂未统计', original_text)
            elif value == 0:
                original_text = re.sub(pattern, r'暂无相关记录', original_text)
            else:
                original_text = re.sub(pattern, r'\1信息暂缺', original_text)

            # 清理剩余的占位符
            original_text = original_text.replace(f"{{{{{placeholder_name}}}}}", "")
            original_text = original_text.replace(f"{{{placeholder_name}}}", "")

        return original_text.strip()

    def _validate_and_convert_placeholder_value(self, name: str, value: Any) -> Optional[str]:
        """
        验证和转换占位符值为字符串

        Args:
            name: 占位符名称
            value: 占位符值

        Returns:
            转换后的字符串，如果验证失败返回 None
        """
        try:
            # 跳过二进制数据
            if isinstance(value, (bytes, bytearray)):
                logger.warning(f"占位符 {name} 包含二进制数据，已跳过")
                return None

            # 跳过不可序列化的复杂对象
            if hasattr(value, '__dict__') and not hasattr(value, '__str__'):
                logger.warning(f"占位符 {name} 包含不可序列化对象，已跳过")
                return None

            # None 转换为空字符串
            if value is None:
                return ""

            # 列表和字典转换为简化表示
            if isinstance(value, (list, dict)):
                # 对于简单的列表/字典，可以转换
                if isinstance(value, list) and len(value) < 10:
                    return ", ".join(str(v) for v in value)
                elif isinstance(value, dict) and len(value) < 5:
                    return ", ".join(f"{k}: {v}" for k, v in value.items())
                else:
                    logger.warning(f"占位符 {name} 包含复杂数据结构，已简化")
                    return f"[包含 {len(value)} 项数据]"

            # 标准类型转换
            return str(value)

        except Exception as e:
            logger.warning(f"占位符 {name} 值转换失败: {e}")
            return None

    def _get_placeholder_value(
        self,
        placeholder_result: Dict[str, Any],
        etl_data: Dict[str, Any]
    ) -> str:
        """获取占位符的实际值"""
        try:
            placeholder_name = placeholder_result.get('placeholder_name', '')
            
            # 从ETL数据中查找对应的值
            for data_source_id, data_info in etl_data.items():
                transform_result = data_info.get('transform', {})
                if transform_result.get('success'):
                    data = transform_result.get('data', [])
                    if data:
                        # 简单的值提取逻辑
                        if isinstance(data, list) and data:
                            first_row = data[0]
                            if isinstance(first_row, dict):
                                # 查找匹配的字段
                                for key, value in first_row.items():
                                    if placeholder_name.lower() in key.lower():
                                        return str(value)
                                # 如果没有找到匹配字段，返回第一个数值字段
                                for key, value in first_row.items():
                                    if isinstance(value, (int, float)):
                                        return str(value)
            
            # 如果没有找到数据，返回占位符名称
            return f"[{placeholder_name}]"
            
        except Exception as e:
            logger.error(f"获取占位符值失败: {e}")
            return f"[{placeholder_name}]"
    
    async def _insert_charts(self, doc, chart_data: Dict[str, Any], config: DocumentConfig):
        """在文档中插入图表"""
        try:
            chart_results = chart_data.get('data', [])
            if not chart_results:
                logger.debug("没有图表数据需要插入")
                return

            inserted_count = 0
            for chart_result in chart_results:
                try:
                    if not chart_result.get('success'):
                        logger.debug(f"跳过未成功的图表: {chart_result.get('error', 'unknown')}")
                        continue

                    chart_path = chart_result.get('file_path')
                    if not chart_path or not os.path.exists(chart_path):
                        logger.warning(f"图表文件不存在: {chart_path}")
                        continue

                    # 验证文件大小
                    file_size = os.path.getsize(chart_path)
                    if file_size == 0:
                        logger.warning(f"图表文件为空: {chart_path}")
                        continue

                    # 验证文件类型
                    if not chart_path.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif')):
                        logger.warning(f"不支持的图表文件格式: {chart_path}")
                        continue

                    # 添加图表段落
                    chart_paragraph = doc.add_paragraph()

                    # 使用配置的对齐方式
                    alignment_map = {
                        "left": self.WD_PARAGRAPH_ALIGNMENT.LEFT,
                        "center": self.WD_PARAGRAPH_ALIGNMENT.CENTER,
                        "right": self.WD_PARAGRAPH_ALIGNMENT.RIGHT
                    }
                    chart_paragraph.alignment = alignment_map.get(
                        config.chart_alignment.lower(),
                        self.WD_PARAGRAPH_ALIGNMENT.CENTER
                    )

                    # 插入图片
                    try:
                        run = chart_paragraph.add_run()
                        # 使用配置的图表宽度
                        run.add_picture(chart_path, width=self.Inches(config.max_chart_width_inches))
                        logger.debug(f"成功插入图表: {chart_path} ({file_size} bytes, 宽度: {config.max_chart_width_inches}英寸)")
                    except Exception as pic_error:
                        logger.error(f"插入图片失败: {pic_error}", exc_info=True)
                        chart_paragraph.add_run().text = f"[图表: {os.path.basename(chart_path)}]"
                        continue
                    
                    # 添加图表标题
                    chart_title = chart_result.get('metadata', {}).get('title')
                    if chart_title:
                        title_paragraph = doc.add_paragraph()
                        title_paragraph.alignment = self.WD_PARAGRAPH_ALIGNMENT.CENTER
                        title_run = title_paragraph.add_run(str(chart_title))
                        title_run.font.bold = True
                    
                    # 添加空行作为间距
                    doc.add_paragraph()
                    inserted_count += 1
                    
                except Exception as chart_error:
                    logger.error(f"处理单个图表失败: {chart_error}", exc_info=True)
                    continue
            
            logger.info(f"已插入 {inserted_count}/{len(chart_results)} 个图表")
            
        except Exception as e:
            logger.error(f"插入图表失败: {e}", exc_info=True)
    
    def _apply_document_formatting(self, doc, config: DocumentConfig):
        """应用文档格式"""
        try:
            # 设置默认样式
            try:
                style = doc.styles['Normal']
                font = style.font
                font.name = config.font_name
                font.size = self.Pt(config.font_size)
                
                # 设置段落格式
                paragraph_format = style.paragraph_format
                paragraph_format.line_spacing = config.line_spacing
            except Exception as style_error:
                logger.warning(f"设置默认样式失败: {style_error}，继续应用格式")
            
            # 应用到所有段落（逐个处理，避免整体替换导致格式丢失）
            applied_count = 0
            for paragraph in doc.paragraphs:
                try:
                    # 只更新字体，不替换样式（保留原有格式如粗体、斜体等）
                    for run in paragraph.runs:
                        try:
                            run.font.name = config.font_name
                            run.font.size = self.Pt(config.font_size)
                        except Exception as run_error:
                            logger.debug(f"更新 run 字体失败: {run_error}")
                            continue
                    applied_count += 1
                except Exception as para_error:
                    logger.debug(f"处理段落格式失败: {para_error}")
                    continue
            
            # 设置页面边距
            try:
                sections = doc.sections
                for section in sections:
                    try:
                        section.top_margin = self.Inches(config.margins['top'] / 2.54)  # 转换cm到英寸
                        section.bottom_margin = self.Inches(config.margins['bottom'] / 2.54)
                        section.left_margin = self.Inches(config.margins['left'] / 2.54)
                        section.right_margin = self.Inches(config.margins['right'] / 2.54)
                    except Exception as section_error:
                        logger.debug(f"设置章节边距失败: {section_error}")
                        continue
            except Exception as sections_error:
                logger.warning(f"设置页面边距失败: {sections_error}")
            
            logger.debug(f"文档格式应用完成，处理了 {applied_count} 个段落")
            
        except Exception as e:
            logger.error(f"应用文档格式失败: {e}", exc_info=True)
    
    async def _save_document(
        self,
        doc,
        template_id: str,
        output_format: DocumentFormat
    ) -> str:
        """保存文档"""
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"report_{template_id}_{timestamp}.{output_format.value}"
            file_path = os.path.join(self.output_dir, filename)
            
            if output_format == DocumentFormat.DOCX:
                doc.save(file_path)
            elif output_format == DocumentFormat.PDF:
                # 需要额外的PDF转换库
                docx_path = file_path.replace('.pdf', '.docx')
                doc.save(docx_path)
                file_path = await self._convert_to_pdf(docx_path, file_path)
            elif output_format == DocumentFormat.HTML:
                # 简单的HTML导出
                file_path = await self._convert_to_html(doc, file_path)
            
            logger.debug(f"文档已保存: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"保存文档失败: {e}")
            raise
    
    async def _convert_to_pdf(self, docx_path: str, pdf_path: str) -> str:
        """转换DOCX到PDF"""
        try:
            # 这里需要使用如python-docx2pdf或类似库
            # 暂时返回DOCX路径作为后备
            logger.warning("PDF转换功能未实现，返回DOCX格式")
            return docx_path
        except Exception as e:
            logger.error(f"PDF转换失败: {e}")
            return docx_path
    
    async def _convert_to_html(self, doc, html_path: str) -> str:
        """转换到HTML"""
        try:
            # 简单的HTML生成
            html_content = ["<!DOCTYPE html>", "<html><head><title>报告</title></head><body>"]
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    html_content.append(f"<p>{paragraph.text}</p>")
            
            html_content.extend(["</body></html>"])
            
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(html_content))
            
            return html_path
            
        except Exception as e:
            logger.error(f"HTML转换失败: {e}")
            raise
    
    def get_document_preview(self, document_path: str) -> Optional[Dict[str, Any]]:
        """获取文档预览信息"""
        try:
            if not os.path.exists(document_path):
                return None
            
            stat = os.stat(document_path)
            
            preview_info = {
                "file_path": document_path,
                "filename": os.path.basename(document_path),
                "size_bytes": stat.st_size,
                "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "format": os.path.splitext(document_path)[1][1:].upper()
            }
            
            # 如果是DOCX文件，获取更多信息
            if document_path.endswith('.docx') and self.Document:
                try:
                    doc = self.Document(document_path)
                    # 简化图片检测，避免xpath的namespaces兼容性问题
                    has_images = False
                    try:
                        for rel in doc.part.rels.values():
                            if "image" in rel.target_ref:
                                has_images = True
                                break
                    except:
                        pass

                    preview_info.update({
                        "paragraph_count": len(doc.paragraphs),
                        "table_count": len(doc.tables),
                        "has_images": has_images
                    })
                except Exception:
                    pass
            
            return preview_info
            
        except Exception as e:
            logger.error(f"获取文档预览失败: {e}")
            return None
    
    def list_exported_documents(self) -> List[Dict[str, Any]]:
        """列出已导出的文档"""
        try:
            documents = []
            for filename in os.listdir(self.output_dir):
                if filename.startswith('report_'):
                    file_path = os.path.join(self.output_dir, filename)
                    preview_info = self.get_document_preview(file_path)
                    if preview_info:
                        documents.append(preview_info)
            
            return sorted(documents, key=lambda x: x['created_at'], reverse=True)
            
        except Exception as e:
            logger.error(f"列出文档失败: {e}")
            return []
    
    def cleanup_old_documents(self, days_old: int = 30) -> int:
        """清理旧文档"""
        try:
            cutoff_time = datetime.now().timestamp() - (days_old * 24 * 3600)
            
            removed_count = 0
            for filename in os.listdir(self.output_dir):
                file_path = os.path.join(self.output_dir, filename)
                if os.path.isfile(file_path) and os.path.getctime(file_path) < cutoff_time:
                    os.remove(file_path)
                    removed_count += 1
            
            logger.info(f"清理了 {removed_count} 个旧文档文件")
            return removed_count
            
        except Exception as e:
            logger.error(f"清理文档文件失败: {e}")
            return 0
    
    async def batch_export_templates(
        self,
        template_ids: List[str],
        base_data: Dict[str, Any],
        config: Optional[DocumentConfig] = None
    ) -> List[DocumentExportResult]:
        """批量导出模板"""
        logger.info(f"开始批量导出 {len(template_ids)} 个模板")
        
        export_tasks = []
        for template_id in template_ids:
            task = self.export_report_document(
                template_id=template_id,
                placeholder_data=base_data.get('placeholder_data', {}),
                etl_data=base_data.get('etl_data', {}),
                chart_data=base_data.get('chart_data', {}),
                config=config
            )
            export_tasks.append(task)
        
        results = await asyncio.gather(*export_tasks, return_exceptions=True)
        
        export_results = []
        for template_id, result in zip(template_ids, results):
            if isinstance(result, Exception):
                export_results.append(DocumentExportResult(
                    success=False,
                    error=str(result)
                ))
            else:
                export_results.append(result)
        
        successful_exports = len([r for r in export_results if r.success])
        logger.info(f"批量导出完成: 成功={successful_exports}, 总计={len(export_results)}")
        
        return export_results


# 工厂函数
def create_word_export_service(user_id: str) -> WordExportService:
    """创建Word导出服务实例"""
    return WordExportService(user_id=user_id)
