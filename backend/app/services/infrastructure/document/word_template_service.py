"""
Word文档模板处理服务

基于用户提供的文档替换逻辑，实现智能的占位符替换和图表生成
"""

import logging
import re
import io
import os
from typing import Dict, Any, List, Optional
from datetime import datetime

# Word文档处理
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 图表生成
try:
    import matplotlib.pyplot as plt
    import matplotlib.font_manager as fm
    plt.rcParams['font.sans-serif'] = ['SimHei', 'Arial Unicode MS']  # 支持中文
    plt.rcParams['axes.unicode_minus'] = False
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False

logger = logging.getLogger(__name__)


class WordTemplateService:
    """Word文档模板处理服务"""

    def __init__(self, font_path: Optional[str] = None):
        self.logger = logging.getLogger(self.__class__.__name__)
        self.font_path = font_path
        self.font_prop = None

        if not DOCX_AVAILABLE:
            self.logger.warning("python-docx 未安装，Word文档功能将受限")
        if not MATPLOTLIB_AVAILABLE:
            self.logger.warning("matplotlib 未安装，图表生成功能将受限")

        # 初始化字体
        if font_path and MATPLOTLIB_AVAILABLE:
            try:
                self.font_prop = fm.FontProperties(fname=font_path)
            except Exception as e:
                self.logger.warning(f"字体文件加载失败: {e}")

    async def process_document_template(
        self,
        template_path: str,
        placeholder_data: Dict[str, Any],
        output_path: str,
        container=None,
        use_agent_charts: bool = True,
        use_agent_optimization: bool = True,
        user_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        处理Word文档模板，替换占位符和生成图表

        Args:
            template_path: 模板文件路径
            placeholder_data: 占位符数据
            output_path: 输出文件路径
            container: 服务容器，用于Agent图表生成和内容优化
            use_agent_charts: 是否使用Agent生成图表
            use_agent_optimization: 是否使用Agent优化文档内容
            user_id: 用户UUID，用于Agent调用

        Returns:
            处理结果
        """
        try:
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx 未安装，无法处理Word文档")

            self.logger.info(f"开始处理Word模板: {template_path} (Agent图表: {use_agent_charts}, Agent优化: {use_agent_optimization})")
            self.logger.info(f"📊 接收到 {len(placeholder_data)} 个占位符数据")

            # 记录前5个占位符的详细信息
            for i, (key, value) in enumerate(list(placeholder_data.items())[:5]):
                value_preview = str(value)[:100] if value is not None else "None"
                self.logger.info(f"  占位符 {i+1}: {key} = {value_preview}")

            # 加载文档
            doc = Document(template_path)

            # 替换文本占位符
            self._replace_text_in_document(doc, placeholder_data)

            # Agent优化文档内容（在替换占位符后，生成图表前）
            if use_agent_optimization and container:
                await self._optimize_document_content_with_agent(doc, placeholder_data, container, user_id)

            # 替换图表占位符 - 优先使用Agent
            if use_agent_charts and container:
                await self._replace_chart_placeholders_with_agent(doc, placeholder_data, container, user_id)
            else:
                await self._replace_chart_placeholders_fallback(doc, placeholder_data)

            # 保存文档
            doc.save(output_path)

            self.logger.info(f"✅ Word文档处理完成: {output_path}")

            return {
                "success": True,
                "output_path": output_path,
                "placeholders_processed": len(placeholder_data),
                "chart_generation_method": "agent" if use_agent_charts and container else "traditional",
                "content_optimization": "enabled" if use_agent_optimization and container else "disabled",
                "message": "Word文档处理成功"
            }

        except Exception as e:
            self.logger.error(f"❌ Word文档处理失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "message": "Word文档处理失败"
            }

    def process_document_template_sync(
        self,
        template_path: str,
        placeholder_data: Dict[str, Any],
        output_path: str
    ) -> Dict[str, Any]:
        """
        同步版本的文档处理方法，用于向后兼容
        使用传统图表生成方法
        """
        import asyncio

        try:
            loop = asyncio.get_event_loop()
            return loop.run_until_complete(
                self.process_document_template(
                    template_path, placeholder_data, output_path,
                    container=None, use_agent_charts=False
                )
            )
        except RuntimeError:
            # 如果没有事件循环，创建新的
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                return loop.run_until_complete(
                    self.process_document_template(
                        template_path, placeholder_data, output_path,
                        container=None, use_agent_charts=False
                    )
                )
            finally:
                loop.close()

    def _replace_text_in_document(self, doc, data: Dict[str, Any]):
        """
        替换文档中的文本占位符
        参考用户提供的replace_text_in_document逻辑
        """
        replaced_count = 0
        self.logger.info(f"🔄 开始替换文本占位符，数据字典包含 {len(data)} 个键")

        # 处理段落
        for p in doc.paragraphs:
            if '{{' in p.text and '}}' in p.text:
                placeholders_in_paragraph = re.findall(r"\{\{.*?\}\}", p.text)
                self.logger.debug(f"段落中发现 {len(placeholders_in_paragraph)} 个占位符: {p.text[:100]}")

                for placeholder in placeholders_in_paragraph:
                    # 跳过图表占位符
                    if placeholder.startswith("{{图表："):
                        continue

                    # 尝试多种格式查找数据
                    value = None
                    matched_key = None
                    if placeholder in data:
                        value = data[placeholder]
                        matched_key = placeholder
                    else:
                        # 尝试去掉花括号的格式
                        placeholder_without_braces = placeholder.replace("{{", "").replace("}}", "")
                        if placeholder_without_braces in data:
                            value = data[placeholder_without_braces]
                            matched_key = placeholder_without_braces

                    if value is None:
                        self.logger.warning(f"⚠️ 未找到占位符数据: {placeholder}")
                        continue

                    self.logger.info(f"✅ 找到匹配: {placeholder} -> {matched_key} = {str(value)[:50]}")
                    replaced_count += 1

                    str_value = str(value) if value is not None else ""

                    # 获取段落的完整文本
                    full_text = "".join(run.text for run in p.runs)

                    if placeholder in full_text:
                        start_index = full_text.find(placeholder)
                        end_index = start_index + len(placeholder)

                        # 找到占位符在哪些runs中
                        current_pos = 0
                        start_run = None
                        end_run = None
                        start_run_char_index = 0
                        end_run_char_index = 0

                        for i, run in enumerate(p.runs):
                            run_len = len(run.text)

                            if start_run is None and start_index < current_pos + run_len:
                                start_run = i
                                start_run_char_index = start_index - current_pos

                            if end_run is None and end_index <= current_pos + run_len:
                                end_run = i
                                end_run_char_index = end_index - current_pos
                                break

                            current_pos += run_len

                        # 执行替换
                        if start_run is not None and end_run is not None:
                            start_run_obj = p.runs[start_run]
                            start_run_obj.text = start_run_obj.text[:start_run_char_index] + str_value

                            end_run_obj = p.runs[end_run]
                            end_run_obj.text = end_run_obj.text[end_run_char_index:]

                            # 清空中间的runs
                            for i in range(start_run + 1, end_run):
                                p.runs[i].text = ""

        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self._replace_text_in_document(cell, data)

        self.logger.info(f"📝 文本占位符替换完成，共替换 {replaced_count} 个占位符")

    async def _optimize_document_content_with_agent(self, doc, data: Dict[str, Any], container=None, user_id: Optional[str] = None):
        """
        使用Agent优化文档内容 - 根据实际数据智能调整占位符周围的文字描述

        Args:
            doc: Word文档对象
            data: 占位符数据
            container: 服务容器
        """
        if not container:
            self.logger.warning("没有提供服务容器，跳过文档内容优化")
            return

        from app.services.infrastructure.agents import AgentService
        from app.services.infrastructure.agents.types import AgentInput, PlaceholderSpec, SchemaInfo, TaskContext, AgentConstraints

        try:
            agent_service = AgentService(container=container)

            total_paragraphs = len(doc.paragraphs)
            self.logger.info(f"📄 开始文档优化，共 {total_paragraphs} 个段落")

            # 遍历所有段落，找到需要优化的内容
            optimized_count = 0
            for i, p in enumerate(doc.paragraphs):
                paragraph_text = p.text.strip()

                # 跳过空段落和图表占位符
                if not paragraph_text or paragraph_text.startswith("{{图表："):
                    continue

                # 检查段落中是否包含已替换的数据值
                has_data_value = False
                related_placeholders = []

                for placeholder_key, placeholder_value in data.items():
                    # 跳过图表占位符
                    if "图表" in placeholder_key or placeholder_key.startswith("{{图表："):
                        continue

                    # 检查段落是否包含这个占位符的值
                    str_value = str(placeholder_value) if placeholder_value is not None else ""
                    # 对于较短的值（如单个数字），需要更严格的匹配
                    if str_value and len(str_value) >= 2 and str_value in paragraph_text:
                        has_data_value = True
                        related_placeholders.append({
                            "key": placeholder_key,
                            "value": placeholder_value
                        })

                # 如果段落包含数据值，使用Agent优化
                if has_data_value and related_placeholders:
                    self.logger.info(f"🤖 使用Agent优化段落 {i+1}: {paragraph_text[:50]}...")

                    try:
                        # 构建优化提示
                        context_info = "\n".join([
                            f"- {ph['key']}: {ph['value']}"
                            for ph in related_placeholders[:5]  # 最多5个占位符
                        ])

                        optimization_prompt = f"""请优化以下报告段落，使其更符合数据特征和专业性要求。

原始段落:
{paragraph_text}

相关数据:
{context_info}

要求:
1. 保持段落的核心意思和数据准确性
2. 使用更专业、流畅的表达方式
3. 根据数据值调整描述的语气（如数值高低、趋势等）
4. 保持简洁，不要过度冗长
5. 只返回优化后的段落文本，不要添加任何解释

请直接输出优化后的段落文本："""

                        # 准备Agent输入
                        placeholder_spec = PlaceholderSpec(
                            id=f"paragraph_{i}",
                            description=f"段落优化: {paragraph_text[:30]}",
                            type="text"
                        )

                        agent_input = AgentInput(
                            user_prompt=optimization_prompt,
                            placeholder=placeholder_spec,
                            schema=SchemaInfo(tables=[], columns={}),
                            context=TaskContext(task_time=None, timezone="Asia/Shanghai"),
                            constraints=AgentConstraints(output_kind="text", max_attempts=1),
                            data_source={"id": "", "type": "generated"},
                            task_driven_context={
                                "paragraph_text": paragraph_text,
                                "related_data": related_placeholders
                            },
                            user_id=user_id or "system"
                        )

                        # 调用Agent
                        agent_result = await agent_service.execute(agent_input)

                        if agent_result.success and hasattr(agent_result, 'result') and agent_result.result:
                            optimized_text = str(agent_result.result).strip()

                            # 清理可能的JSON或Markdown包裹
                            import json
                            try:
                                # 如果返回的是JSON，提取文本
                                parsed = json.loads(optimized_text)
                                if isinstance(parsed, dict):
                                    # 优先查找 optimized_paragraph 字段，然后是其他常见字段
                                    optimized_text = parsed.get('optimized_paragraph') or \
                                                    parsed.get('result') or \
                                                    parsed.get('text') or \
                                                    parsed.get('content') or \
                                                    optimized_text

                                    # 如果仍然是整个JSON（没找到有效字段），检查是否是错误响应
                                    if optimized_text == str(parsed) and parsed.get('success') == False:
                                        # 这是一个错误响应，不应该插入文档
                                        self.logger.warning(f"⚠️ Agent返回错误响应，跳过优化: {parsed.get('error', 'unknown')}")
                                        optimized_text = None
                            except json.JSONDecodeError:
                                # 不是JSON格式，保持原样
                                pass
                            except Exception as e:
                                self.logger.warning(f"JSON解析异常: {e}")
                                pass

                            # 移除可能的markdown代码块标记
                            if optimized_text:
                                optimized_text = optimized_text.replace('```', '').strip()

                            if optimized_text:
                                if optimized_text != paragraph_text:
                                    self.logger.info(f"✅ 段落优化成功: {optimized_text[:50]}...")
                                    optimized_count += 1

                                    # 保持原有的格式，只替换文本
                                    if p.runs:
                                        # 保留第一个run的格式
                                        first_run = p.runs[0]
                                        # 清空所有runs
                                        for run in p.runs:
                                            run.text = ""
                                        # 在第一个run中设置新文本
                                        first_run.text = optimized_text
                                    else:
                                        p.text = optimized_text
                                else:
                                    self.logger.debug("优化结果与原文相同，保持不变")
                            else:
                                self.logger.debug("Agent返回无效内容，保持原文不变")
                        else:
                            error_msg = getattr(agent_result, 'metadata', {}).get('error', '优化失败')
                            self.logger.warning(f"⚠️ 段落优化失败: {error_msg}")

                    except Exception as opt_error:
                        self.logger.warning(f"⚠️ 段落优化异常: {opt_error}, 保持原文")
                        continue

            self.logger.info(f"✅ 文档内容优化完成，共优化 {optimized_count} 个段落")

        except Exception as e:
            self.logger.error(f"❌ 文档内容优化失败: {e}")
            # 优化失败不影响整体流程，继续执行

    async def _replace_chart_placeholders_with_agent(self, doc, data: Dict[str, Any], container=None, user_id: Optional[str] = None):
        """
        使用Agent替换图表占位符 - 更智能的图表生成

        Args:
            doc: Word文档对象
            data: 占位符数据
            container: 服务容器
            user_id: 用户UUID
        """
        if not DOCX_AVAILABLE:
            return

        from app.services.infrastructure.agents import AgentService
        from app.services.infrastructure.agents.types import AgentInput

        # 如果没有容器，回退到传统方法
        if not container:
            self.logger.warning("没有提供服务容器，回退到传统图表生成")
            return await self._replace_chart_placeholders_fallback(doc, data)

        agent_service = AgentService(container=container)

        for p in doc.paragraphs:
            placeholder = p.text.strip()

            if placeholder.startswith("{{图表："):
                # 尝试多种key格式查找数据
                chart_data = data.get(placeholder)  # 先尝试完整格式 {{图表：xxx}}

                if chart_data is None:
                    # 尝试去掉花括号的格式 图表：xxx
                    placeholder_without_braces = placeholder.replace("{{", "").replace("}}", "")
                    chart_data = data.get(placeholder_without_braces)

                if chart_data is None:
                    self.logger.warning(f"没有找到图表数据: {placeholder} (也尝试了 {placeholder_without_braces})")
                    continue

                self.logger.info(f"🤖 使用Agent为 '{placeholder}' 生成图表...")

                title = placeholder.replace("{{图表：", "").replace("}}", "")
                p.text = ""

                try:
                    # 准备Agent输入 - 使用正确的数据结构
                    from app.services.infrastructure.agents.types import PlaceholderSpec, SchemaInfo, TaskContext, AgentConstraints

                    placeholder_spec = PlaceholderSpec(
                        id=placeholder,
                        description=title,
                        type="chart"
                    )

                    schema_info = SchemaInfo(
                        tables=[],
                        columns={}
                    )

                    task_context = TaskContext(
                        task_time=None,
                        timezone="Asia/Shanghai"
                    )

                    constraints = AgentConstraints(
                        output_kind="chart",
                        max_attempts=3
                    )

                    data_rows = self._convert_data_to_rows(chart_data)
                    data_columns = self._extract_columns_from_data(chart_data)

                    agent_input = AgentInput(
                        user_prompt=f"为以下数据生成图表：{title}",
                        placeholder=placeholder_spec,
                        schema=schema_info,
                        context=task_context,
                        constraints=constraints,
                        data_source={"id": "", "type": "generated"},
                        task_driven_context={
                            "execution_result": {
                                "rows": data_rows,
                                "columns": data_columns
                            },
                            "chart_requirements": {
                                "title": title,
                                "placeholder": placeholder,
                                "suggested_type": self._suggest_chart_type(placeholder, chart_data)
                            }
                        },
                        user_id=user_id or "system"
                    )

                    # 使用Agent生成图表
                    agent_result = await agent_service.execute(agent_input)

                    if agent_result.success and hasattr(agent_result, 'result') and agent_result.result:
                        # Agent成功生成了图表
                        chart_image_path = agent_result.result
                        self.logger.info(f"✅ Agent图表生成成功: {chart_image_path}")

                        # 插入图表到文档
                        run = p.add_run()
                        run.add_picture(chart_image_path, width=Inches(6.0))
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        # Agent失败，记录错误并添加占位符
                        error_msg = getattr(agent_result, 'metadata', {}).get('error', '图表生成失败')
                        self.logger.error(f"❌ Agent图表生成失败: {error_msg}")

                        p.add_run().text = f"[{title} - Agent图表生成失败: {error_msg}]"
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                except Exception as e:
                    self.logger.error(f"❌ Agent图表生成异常: {e}")
                    # 异常情况下回退到传统方法
                    chart_buffer = self._create_chart_fallback(chart_data, title)
                    if chart_buffer and chart_buffer.getbuffer().nbytes > 0:
                        run = p.add_run()
                        run.add_picture(chart_buffer, width=Inches(6.0))
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.add_run().text = f"[{title} - 图表生成异常]"
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    async def _replace_chart_placeholders_fallback(self, doc, data: Dict[str, Any]):
        """
        传统图表替换方法作为回退
        """
        if not MATPLOTLIB_AVAILABLE:
            self.logger.warning("matplotlib 未安装，跳过图表生成")
            return

        for p in doc.paragraphs:
            placeholder = p.text.strip()

            if placeholder.startswith("{{图表："):
                # 尝试多种key格式查找数据
                chart_data = data.get(placeholder)  # 先尝试完整格式 {{图表：xxx}}

                if chart_data is None:
                    # 尝试去掉花括号的格式 图表：xxx
                    placeholder_without_braces = placeholder.replace("{{", "").replace("}}", "")
                    chart_data = data.get(placeholder_without_braces)

                if chart_data is None:
                    continue

                title = placeholder.replace("{{图表：", "").replace("}}", "")
                p.text = ""

                # 生成图表
                chart_buffer = self._create_chart_fallback(chart_data, title)

                if chart_buffer and chart_buffer.getbuffer().nbytes > 0:
                    run = p.add_run()
                    run.add_picture(chart_buffer, width=Inches(6.0))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.add_run().text = f"[{title} - 图表生成失败]"
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _convert_data_to_rows(self, chart_data) -> List[List]:
        """
        将图表数据转换为Agent期望的行格式
        """
        if not chart_data:
            return []

        if isinstance(chart_data, list) and chart_data:
            if isinstance(chart_data[0], dict):
                # 字典列表格式，转换为行列表
                keys = list(chart_data[0].keys())
                rows = []
                for item in chart_data:
                    row = [item.get(key, "") for key in keys]
                    rows.append(row)
                return rows
            elif isinstance(chart_data[0], (list, tuple)):
                # 已经是行格式
                return [list(row) for row in chart_data]

        # 其他格式，尝试转换
        return [[str(chart_data)]]

    def _extract_columns_from_data(self, chart_data) -> List[str]:
        """
        从图表数据中提取列名
        """
        if not chart_data:
            return []

        if isinstance(chart_data, list) and chart_data:
            if isinstance(chart_data[0], dict):
                return list(chart_data[0].keys())
            elif len(chart_data[0]) >= 2:
                return ["标签", "数值"]

        return ["值"]

    def _suggest_chart_type(self, placeholder: str, chart_data) -> str:
        """
        基于占位符和数据建议图表类型
        """
        placeholder_lower = placeholder.lower()

        if any(word in placeholder_lower for word in ["饼图", "pie", "占比", "比例"]):
            return "pie"
        elif any(word in placeholder_lower for word in ["线图", "line", "趋势", "变化"]):
            return "line"
        elif any(word in placeholder_lower for word in ["柱状图", "bar", "柱图", "对比"]):
            return "bar"
        elif any(word in placeholder_lower for word in ["散点图", "scatter", "分布"]):
            return "scatter"
        else:
            # 根据数据特征判断
            if isinstance(chart_data, list) and len(chart_data) <= 8:
                return "pie"  # 少量数据适合饼图
            else:
                return "bar"  # 默认柱状图

    def _create_chart_fallback(self, chart_data, title: str):
        """
        传统图表生成作为回退方案
        """
        return self._create_chart(chart_data, title, "bar")

    def _create_chart(
        self,
        data: List[Dict[str, Any]],
        title: str,
        chart_type: str
    ) -> Optional[io.BytesIO]:
        """
        创建图表
        参考用户提供的create_chart逻辑
        """
        try:
            if not data or not isinstance(data, list) or not all(isinstance(i, dict) for i in data):
                self.logger.warning(f"⚠️ 警告: '{title}' 的图表数据格式不正确，跳过生成。")
                return None

            first_item = data[0]
            label_key = None
            value_key = None

            # 自动识别标签和数值列
            for key, value in first_item.items():
                if isinstance(value, str):
                    label_key = key
                elif isinstance(value, (int, float)):
                    value_key = key

            if label_key is None or value_key is None:
                self.logger.warning(f"⚠️ 警告: 无法从 '{title}' 的数据中识别标签和数值列，跳过生成。")
                return None

            # 提取数据
            labels = [item.get(label_key, '') for item in data]
            values = [float(item.get(value_key, 0)) for item in data]

            # 创建图表
            fig, ax = plt.subplots(figsize=(10, 6) if chart_type == 'bar' else (8, 8))
            ax.set_title(title, fontsize=16, fontproperties=self.font_prop)

            if chart_type == 'bar':
                ax.bar(labels, values)
                ax.set_ylabel("数量", fontsize=12, fontproperties=self.font_prop)
                plt.xticks(rotation=45, ha="right")

                # 设置x轴标签字体
                for tick_label in ax.get_xticklabels():
                    if self.font_prop:
                        tick_label.set_fontproperties(self.font_prop)

                fig.tight_layout()

            elif chart_type == 'pie':
                text_props = {'fontproperties': self.font_prop} if self.font_prop else {}
                ax.pie(values, labels=labels, autopct='%1.1f%%', startangle=90, textprops=text_props)
                ax.axis('equal')

            # 保存到内存缓冲区
            img_buffer = io.BytesIO()
            fig.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
            img_buffer.seek(0)
            plt.close(fig)

            return img_buffer

        except Exception as e:
            self.logger.error(f"❌ 图表生成失败: {e}")
            return None

    def extract_placeholders_from_template(self, template_path: str) -> List[str]:
        """
        从模板中提取所有占位符

        Args:
            template_path: 模板文件路径

        Returns:
            占位符列表
        """
        try:
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx 未安装")

            doc = Document(template_path)
            placeholders = set()

            # 从段落中提取
            for p in doc.paragraphs:
                found = re.findall(r"\{\{.*?\}\}", p.text)
                placeholders.update(found)

            # 从表格中提取
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            found = re.findall(r"\{\{.*?\}\}", p.text)
                            placeholders.update(found)

            self.logger.info(f"从模板中提取到 {len(placeholders)} 个占位符")
            return list(placeholders)

        except Exception as e:
            self.logger.error(f"❌ 提取占位符失败: {e}")
            return []

    async def process_template_with_data(
        self,
        template_path: str,
        output_path: str,
        placeholder_data: Dict[str, str]
    ) -> Dict[str, Any]:
        """
        处理模板并替换占位符数据

        这是DocAssemblerTool需要的主要接口方法

        Args:
            template_path: 模板文件路径
            output_path: 输出文件路径
            placeholder_data: 占位符数据字典 {placeholder_name: processed_text}

        Returns:
            处理结果
        """
        try:
            if not DOCX_AVAILABLE:
                return {
                    "success": False,
                    "error": "python-docx 未安装，无法处理Word文档",
                    "output_path": None
                }

            self.logger.info(f"开始处理Word模板: {template_path} -> {output_path}")
            self.logger.info(f"占位符数量: {len(placeholder_data)}")

            # 加载文档
            doc = Document(template_path)

            # 替换文本占位符
            self._replace_text_in_document(doc, placeholder_data)

            # 替换图表占位符（使用传统方法，因为这时已经是处理后的文本数据）
            await self._replace_chart_placeholders_fallback(doc, placeholder_data)

            # 保存文档
            doc.save(output_path)

            self.logger.info(f"✅ Word文档处理完成: {output_path}")

            return {
                "success": True,
                "output_path": output_path,
                "placeholders_processed": len(placeholder_data),
                "message": "Word文档处理成功"
            }

        except Exception as e:
            self.logger.error(f"❌ Word文档处理失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "output_path": None
            }

    def validate_template_format(self, template_path: str) -> Dict[str, Any]:
        """
        验证模板格式

        Args:
            template_path: 模板文件路径

        Returns:
            验证结果
        """
        try:
            result = {
                "valid": True,
                "issues": [],
                "warnings": [],
                "placeholders": [],
                "chart_placeholders": []
            }

            placeholders = self.extract_placeholders_from_template(template_path)
            result["placeholders"] = placeholders

            # 分类占位符
            for placeholder in placeholders:
                if placeholder.startswith("{{图表："):
                    result["chart_placeholders"].append(placeholder)

            # 检查是否有悬空的花括号
            if not DOCX_AVAILABLE:
                result["warnings"].append("python-docx 未安装，无法进行详细验证")
                return result

            doc = Document(template_path)
            doc_text = "\n".join([p.text for p in doc.paragraphs])

            unmatched_braces = re.findall(r'(?:^|[^{]){(?:[^{}]|$)', doc_text)
            if unmatched_braces:
                result["warnings"].append("检测到可能不匹配的花括号")

            self.logger.info(f"✅ 模板验证完成: {'通过' if result['valid'] else '失败'}")
            return result

        except Exception as e:
            self.logger.error(f"❌ 模板验证失败: {e}")
            return {
                "valid": False,
                "issues": [f"验证过程异常: {e}"],
                "warnings": [],
                "placeholders": [],
                "chart_placeholders": []
            }


# 全局服务实例
word_template_service = WordTemplateService()


def create_word_template_service(font_path: Optional[str] = None) -> WordTemplateService:
    """创建Word模板服务实例"""
    return WordTemplateService(font_path=font_path)


def create_agent_enhanced_word_service(
    font_path: Optional[str] = None,
    container=None
) -> 'AgentEnhancedWordService':
    """创建Agent增强的Word服务实例"""
    return AgentEnhancedWordService(font_path=font_path, container=container)


class AgentEnhancedWordService(WordTemplateService):
    """Agent增强的Word服务，默认使用Agent生成图表"""

    def __init__(self, font_path: Optional[str] = None, container=None):
        super().__init__(font_path)
        self.container = container

    async def process_document_template_enhanced(
        self,
        template_path: str,
        placeholder_data: Dict[str, Any],
        output_path: str,
        use_intelligent_text: bool = True,
        user_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        增强版文档处理，默认使用Agent图表生成和智能文本处理
        """
        return await self.process_document_template_with_intelligence(
            template_path=template_path,
            placeholder_data=placeholder_data,
            output_path=output_path,
            container=self.container,
            use_agent_charts=True,
            use_intelligent_text=use_intelligent_text,
            user_id=user_id
        )

    async def process_document_template_with_intelligence(
        self,
        template_path: str,
        placeholder_data: Dict[str, Any],
        output_path: str,
        container=None,
        use_agent_charts: bool = False,
        use_intelligent_text: bool = True,
        user_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        带智能文本处理的文档模板处理

        Args:
            template_path: 模板文件路径
            placeholder_data: ETL返回的占位符数据
            output_path: 输出文件路径
            container: 服务容器
            use_agent_charts: 是否使用Agent生成图表
            use_intelligent_text: 是否使用智能文本处理
            user_id: 用户UUID

        Returns:
            处理结果
        """
        if not DOCX_AVAILABLE:
            return {
                "success": False,
                "error": "python-docx未安装，无法处理Word文档",
                "placeholders_processed": 0
            }

        try:
            self.logger.info(f"📄 开始智能文档处理: {template_path} (智能文本: {use_intelligent_text})")

            # 1. 检查模板文件
            if not os.path.exists(template_path):
                return {
                    "success": False,
                    "error": f"模板文件不存在: {template_path}",
                    "placeholders_processed": 0
                }

            # 2. 打开文档
            doc = Document(template_path)
            self.logger.info(f"📄 Word文档加载成功，段落数: {len(doc.paragraphs)}")

            # 3. 智能文本处理 (核心新功能)
            processed_placeholder_data = placeholder_data
            if use_intelligent_text and container:
                processed_placeholder_data = await self._process_placeholder_data_intelligently(
                    doc, placeholder_data, container
                )

            # 4. 替换文本占位符
            self._replace_text_in_document(doc, processed_placeholder_data)

            # 5. 处理图表占位符
            if use_agent_charts and container:
                await self._replace_chart_placeholders_with_agent(doc, placeholder_data, container, user_id)
            else:
                await self._replace_chart_placeholders_fallback(doc, placeholder_data)

            # 6. 保存文档
            doc.save(output_path)
            self.logger.info(f"✅ Word文档保存成功: {output_path}")

            return {
                "success": True,
                "output_path": output_path,
                "placeholders_processed": len(placeholder_data),
                "intelligent_text_used": use_intelligent_text and container,
                "chart_generation_method": "agent" if use_agent_charts and container else "traditional",
                "message": f"智能文档处理完成: {output_path}"
            }

        except Exception as e:
            self.logger.error(f"❌ 智能文档处理失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "placeholders_processed": 0
            }

    async def _process_placeholder_data_intelligently(
        self,
        doc,
        placeholder_data: Dict[str, Any],
        container
    ) -> Dict[str, Any]:
        """
        使用Agent智能处理占位符数据

        这是核心的智能文本处理环节：
        1. 提取Word文档的上下文信息
        2. 对每个占位符进行智能文本生成
        3. 返回优化后的文本映射

        Args:
            doc: Word文档对象
            placeholder_data: ETL返回的原始数据 {placeholder_name: data_value}
            container: 服务容器

        Returns:
            智能处理后的文本映射 {placeholder_name: "intelligent_text"}
        """
        try:
            from app.services.infrastructure.agents.placeholder_intelligent_processor import create_placeholder_intelligent_processor

            self.logger.info(f"🤖 开始智能文本处理 {len(placeholder_data)} 个占位符")

            # 创建占位符智能处理器
            processor = create_placeholder_intelligent_processor(container)

            # 提取文档中的上下文信息
            document_text = self._extract_document_text(doc)
            template_context = processor.extract_template_context(document_text)

            self.logger.debug(f"提取到 {len(template_context)} 个占位符的上下文信息")

            # 智能处理占位符数据
            processed_data = await processor.process_placeholder_data(
                placeholder_data=placeholder_data,
                template_context=template_context
            )

            # 记录处理结果
            for name, original in placeholder_data.items():
                processed = processed_data.get(name, str(original))
                if str(original) != processed:
                    self.logger.info(f"📝 占位符智能优化: {name}")
                    self.logger.debug(f"   原始: {original}")
                    self.logger.debug(f"   优化: {processed}")

            self.logger.info(f"✅ 智能文本处理完成")
            return processed_data

        except Exception as e:
            self.logger.error(f"❌ 智能文本处理失败: {e}")
            # 降级到原始数据
            return placeholder_data

    def _extract_document_text(self, doc) -> str:
        """提取Word文档的文本内容，用于上下文分析"""
        if not DOCX_AVAILABLE:
            return ""

        try:
            text_parts = []

            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            document_text = "\n".join(text_parts)
            self.logger.debug(f"提取文档文本长度: {len(document_text)} 字符")
            return document_text

        except Exception as e:
            self.logger.warning(f"提取文档文本失败: {e}")
            return ""
