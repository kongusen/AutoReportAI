"""
Word文档生成流水线系统

基于模板的稳定文档生成流水线，支持：
1. 模板解析和验证
2. 占位符替换
3. 格式化和样式应用
4. Word文档生成和导出
5. 文档质量检查
"""

import asyncio
import json
import logging
import os
import re
import time
from dataclasses import dataclass, field
from enum import Enum
from typing import Any, Dict, List, Optional, Union
from datetime import datetime
import uuid
import base64

# 尝试导入python-docx，如果没有则使用模拟
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)


class DocumentFormat(Enum):
    """文档格式枚举"""
    DOCX = "docx"
    PDF = "pdf"
    HTML = "html"
    MARKDOWN = "md"


class ContentType(Enum):
    """内容类型枚举"""
    TEXT = "text"
    TABLE = "table"
    CHART = "chart"
    IMAGE = "image"
    LIST = "list"
    PARAGRAPH = "paragraph"


@dataclass
class PlaceholderInfo:
    """占位符信息"""
    name: str
    type: str
    description: str
    content_type: ContentType
    required: bool = True
    default_value: str = ""
    format_rules: Dict[str, Any] = field(default_factory=dict)


@dataclass
class DocumentTemplate:
    """文档模板"""
    template_id: str
    name: str
    content: str
    placeholders: List[PlaceholderInfo]
    styles: Dict[str, Any] = field(default_factory=dict)
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ProcessedContent:
    """处理后的内容"""
    placeholder_name: str
    content_type: ContentType
    raw_content: Any
    formatted_content: str
    style_info: Dict[str, Any] = field(default_factory=dict)
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class DocumentGenerationResult:
    """文档生成结果"""
    success: bool
    document_path: str
    document_format: DocumentFormat
    generation_time: float
    placeholder_count: int
    processed_placeholders: int
    failed_placeholders: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    error_message: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


class TemplateParser:
    """模板解析器"""
    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
        # 支持单花括号和双花括号格式的占位符
        self.placeholder_pattern = r'\{\{([^}]+)\}\}|\{([^}]+)\}'
    
    def parse_template(self, template_content: str) -> DocumentTemplate:
        """解析模板内容"""
        template_id = str(uuid.uuid4())
        
        # 提取占位符
        placeholders = self._extract_placeholders(template_content)
        
        # 分析模板结构
        styles = self._analyze_template_styles(template_content)
        
        return DocumentTemplate(
            template_id=template_id,
            name="Parsed Template",
            content=template_content,
            placeholders=placeholders,
            styles=styles,
            metadata={
                "parsed_at": datetime.now().isoformat(),
                "placeholder_count": len(placeholders)
            }
        )
    
    def extract_placeholders(self, content: str) -> List[Dict[str, Any]]:
        """提取占位符（公共方法）"""
        # 首先尝试检测是否为二进制Word文档
        text_content = self._extract_text_from_content(content)
        placeholders_info = self._extract_placeholders(text_content)
        
        # 转换为字典格式以兼容现有代码
        result = []
        for placeholder in placeholders_info:
            result.append({
                "name": placeholder.name,
                "type": placeholder.type,
                "description": placeholder.description,
                "content_type": placeholder.content_type.value if hasattr(placeholder.content_type, 'value') else str(placeholder.content_type),
                "required": placeholder.required
            })
        
        return result
    
    def _extract_text_from_content(self, content: str) -> str:
        """从内容中提取文本（支持二进制Word文档）"""
        # 检查是否为二进制数据（Word文档的十六进制表示）
        if self._is_binary_content(content):
            self.logger.info("检测到二进制Word文档，尝试解析...")
            try:
                # 尝试将十六进制字符串转换为二进制数据
                binary_data = bytes.fromhex(content)
                
                # 检查是否为Word文档（ZIP格式）
                if binary_data.startswith(b'PK'):
                    return self._extract_text_from_word_doc(binary_data)
                else:
                    self.logger.warning("无法识别的二进制格式")
                    return content
            except Exception as e:
                self.logger.error(f"解析二进制文档失败: {e}")
                return content
        else:
            # 纯文本内容，直接返回
            return content
    
    def _is_binary_content(self, content: str) -> bool:
        """检查内容是否为二进制数据的十六进制表示"""
        if len(content) > 20 and all(c in '0123456789abcdef' for c in content.lower()):
            # 检查是否以常见的文档格式魔术字节开头
            hex_prefix = content[:20].lower()
            # PK（ZIP/Word文档）= 504b
            # PDF = 255044462d
            if hex_prefix.startswith('504b') or hex_prefix.startswith('255044462d'):
                return True
        return False
    
    def _extract_text_from_word_doc(self, binary_data: bytes) -> str:
        """从Word文档二进制数据中提取文本"""
        try:
            import tempfile
            import os
            
            # 保存二进制数据到临时文件
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(binary_data)
                temp_file_path = temp_file.name
            
            try:
                if HAS_DOCX:
                    # 使用python-docx解析
                    from docx import Document
                    doc = Document(temp_file_path)
                    
                    # 提取所有段落文本
                    text_parts = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)
                    
                    # 提取表格文本
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    text_parts.append(cell.text)
                    
                    extracted_text = '\n'.join(text_parts)
                    self.logger.info(f"成功从Word文档提取文本，长度: {len(extracted_text)}")
                    return extracted_text
                else:
                    self.logger.warning("python-docx不可用，无法解析Word文档")
                    return ""
            finally:
                # 清理临时文件
                try:
                    os.unlink(temp_file_path)
                except:
                    pass
                    
        except Exception as e:
            self.logger.error(f"解析Word文档失败: {e}")
            return ""
    
    def _extract_placeholders(self, content: str) -> List[PlaceholderInfo]:
        """提取占位符（私有方法）"""
        placeholders = []
        matches = re.findall(self.placeholder_pattern, content)
        
        for match in matches:
            # match是一个元组，包含两个组：(双花括号匹配, 单花括号匹配)
            # 取非空的那个
            placeholder_name = (match[0] or match[1]).strip()
            
            if not placeholder_name:
                continue
            
            # 分析占位符类型
            content_type = self._infer_content_type(placeholder_name)
            
            placeholder = PlaceholderInfo(
                name=placeholder_name,
                type=self._infer_placeholder_type(placeholder_name),
                description=f"占位符: {placeholder_name}",
                content_type=content_type,
                required=True
            )
            
            placeholders.append(placeholder)
        
        # 去重
        unique_placeholders = []
        seen_names = set()
        
        for placeholder in placeholders:
            if placeholder.name not in seen_names:
                unique_placeholders.append(placeholder)
                seen_names.add(placeholder.name)
        
        return unique_placeholders
    
    def _infer_content_type(self, placeholder_name: str) -> ContentType:
        """推断内容类型"""
        name_lower = placeholder_name.lower()
        
        if any(keyword in name_lower for keyword in ['图表', 'chart', '图', '可视化']):
            return ContentType.CHART
        elif any(keyword in name_lower for keyword in ['表格', 'table', '列表', 'list']):
            return ContentType.TABLE
        elif any(keyword in name_lower for keyword in ['图片', 'image', '照片']):
            return ContentType.IMAGE
        elif any(keyword in name_lower for keyword in ['清单', '项目', 'item']):
            return ContentType.LIST
        else:
            return ContentType.TEXT
    
    def _infer_placeholder_type(self, placeholder_name: str) -> str:
        """推断占位符类型"""
        name_lower = placeholder_name.lower()
        
        if any(keyword in name_lower for keyword in ['统计', '数量', '总数', '计算']):
            return "statistic"
        elif any(keyword in name_lower for keyword in ['分析', '洞察', '趋势']):
            return "analysis"
        elif any(keyword in name_lower for keyword in ['图表', 'chart']):
            return "chart"
        else:
            return "text"
    
    def _analyze_template_styles(self, content: str) -> Dict[str, Any]:
        """分析模板样式"""
        styles = {
            "default_font": "Arial",
            "default_size": 12,
            "heading_styles": {},
            "paragraph_styles": {},
            "table_styles": {}
        }
        
        # 分析标题样式
        heading_pattern = r'^(#{1,6})\s+(.+)$'
        headings = re.findall(heading_pattern, content, re.MULTILINE)
        
        for heading_level, heading_text in headings:
            level = len(heading_level)
            styles["heading_styles"][f"heading_{level}"] = {
                "font_size": max(16 - level, 10),
                "bold": True,
                "color": "black"
            }
        
        return styles
    
    def parse_doc_placeholders(self, doc_path: str) -> Dict[str, Any]:
        """解析DOC文档中的占位符（API兼容方法）"""
        try:
            # 读取文档内容
            if HAS_DOCX:
                try:
                    doc = Document(doc_path)
                    text_content = ""
                    for paragraph in doc.paragraphs:
                        text_content += paragraph.text + "\n"
                except Exception as docx_error:
                    self.logger.warning(f"python-docx解析失败，使用二进制方式: {docx_error}")
                    # 降级到二进制读取
                    with open(doc_path, 'rb') as f:
                        binary_data = f.read()
                    text_content = self._extract_text_from_word_doc(binary_data)
            else:
                # 如果没有python-docx，尝试二进制读取
                with open(doc_path, 'rb') as f:
                    binary_data = f.read()
                text_content = self._extract_text_from_word_doc(binary_data)
            
            # 提取占位符
            placeholders_info = self._extract_placeholders(text_content)
            
            # 分类占位符
            stats_placeholders = []
            chart_placeholders = []
            
            for placeholder in placeholders_info:
                placeholder_dict = {
                    "description": placeholder.description,
                    "placeholder_text": f"{{{{{placeholder.name}}}}}",
                    "placeholder_name": placeholder.name,
                    "placeholder_type": placeholder.type,
                    "content_type": placeholder.content_type.value if hasattr(placeholder.content_type, 'value') else str(placeholder.content_type)
                }
                
                # 根据类型分类
                if placeholder.type in ["统计", "statistic", "count", "sum", "average"]:
                    placeholder_dict["analysis_requirements"] = {
                        "data_operation": self._infer_data_operation(placeholder.name),
                        "aggregation_type": self._infer_aggregation_type(placeholder.name),
                        "time_dimension": "时间" in placeholder.name or "年" in placeholder.name or "月" in placeholder.name,
                        "geographic_dimension": "地区" in placeholder.name or "区域" in placeholder.name,
                        "requires_grouping": "分组" in placeholder.name or "按" in placeholder.name
                    }
                    stats_placeholders.append(placeholder_dict)
                elif placeholder.type in ["图表", "chart", "graph"]:
                    placeholder_dict["chart_requirements"] = {
                        "chart_type": self._infer_chart_type(placeholder.name),
                        "data_series": 1,
                        "x_axis": "category",
                        "y_axis": "value",
                        "show_legend": True
                    }
                    chart_placeholders.append(placeholder_dict)
                else:
                    # 默认归类为统计
                    placeholder_dict["analysis_requirements"] = {
                        "data_operation": "unknown",
                        "aggregation_type": None,
                        "time_dimension": False,
                        "geographic_dimension": False,
                        "requires_grouping": False
                    }
                    stats_placeholders.append(placeholder_dict)
            
            return {
                "success": True,
                "stats_placeholders": stats_placeholders,
                "chart_placeholders": chart_placeholders,
                "total_placeholders": len(placeholders_info),
                "document_path": doc_path
            }
            
        except Exception as e:
            self.logger.error(f"解析DOC占位符失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "stats_placeholders": [],
                "chart_placeholders": [],
                "total_placeholders": 0,
                "document_path": doc_path
            }
    
    def _infer_data_operation(self, placeholder_name: str) -> str:
        """推断数据操作类型"""
        name_lower = placeholder_name.lower()
        if any(keyword in name_lower for keyword in ['总数', '数量', 'count', '个数']):
            return 'count'
        elif any(keyword in name_lower for keyword in ['总和', '合计', 'sum', '总计']):
            return 'sum'
        elif any(keyword in name_lower for keyword in ['平均', 'avg', 'average', '均值']):
            return 'average'
        elif any(keyword in name_lower for keyword in ['最大', 'max', '最高']):
            return 'max'
        elif any(keyword in name_lower for keyword in ['最小', 'min', '最低']):
            return 'min'
        else:
            return 'unknown'
    
    def _infer_aggregation_type(self, placeholder_name: str) -> Optional[str]:
        """推断聚合类型"""
        name_lower = placeholder_name.lower()
        if any(keyword in name_lower for keyword in ['按月', '月度', 'monthly']):
            return 'monthly'
        elif any(keyword in name_lower for keyword in ['按年', '年度', 'yearly']):
            return 'yearly'
        elif any(keyword in name_lower for keyword in ['按日', '日度', 'daily']):
            return 'daily'
        elif any(keyword in name_lower for keyword in ['按区域', '地区', 'region']):
            return 'regional'
        else:
            return None
    
    def _infer_chart_type(self, placeholder_name: str) -> str:
        """推断图表类型"""
        name_lower = placeholder_name.lower()
        if any(keyword in name_lower for keyword in ['柱状图', 'bar', '条形图']):
            return 'bar'
        elif any(keyword in name_lower for keyword in ['折线图', 'line', '趋势图']):
            return 'line'
        elif any(keyword in name_lower for keyword in ['饼图', 'pie', '饼状图']):
            return 'pie'
        elif any(keyword in name_lower for keyword in ['散点图', 'scatter']):
            return 'scatter'
        else:
            return 'bar'  # 默认柱状图


class ContentFormatter:
    """内容格式化器"""
    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
    
    def format_content(self, raw_content: Any, content_type: ContentType, 
                      format_rules: Dict[str, Any] = None) -> ProcessedContent:
        """格式化内容"""
        format_rules = format_rules or {}
        
        if content_type == ContentType.TEXT:
            formatted = self._format_text(raw_content, format_rules)
        elif content_type == ContentType.TABLE:
            formatted = self._format_table(raw_content, format_rules)
        elif content_type == ContentType.CHART:
            formatted = self._format_chart(raw_content, format_rules)
        elif content_type == ContentType.LIST:
            formatted = self._format_list(raw_content, format_rules)
        elif content_type == ContentType.IMAGE:
            formatted = self._format_image(raw_content, format_rules)
        else:
            formatted = str(raw_content)
        
        return ProcessedContent(
            placeholder_name="",
            content_type=content_type,
            raw_content=raw_content,
            formatted_content=formatted,
            metadata={"formatted_at": datetime.now().isoformat()}
        )
    
    def _format_text(self, content: Any, format_rules: Dict) -> str:
        """格式化文本内容"""
        if isinstance(content, (dict, list)):
            # 如果是结构化数据，转换为文本
            if isinstance(content, dict):
                if "value" in content:
                    return str(content["value"])
                elif "results" in content:
                    results = content["results"]
                    if isinstance(results, dict):
                        # 提取主要信息
                        if "numeric_statistics" in results:
                            stats = results["numeric_statistics"]
                            text_parts = []
                            for field, field_stats in stats.items():
                                mean = field_stats.get("mean", 0)
                                count = field_stats.get("count", 0)
                                text_parts.append(f"{field}: 平均值 {mean:.2f}，共 {count} 项")
                            return "; ".join(text_parts)
                        else:
                            return json.dumps(results, ensure_ascii=False, indent=2)
                    else:
                        return str(results)
                else:
                    # 选择关键字段
                    key_fields = ["total", "count", "sum", "average", "result", "value"]
                    for field in key_fields:
                        if field in content:
                            return str(content[field])
                    return json.dumps(content, ensure_ascii=False)
            elif isinstance(content, list) and content:
                if isinstance(content[0], dict):
                    # 如果是单值结果
                    if len(content) == 1:
                        return self._format_text(content[0], format_rules)
                    # 如果是多行数据，格式化为摘要
                    return f"共 {len(content)} 项数据"
                else:
                    return ", ".join(str(item) for item in content[:5])
        
        return str(content)
    
    def _format_table(self, content: Any, format_rules: Dict) -> str:
        """格式化表格内容"""
        if not isinstance(content, list) or not content:
            return "无表格数据"
        
        # 如果是查询结果格式
        if isinstance(content, dict) and "data" in content:
            content = content["data"]
        
        if not isinstance(content, list) or not content:
            return "无表格数据"
        
        # 获取列名
        if isinstance(content[0], dict):
            columns = list(content[0].keys())
        else:
            return "无效表格数据格式"
        
        # 限制显示行数
        max_rows = format_rules.get("max_rows", 10)
        display_data = content[:max_rows]
        
        # 生成表格字符串
        table_lines = []
        
        # 表头
        header = " | ".join(columns)
        table_lines.append(header)
        table_lines.append(" | ".join(["---"] * len(columns)))
        
        # 数据行
        for row in display_data:
            if isinstance(row, dict):
                values = []
                for col in columns:
                    value = row.get(col, "")
                    # 格式化数值
                    if isinstance(value, float):
                        values.append(f"{value:.2f}")
                    else:
                        values.append(str(value))
                table_lines.append(" | ".join(values))
        
        # 如果有更多数据，添加省略号
        if len(content) > max_rows:
            table_lines.append(f"... (共 {len(content)} 行)")
        
        return "\n".join(table_lines)
    
    def _format_chart(self, content: Any, format_rules: Dict) -> str:
        """格式化图表内容"""
        if isinstance(content, dict):
            chart_type = content.get("chart_type", "unknown")
            data_points = 0
            
            if "chart_data" in content:
                data_points = len(content["chart_data"])
            elif "data" in content:
                data_points = len(content["data"]) if isinstance(content["data"], list) else 0
            
            return f"[{chart_type}图表，包含{data_points}个数据点]"
        
        return "[图表内容]"
    
    def _format_list(self, content: Any, format_rules: Dict) -> str:
        """格式化列表内容"""
        if isinstance(content, list):
            list_items = []
            for i, item in enumerate(content[:10]):  # 限制10项
                if isinstance(item, dict):
                    # 提取主要信息
                    if "name" in item and "value" in item:
                        list_items.append(f"{i+1}. {item['name']}: {item['value']}")
                    else:
                        list_items.append(f"{i+1}. {json.dumps(item, ensure_ascii=False)}")
                else:
                    list_items.append(f"{i+1}. {item}")
            
            if len(content) > 10:
                list_items.append(f"... (共 {len(content)} 项)")
            
            return "\n".join(list_items)
        
        return str(content)
    
    def _format_image(self, content: Any, format_rules: Dict) -> str:
        """格式化图片内容"""
        if isinstance(content, dict) and "image_path" in content:
            return f"[图片: {content['image_path']}]"
        elif isinstance(content, str) and content.startswith("data:image"):
            return "[Base64图片内容]"
        else:
            return "[图片内容]"


class WordDocumentGenerator:
    """Word文档生成器"""
    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
    
    def generate_document(self, template: DocumentTemplate, 
                         processed_contents: List[ProcessedContent],
                         output_path: str = None) -> DocumentGenerationResult:
        """生成Word文档"""
        start_time = time.time()
        
        try:
            # 生成输出路径
            if not output_path:
                timestamp = int(time.time())
                output_path = f"/tmp/report_{timestamp}.docx"
            
            if HAS_DOCX:
                # 使用python-docx生成真实文档
                result = self._generate_docx_document(template, processed_contents, output_path)
            else:
                # 使用模拟生成
                result = self._generate_mock_document(template, processed_contents, output_path)
            
            generation_time = time.time() - start_time
            result.generation_time = generation_time
            
            return result
            
        except Exception as e:
            generation_time = time.time() - start_time
            self.logger.error(f"文档生成失败: {e}")
            
            return DocumentGenerationResult(
                success=False,
                document_path="",
                document_format=DocumentFormat.DOCX,
                generation_time=generation_time,
                placeholder_count=len(template.placeholders),
                processed_placeholders=0,
                error_message=str(e)
            )
    
    def _generate_docx_document(self, template: DocumentTemplate,
                               processed_contents: List[ProcessedContent],
                               output_path: str) -> DocumentGenerationResult:
        """使用python-docx生成真实文档"""
        doc = Document()
        
        # 设置文档标题
        title = doc.add_heading(template.name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 处理模板内容
        content_lines = template.content.split('\n')
        content_map = {pc.placeholder_name: pc for pc in processed_contents}
        
        failed_placeholders = []
        processed_count = 0
        warnings = []
        
        for line in content_lines:
            line = line.strip()
            if not line:
                # 添加空行
                doc.add_paragraph()
                continue
            
            # 检查是否包含占位符
            placeholders_in_line = re.findall(r'\{([^}]+)\}', line)
            
            if placeholders_in_line:
                # 替换占位符
                processed_line = line
                
                for placeholder_name in placeholders_in_line:
                    if placeholder_name in content_map:
                        processed_content = content_map[placeholder_name]
                        replacement = processed_content.formatted_content
                        processed_line = processed_line.replace(f"{{{placeholder_name}}}", replacement)
                        processed_count += 1
                    else:
                        processed_line = processed_line.replace(f"{{{placeholder_name}}}", f"[未处理: {placeholder_name}]")
                        failed_placeholders.append(placeholder_name)
                        warnings.append(f"占位符 {placeholder_name} 未找到对应内容")
                
                # 根据内容类型添加到文档
                if any(pc.content_type == ContentType.TABLE for pc in content_map.values()
                       if pc.placeholder_name in placeholders_in_line):
                    # 如果包含表格内容，添加为表格
                    self._add_table_to_doc(doc, processed_line, content_map, placeholders_in_line)
                else:
                    # 添加为段落
                    self._add_paragraph_to_doc(doc, processed_line, line)
            else:
                # 普通文本行
                self._add_paragraph_to_doc(doc, line)
        
        # 保存文档
        doc.save(output_path)
        
        return DocumentGenerationResult(
            success=True,
            document_path=output_path,
            document_format=DocumentFormat.DOCX,
            generation_time=0,  # 会在上级函数中设置
            placeholder_count=len(template.placeholders),
            processed_placeholders=processed_count,
            failed_placeholders=failed_placeholders,
            warnings=warnings
        )
    
    def _generate_mock_document(self, template: DocumentTemplate,
                               processed_contents: List[ProcessedContent],
                               output_path: str) -> DocumentGenerationResult:
        """生成模拟文档（当没有python-docx时）"""
        content_map = {pc.placeholder_name: pc for pc in processed_contents}
        
        # 替换模板中的占位符
        final_content = template.content
        processed_count = 0
        failed_placeholders = []
        
        for placeholder in template.placeholders:
            placeholder_pattern = f"{{{placeholder.name}}}"
            
            if placeholder.name in content_map:
                replacement = content_map[placeholder.name].formatted_content
                final_content = final_content.replace(placeholder_pattern, replacement)
                processed_count += 1
            else:
                final_content = final_content.replace(placeholder_pattern, f"[未处理: {placeholder.name}]")
                failed_placeholders.append(placeholder.name)
        
        # 保存为文本文件（模拟）
        text_output_path = output_path.replace('.docx', '.txt')
        
        try:
            with open(text_output_path, 'w', encoding='utf-8') as f:
                f.write(f"# {template.name}\n\n")
                f.write(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"占位符总数: {len(template.placeholders)}\n")
                f.write(f"已处理: {processed_count}\n")
                f.write(f"未处理: {len(failed_placeholders)}\n\n")
                f.write("=" * 50 + "\n\n")
                f.write(final_content)
            
            self.logger.info(f"模拟文档已保存: {text_output_path}")
            
        except Exception as e:
            self.logger.warning(f"保存模拟文档失败: {e}")
            text_output_path = output_path
        
        return DocumentGenerationResult(
            success=True,
            document_path=text_output_path,
            document_format=DocumentFormat.DOCX,
            generation_time=0,
            placeholder_count=len(template.placeholders),
            processed_placeholders=processed_count,
            failed_placeholders=failed_placeholders,
            warnings=["使用模拟文档生成器（未安装python-docx）"] if not HAS_DOCX else [],
            metadata={"mock_generation": not HAS_DOCX}
        )
    
    def _add_paragraph_to_doc(self, doc, content: str, original_line: str = None):
        """添加段落到文档"""
        if content.startswith('#'):
            # 标题
            level = min(len(content) - len(content.lstrip('#')), 6)
            title_text = content.lstrip('#').strip()
            doc.add_heading(title_text, level)
        else:
            # 普通段落
            para = doc.add_paragraph(content)
            
            # 如果内容包含数字，设置为右对齐
            if re.search(r'\d+\.?\d*', content):
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _add_table_to_doc(self, doc, content: str, content_map: Dict,
                          placeholders_in_line: List[str]):
        """添加表格到文档"""
        # 简化实现：将表格内容作为格式化文本添加
        doc.add_paragraph(content)
        
        # 这里可以扩展为真正的表格插入逻辑
        for placeholder_name in placeholders_in_line:
            if (placeholder_name in content_map and 
                content_map[placeholder_name].content_type == ContentType.TABLE):
                
                table_content = content_map[placeholder_name].formatted_content
                # 简单添加为代码块样式
                para = doc.add_paragraph(table_content)
                para.style = 'Normal'


class DocumentPipeline:
    """文档生成流水线"""
    
    def __init__(self):
        self.pipeline_id = str(uuid.uuid4())
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
        
        # 初始化组件
        self.template_parser = TemplateParser()
        self.content_formatter = ContentFormatter()
        self.document_generator = WordDocumentGenerator()
        
        # 流水线统计
        self.stats = {
            "templates_processed": 0,
            "documents_generated": 0,
            "total_processing_time": 0,
            "success_rate": 0
        }
    
    async def process_template(self, template_content: str, placeholder_data: Dict[str, Any],
                              output_path: str = None, template_name: str = "Generated Report") -> DocumentGenerationResult:
        """处理模板并生成文档"""
        start_time = time.time()
        
        try:
            self.logger.info(f"开始处理模板: {template_name}")
            
            # 1. 解析模板
            template = self.template_parser.parse_template(template_content)
            template.name = template_name
            
            # 2. 处理占位符数据
            processed_contents = []
            
            for placeholder in template.placeholders:
                if placeholder.name in placeholder_data:
                    raw_content = placeholder_data[placeholder.name]
                    
                    # 格式化内容
                    processed_content = self.content_formatter.format_content(
                        raw_content=raw_content,
                        content_type=placeholder.content_type,
                        format_rules=placeholder.format_rules
                    )
                    processed_content.placeholder_name = placeholder.name
                    
                    processed_contents.append(processed_content)
            
            # 3. 生成文档
            result = self.document_generator.generate_document(
                template=template,
                processed_contents=processed_contents,
                output_path=output_path
            )
            
            # 4. 更新统计信息
            self.stats["templates_processed"] += 1
            if result.success:
                self.stats["documents_generated"] += 1
            
            total_time = time.time() - start_time
            self.stats["total_processing_time"] += total_time
            self.stats["success_rate"] = self.stats["documents_generated"] / self.stats["templates_processed"]
            
            self.logger.info(f"模板处理完成: {template_name}, 耗时: {total_time:.2f}秒")
            
            return result
            
        except Exception as e:
            self.logger.error(f"模板处理失败: {e}")
            return DocumentGenerationResult(
                success=False,
                document_path="",
                document_format=DocumentFormat.DOCX,
                generation_time=time.time() - start_time,
                placeholder_count=0,
                processed_placeholders=0,
                error_message=str(e)
            )
    
    def get_pipeline_stats(self) -> Dict[str, Any]:
        """获取流水线统计信息"""
        avg_processing_time = (
            self.stats["total_processing_time"] / self.stats["templates_processed"]
            if self.stats["templates_processed"] > 0 else 0
        )
        
        return {
            **self.stats,
            "pipeline_id": self.pipeline_id,
            "average_processing_time": avg_processing_time,
            "uptime": datetime.now().isoformat()
        }
    
    async def batch_process_templates(self, template_requests: List[Dict[str, Any]]) -> List[DocumentGenerationResult]:
        """批量处理模板"""
        self.logger.info(f"开始批量处理 {len(template_requests)} 个模板")
        
        results = []
        
        for i, request in enumerate(template_requests):
            self.logger.info(f"处理模板 {i+1}/{len(template_requests)}")
            
            result = await self.process_template(
                template_content=request.get("template_content", ""),
                placeholder_data=request.get("placeholder_data", {}),
                output_path=request.get("output_path"),
                template_name=request.get("template_name", f"Template_{i+1}")
            )
            
            results.append(result)
        
        self.logger.info(f"批量处理完成，成功: {sum(1 for r in results if r.success)}/{len(results)}")
        
        return results


# 全局流水线实例
document_pipeline = DocumentPipeline()


# 便捷函数
async def generate_document_from_template(template_content: str, placeholder_data: Dict[str, Any],
                                        output_path: str = None, template_name: str = "Report") -> DocumentGenerationResult:
    """从模板生成文档的便捷函数"""
    return await document_pipeline.process_template(
        template_content=template_content,
        placeholder_data=placeholder_data,
        output_path=output_path,
        template_name=template_name
    )


def create_sample_template() -> str:
    """创建示例模板"""
    return """
# {报告标题}

## 执行摘要
{执行摘要}

## 数据概览
在本报告期间，我们收集和分析了以下数据：

### 基础统计
- 总记录数：{总记录数}
- 平均值：{平均值}
- 数据完整性：{数据完整性}

## 详细分析

### 分类统计
{分类统计表格}

### 趋势分析
{趋势分析图表}

### 关键洞察
{关键洞察列表}

## 结论和建议
基于以上分析，我们得出以下结论：

{结论和建议}

## 附录
- 数据来源：{数据来源}
- 生成时间：{生成时间}
- 报告版本：{报告版本}
"""


if __name__ == "__main__":
    # 测试用例
    async def test_document_pipeline():
        """测试文档生成流水线"""
        print("🧪 测试文档生成流水线...")
        
        # 创建示例模板
        template_content = create_sample_template()
        
        # 准备占位符数据
        placeholder_data = {
            "报告标题": "月度数据分析报告",
            "执行摘要": "本报告分析了本月的业务数据，发现了多个重要趋势和机会点。",
            "总记录数": 1234,
            "平均值": 89.5,
            "数据完整性": "95.2%",
            "分类统计表格": [
                {"类别": "产品A", "数量": 150, "占比": "30%"},
                {"类别": "产品B", "数量": 120, "占比": "24%"},
                {"类别": "产品C", "数量": 230, "占比": "46%"}
            ],
            "趋势分析图表": {
                "chart_type": "line",
                "data": [{"month": "1月", "value": 100}, {"month": "2月", "value": 120}]
            },
            "关键洞察列表": [
                "产品C表现最佳，占比达到46%",
                "整体趋势呈上升态势",
                "数据质量良好，完整性超过95%"
            ],
            "结论和建议": "建议继续关注产品C的发展，同时改进产品A和B的市场策略。",
            "数据来源": "内部业务系统",
            "生成时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "报告版本": "v1.0"
        }
        
        # 生成文档
        result = await generate_document_from_template(
            template_content=template_content,
            placeholder_data=placeholder_data,
            template_name="测试报告"
        )
        
        print(f"✅ 文档生成完成!")
        print(f"成功: {result.success}")
        print(f"文档路径: {result.document_path}")
        print(f"生成时间: {result.generation_time:.2f}秒")
        print(f"占位符处理: {result.processed_placeholders}/{result.placeholder_count}")
        
        if result.warnings:
            print("⚠️ 警告:")
            for warning in result.warnings:
                print(f"  - {warning}")
        
        if result.failed_placeholders:
            print("❌ 失败的占位符:")
            for failed in result.failed_placeholders:
                print(f"  - {failed}")
        
        # 显示流水线统计
        stats = document_pipeline.get_pipeline_stats()
        print(f"\n📊 流水线统计:")
        print(f"  - 处理模板数: {stats['templates_processed']}")
        print(f"  - 生成文档数: {stats['documents_generated']}")
        print(f"  - 成功率: {stats['success_rate']:.2%}")
        print(f"  - 平均处理时间: {stats['average_processing_time']:.2f}秒")
    
    # 运行测试
    asyncio.run(test_document_pipeline())