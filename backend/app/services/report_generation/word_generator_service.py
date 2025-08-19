import base64
import io
import logging
import re
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, Optional

import docx
from docx.shared import Inches

from app.services.storage.file_storage_service import file_storage_service

logger = logging.getLogger(__name__)


class WordGeneratorService:
    # Regex to find <img ...> tags with base64 data
    IMG_REGEX = re.compile(r'<img src="data:image/png;base64,([^"]+)">')

    def generate_report_from_template(
        self,
        template_content: str,
        placeholder_values: Dict[str, Any],
        title: str = "自动生成报告",
        format: str = "docx"
    ) -> str:
        """
        基于模板生成报告文件
        
        Args:
            template_content: 模板内容 (hex编码的二进制数据或文本)
            placeholder_values: 占位符值字典
            title: 报告标题
            format: 文件格式 (docx)
            
        Returns:
            存储后的文件路径
        """
        try:
            # 1. 处理模板内容
            if self._is_binary_template(template_content):
                # 处理二进制模板 (hex编码)
                return self._generate_from_binary_template(
                    template_content, placeholder_values, title
                )
            else:
                # 处理文本模板
                return self._generate_from_text_template(
                    template_content, placeholder_values, title
                )
                
        except Exception as e:
            logger.error(f"基于模板的报告生成失败: {str(e)}")
            # 降级到普通生成
            processed_content = self._replace_placeholders_in_text(
                template_content, placeholder_values
            )
            return self.generate_report(processed_content, title, format)

    def _is_binary_template(self, content: str) -> bool:
        """判断是否为二进制模板内容"""
        if not content:
            return False
        # 检查是否为hex编码
        content_clean = content.replace(' ', '').replace('\n', '')
        return len(content_clean) > 100 and all(c in '0123456789ABCDEFabcdef' for c in content_clean)

    def _generate_from_binary_template(
        self,
        template_content: str,
        placeholder_values: Dict[str, Any],
        title: str
    ) -> str:
        """从二进制模板生成报告"""
        try:
            # 解析hex编码的二进制数据
            binary_data = bytes.fromhex(template_content.replace(' ', '').replace('\n', ''))
            
            # 加载为Word文档
            template_buffer = BytesIO(binary_data)
            doc = docx.Document(template_buffer)
            
            # 替换占位符
            self._replace_placeholders_in_doc(doc, placeholder_values)
            
            # 保存到内存
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            # 生成文件名并上传 (保持DOCX格式)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{title}_{timestamp}.docx"
            
            file_info = file_storage_service.upload_file(
                file_data=doc_buffer,
                original_filename=filename,
                file_type="report",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            logger.info(f"基于二进制模板的报告生成成功: {file_info['file_path']}")
            return file_info['file_path']
            
        except Exception as e:
            logger.warning(f"二进制模板处理失败，降级到文本处理: {e}")
            raise

    def _generate_from_text_template(
        self,
        template_content: str,
        placeholder_values: Dict[str, Any],
        title: str
    ) -> str:
        """从文本模板生成报告"""
        # 替换占位符
        processed_content = self._replace_placeholders_in_text(template_content, placeholder_values)
        
        # 使用普通生成方法
        return self.generate_report(processed_content, title)

    def _replace_placeholders_in_doc(self, doc, placeholder_values: Dict[str, Any]):
        """在Word文档中替换占位符"""
        # 处理占位符值，提取实际值
        processed_values = {}
        for key, value_info in placeholder_values.items():
            if isinstance(value_info, dict) and "value" in value_info:
                extracted_value = self._extract_value_from_result(value_info["value"])
                processed_values[key] = str(extracted_value)
            else:
                extracted_value = self._extract_value_from_result(value_info)
                processed_values[key] = str(extracted_value)
        
        # 替换段落中的占位符
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            for key, value in processed_values.items():
                # 支持多种占位符格式
                patterns = [
                    f"{{{{{key}}}}}",  # {{key}}
                    f"{{{key}}}",      # {key}
                ]
                for pattern in patterns:
                    if pattern in paragraph.text:
                        paragraph.text = paragraph.text.replace(pattern, value)
        
        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in processed_values.items():
                        # 支持多种占位符格式
                        patterns = [
                            f"{{{{{key}}}}}",  # {{key}}
                            f"{{{key}}}",      # {key}
                        ]
                        for pattern in patterns:
                            if pattern in cell.text:
                                cell.text = cell.text.replace(pattern, value)

    def _extract_value_from_result(self, value: Any) -> Any:
        """从查询结果中提取实际数值"""
        try:
            # 如果是DorisQueryResult对象
            if hasattr(value, 'data') and hasattr(value, 'execution_time'):
                # 提取DataFrame中的数据
                if hasattr(value.data, 'iloc') and len(value.data) > 0:
                    # pandas DataFrame，获取第一行第一列的值
                    return value.data.iloc[0, 0] if not value.data.empty else 0
                elif hasattr(value.data, '__iter__'):
                    # 其他可迭代对象
                    data_list = list(value.data)
                    if data_list and len(data_list) > 0:
                        first_row = data_list[0]
                        if isinstance(first_row, dict) and first_row:
                            # 获取第一个值
                            return next(iter(first_row.values()))
                        return first_row
                return 0
            
            # 如果是包含DorisQueryResult的字典
            if isinstance(value, dict):
                if 'data' in value and hasattr(value['data'], 'data'):
                    return self._extract_value_from_result(value['data'])
                # 普通字典，尝试获取常见的数值字段
                for key in ['count', 'total', 'value', 'result']:
                    if key in value:
                        return self._extract_value_from_result(value[key])
            
            # 如果是pandas DataFrame
            if hasattr(value, 'iloc') and len(value) > 0:
                return value.iloc[0, 0] if not value.empty else 0
                
            # 如果是列表或元组
            if isinstance(value, (list, tuple)) and value:
                first_item = value[0]
                if isinstance(first_item, dict) and first_item:
                    return next(iter(first_item.values()))
                return first_item
                
            # 直接返回原值
            return value
            
        except Exception as e:
            logger.warning(f"提取查询结果数值失败: {e}, 使用原值: {value}")
            return value

    def _replace_placeholders_in_text(self, content: str, placeholder_values: Dict[str, Any]) -> str:
        """在文本中替换占位符"""
        processed_content = content
        
        # 处理占位符值，提取实际值
        for key, value_info in placeholder_values.items():
            if isinstance(value_info, dict) and "value" in value_info:
                extracted_value = self._extract_value_from_result(value_info["value"])
                value = str(extracted_value)
            else:
                extracted_value = self._extract_value_from_result(value_info)
                value = str(extracted_value)
                
            # 支持多种占位符格式
            patterns = [
                f"{{{{{key}}}}}",  # {{key}}
                f"{{{key}}}",      # {key}
            ]
            for pattern in patterns:
                processed_content = processed_content.replace(pattern, value)
        return processed_content

    def _get_file_format_info(self, format_type: str) -> Dict[str, str]:
        """获取文件格式信息"""
        format_map = {
            'docx': {
                'extension': '.docx',
                'content_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            },
            'xlsx': {
                'extension': '.xlsx', 
                'content_type': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            },
            'txt': {
                'extension': '.txt',
                'content_type': 'text/plain'
            },
            'html': {
                'extension': '.html',
                'content_type': 'text/html'
            }
        }
        return format_map.get(format_type, format_map['docx'])

    def generate_report(
        self,
        content: str,
        title: str = "自动生成报告",
        format: str = "docx"
    ) -> str:
        """
        生成报告文件并存储到MinIO或本地
        
        Args:
            content: 报告内容
            title: 报告标题  
            format: 文件格式 (docx)
            
        Returns:
            存储后的文件路径
        """
        try:
            # 创建Word文档
            doc = docx.Document()
            
            # 添加标题
            doc.add_heading(title, 0)
            
            # 添加生成时间
            doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("")  # 空行
            
            # 处理内容
            self._process_content(doc, content)
            
            # 将文档保存到内存
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            # 获取格式信息并生成文件名
            format_info = self._get_file_format_info(format)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{title}_{timestamp}{format_info['extension']}"
            
            # 上传到存储服务
            file_info = file_storage_service.upload_file(
                file_data=doc_buffer,
                original_filename=filename,
                file_type="report",
                content_type=format_info['content_type']
            )
            
            logger.info(f"报告生成成功: {file_info['file_path']}")
            return file_info['file_path']
            
        except Exception as e:
            logger.error(f"报告生成失败: {str(e)}")
            raise

    def generate_report_from_content(self, composed_content: str, output_path: str):
        """
        Generates a .docx report from a string containing the fully composed content.
        This content can include text and special <img> tags for base64 images.
        
        保留原有方法以保持向后兼容
        """
        doc = docx.Document()

        # Split the content into paragraphs based on newlines
        paragraphs = composed_content.split("\n")

        for para_text in paragraphs:
            self._process_paragraph(doc, para_text)

        doc.save(output_path)

    def _process_content(self, doc, content: str):
        """处理报告内容，支持文本和图像"""
        # 分段处理
        paragraphs = content.split("\n")
        
        for para_text in paragraphs:
            self._process_paragraph(doc, para_text)

    def _process_paragraph(self, doc, para_text: str):
        """
        Processes a single paragraph of text, adding text and images to the document.
        """
        # Find all image tags in the paragraph
        img_matches = list(self.IMG_REGEX.finditer(para_text))

        if not img_matches:
            # If no images, add the whole paragraph as text
            doc.add_paragraph(para_text)
            return

        current_pos = 0
        for match in img_matches:
            # Add text before the image
            start, end = match.span()
            if start > current_pos:
                doc.add_paragraph(para_text[current_pos:start])

            # Add the image
            base64_data = match.group(1)
            try:
                image_stream = io.BytesIO(base64.b64decode(base64_data))
                doc.add_picture(image_stream, width=Inches(6.0))
            except Exception as e:
                print(f"Error decoding or adding picture: {e}")
                # Add a placeholder text on error
                doc.add_paragraph(f"[Image could not be loaded: {e}]")

            current_pos = end

        # Add any remaining text after the last image
        if current_pos < len(para_text):
            doc.add_paragraph(para_text[current_pos:])


word_generator_service = WordGeneratorService()
