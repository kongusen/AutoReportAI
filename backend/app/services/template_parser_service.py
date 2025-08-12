import logging
import re
from dataclasses import dataclass
from datetime import datetime
from typing import Any, Dict, List, Optional, Union

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from .ai_integration import ChartGenerator, ChartConfig, ChartResult
from .ai_integration import ContentGenerator, FormatConfig, GeneratedContent
# Enhanced Agent system integration
from .agents.orchestrator import orchestrator
from .agents.base import AgentResult
from .agents.core.intelligent_pipeline_orchestrator import pipeline_orchestrator, PipelineContext

logger = logging.getLogger(__name__)


@dataclass
class ReplacementResult:
    """替换结果"""

    success: bool
    original_placeholder: str
    replacement_content: str
    content_type: str
    error_message: Optional[str] = None
    chart_path: Optional[str] = None  # 图表文件路径


@dataclass
class TemplateProcessingResult:
    """模板处理结果"""

    success: bool
    processed_file_path: Optional[str] = None
    replacements: List[ReplacementResult] = None
    total_placeholders: int = 0
    successful_replacements: int = 0
    processing_time: float = 0.0
    error_message: Optional[str] = None


class TemplateParser:
    # 原有的占位符正则表达式
    PLACEHOLDER_REGEX = re.compile(
        r"\{\{(?P<scalar>[\w\s]+?)\s*(?:\s+\"(?P<s_desc>.*?)\")?\s*\}\}|"
        r"\[(?P<type>chart|table):(?P<name>[\w\s]+?)\s*(?:\s+\"(?P<ct_desc>.*?)\")?\s*\]"
    )

    # 统一的智能占位符正则表达式 {{类型:描述}}
    INTELLIGENT_PLACEHOLDER_REGEX = re.compile(r"\{\{([^:]+):([^}]+)\}\}")
    
    # 专用的统计和图表占位符正则表达式
    STATS_PLACEHOLDER_REGEX = re.compile(r"\{\{统计\s*:\s*([^}]+)\}\}")
    CHART_PLACEHOLDER_REGEX = re.compile(r"\{\{图表\s*:\s*([^}]+)\}\}")

    def __init__(self):
        self.format_config = FormatConfig()
        self.chart_config = ChartConfig()

    def parse(self, file_path: str) -> Dict[str, List[Dict[str, Any]]]:
        """
        Parses a .docx file to extract placeholders for scalar values, charts, and tables.

        Each placeholder can have an optional description.
        """
        doc = docx.Document(file_path)
        placeholders = []
        found_keys = set()

        # Combine text from paragraphs and tables for parsing
        full_text = "\n".join([p.text for p in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += "\n" + cell.text

        for match in self.PLACEHOLDER_REGEX.finditer(full_text):
            if match.group("scalar"):
                key = match.group("scalar").strip()
                if key not in found_keys:
                    placeholders.append(
                        {
                            "name": key,
                            "type": "scalar",
                            "description": match.group("s_desc") or "",
                        }
                    )
                    found_keys.add(key)
            else:
                key = match.group("name").strip()
                if key not in found_keys:
                    placeholders.append(
                        {
                            "name": key,
                            "type": match.group("type"),
                            "description": match.group("ct_desc") or "",
                        }
                    )
                    found_keys.add(key)

        return {"placeholders": placeholders}

    def parse_intelligent_placeholders(
        self, file_path: str
    ) -> Dict[str, List[Dict[str, Any]]]:
        """
        解析智能占位符 {{类型:描述}} 格式
        """
        doc = docx.Document(file_path)
        placeholders = []
        found_keys = set()

        # 获取完整文本
        full_text = "\n".join([p.text for p in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += "\n" + cell.text

        # 查找智能占位符
        for match in self.INTELLIGENT_PLACEHOLDER_REGEX.finditer(full_text):
            placeholder_type = match.group(1).strip()
            description = match.group(2).strip()
            placeholder_text = match.group(0)

            if placeholder_text not in found_keys:
                placeholders.append(
                    {
                        "placeholder_text": placeholder_text,
                        "placeholder_type": placeholder_type,
                        "description": description,
                        "position": match.start(),
                    }
                )
                found_keys.add(placeholder_text)

        return {"intelligent_placeholders": placeholders}

    def parse_doc_placeholders(
        self, file_path: str
    ) -> Dict[str, List[Dict[str, Any]]]:
        """
        专门解析DOC文件中的统计和图表占位符
        支持格式：{{统计:描述}} 和 {{图表:描述}}
        """
        doc = docx.Document(file_path)
        stats_placeholders = []
        chart_placeholders = []
        found_keys = set()

        # 获取完整文本
        full_text = "\n".join([p.text for p in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += "\n" + cell.text

        # 查找统计占位符
        for match in self.STATS_PLACEHOLDER_REGEX.finditer(full_text):
            description = match.group(1).strip()
            placeholder_text = match.group(0)

            if placeholder_text not in found_keys:
                stats_placeholders.append(
                    {
                        "placeholder_text": placeholder_text,
                        "placeholder_type": "统计",
                        "description": description,
                        "position": match.start(),
                        "analysis_requirements": self._parse_stats_description(description)
                    }
                )
                found_keys.add(placeholder_text)

        # 查找图表占位符
        for match in self.CHART_PLACEHOLDER_REGEX.finditer(full_text):
            description = match.group(1).strip()
            placeholder_text = match.group(0)

            if placeholder_text not in found_keys:
                chart_placeholders.append(
                    {
                        "placeholder_text": placeholder_text,
                        "placeholder_type": "图表",
                        "description": description,
                        "position": match.start(),
                        "chart_requirements": self._parse_chart_description(description)
                    }
                )
                found_keys.add(placeholder_text)

        return {
            "stats_placeholders": stats_placeholders,
            "chart_placeholders": chart_placeholders,
            "total_count": len(stats_placeholders) + len(chart_placeholders)
        }

    def _parse_stats_description(self, description: str) -> Dict[str, Any]:
        """
        解析统计描述，提取统计需求
        """
        requirements = {
            "operation": "sum",  # 默认求和
            "field_hints": [],
            "groupby_hints": [],
            "filter_hints": [],
            "format_hints": {}
        }
        
        desc_lower = description.lower()
        
        # 识别统计操作类型
        if any(keyword in desc_lower for keyword in ["求和", "总和", "合计", "sum"]):
            requirements["operation"] = "sum"
        elif any(keyword in desc_lower for keyword in ["平均", "均值", "average", "avg"]):
            requirements["operation"] = "avg"
        elif any(keyword in desc_lower for keyword in ["计数", "数量", "count"]):
            requirements["operation"] = "count"
        elif any(keyword in desc_lower for keyword in ["最大", "最高", "max"]):
            requirements["operation"] = "max"
        elif any(keyword in desc_lower for keyword in ["最小", "最低", "min"]):
            requirements["operation"] = "min"
        elif any(keyword in desc_lower for keyword in ["占比", "比例", "percentage", "率"]):
            requirements["operation"] = "percentage"
            
        # 识别分组字段
        if any(keyword in desc_lower for keyword in ["按", "分组", "group by"]):
            # 提取可能的分组字段提示
            if "按" in description:
                group_part = description.split("按")[1].split("的")[0] if "的" in description else description.split("按")[1]
                requirements["groupby_hints"].append(group_part.strip())
                
        # 识别格式要求
        if any(keyword in desc_lower for keyword in ["万", "千", "万元"]):
            requirements["format_hints"]["unit"] = "万"
        elif any(keyword in desc_lower for keyword in ["百分比", "%"]):
            requirements["format_hints"]["format"] = "percentage"
            
        return requirements

    def _parse_chart_description(self, description: str) -> Dict[str, Any]:
        """
        解析图表描述，提取图表需求
        """
        requirements = {
            "chart_type": "auto",  # 自动选择
            "x_field_hints": [],
            "y_field_hints": [],
            "groupby_hints": [],
            "title_hints": "",
            "style_hints": {}
        }
        
        desc_lower = description.lower()
        
        # 识别图表类型
        if any(keyword in desc_lower for keyword in ["柱状图", "条形图", "bar"]):
            requirements["chart_type"] = "bar"
        elif any(keyword in desc_lower for keyword in ["折线图", "线图", "line"]):
            requirements["chart_type"] = "line"
        elif any(keyword in desc_lower for keyword in ["饼图", "pie"]):
            requirements["chart_type"] = "pie"
        elif any(keyword in desc_lower for keyword in ["散点图", "scatter"]):
            requirements["chart_type"] = "scatter"
        elif any(keyword in desc_lower for keyword in ["趋势", "变化"]):
            requirements["chart_type"] = "line"
        elif any(keyword in desc_lower for keyword in ["分布", "构成"]):
            requirements["chart_type"] = "pie"
            
        # 提取标题提示
        if "图" in description:
            requirements["title_hints"] = description
            
        return requirements

    async def process_template_with_intelligent_replacement(
        self,
        template_path: str,
        output_path: str,
        data_source_id: int,
        task_config: Optional[Dict[str, Any]] = None,
    ) -> TemplateProcessingResult:
        """
        使用智能管道处理模板 - 通过完整的Agent pipeline处理

        Args:
            template_path: 模板文件路径
            output_path: 输出文件路径
            data_source_id: 数据源ID
            task_config: 任务配置

        Returns:
            模板处理结果
        """
        start_time = datetime.now()

        try:
            logger.info(f"开始智能管道模板处理: {template_path}")

            # 创建管道上下文
            pipeline_context = PipelineContext(
                template_id=template_path,  # 使用路径作为临时ID
                data_source_id=str(data_source_id),
                user_id=task_config.get("user_id", "system") if task_config else "system",
                template_type=self._detect_template_type(template_path),
                output_format="docx",
                optimization_level=task_config.get("optimization_level", "standard") if task_config else "standard",
                batch_size=task_config.get("batch_size", 10000) if task_config else 10000,
                custom_config={
                    "template_path": template_path,
                    "output_path": output_path,
                    **(task_config or {})
                }
            )

            # 执行完整的智能管道
            pipeline_result = await pipeline_orchestrator.execute(pipeline_context)

            if pipeline_result.success:
                pipeline_data = pipeline_result.data
                
                # 从管道结果构建模板处理结果
                replacements = []
                successful_count = 0
                total_placeholders = 0

                # 提取占位符处理信息
                if hasattr(pipeline_data, 'stage_results'):
                    # 从各个阶段提取处理信息
                    for stage, stage_result in pipeline_data.stage_results.items():
                        if stage_result.success and hasattr(stage_result, 'metadata'):
                            metadata = stage_result.metadata
                            if "placeholders_processed" in metadata:
                                total_placeholders += metadata.get("placeholders_processed", 0)
                                successful_count += metadata.get("successful_replacements", 0)
                
                # 处理输出文件
                success = False
                final_output_path = output_path
                
                if pipeline_data.final_output:
                    # 写入最终输出
                    try:
                        with open(output_path, 'wb') as f:
                            f.write(pipeline_data.final_output)
                        success = True
                        logger.info(f"管道输出已写入: {output_path}")
                    except Exception as e:
                        logger.error(f"写入输出文件失败: {e}")
                        success = False
                elif pipeline_data.output_path:
                    # 复制现有输出文件
                    try:
                        import shutil
                        shutil.copy2(pipeline_data.output_path, output_path)
                        success = True
                        logger.info(f"管道输出已复制到: {output_path}")
                    except Exception as e:
                        logger.error(f"复制输出文件失败: {e}")
                        success = False

                processing_time = (datetime.now() - start_time).total_seconds()

                return TemplateProcessingResult(
                    success=success,
                    processed_file_path=final_output_path if success else None,
                    replacements=replacements,
                    total_placeholders=total_placeholders,
                    successful_replacements=successful_count,
                    processing_time=processing_time,
                )
            
            else:
                # 管道执行失败，回退到传统处理方式
                logger.warning(f"管道处理失败，回退到传统方式: {pipeline_result.error_message}")
                return await self._fallback_to_traditional_processing(
                    template_path, output_path, data_source_id, task_config, start_time
                )

        except Exception as e:
            logger.error(f"智能管道处理失败: {e}")
            # 回退到传统处理方式
            return await self._fallback_to_traditional_processing(
                template_path, output_path, data_source_id, task_config, start_time
            )

    async def _fallback_to_traditional_processing(
        self,
        template_path: str,
        output_path: str,
        data_source_id: int,
        task_config: Optional[Dict[str, Any]],
        start_time: datetime
    ) -> TemplateProcessingResult:
        """回退到传统的占位符处理方式"""
        try:
            logger.info("使用传统方式处理模板占位符")

            # 1. 解析智能占位符
            placeholder_data = self.parse_intelligent_placeholders(template_path)
            placeholders = placeholder_data.get("intelligent_placeholders", [])

            if not placeholders:
                return TemplateProcessingResult(
                    success=False,
                    error_message="未找到智能占位符",
                    processing_time=(datetime.now() - start_time).total_seconds(),
                )

            logger.info(f"找到 {len(placeholders)} 个智能占位符")

            # 2. 处理每个占位符
            replacements = []
            successful_count = 0

            for placeholder in placeholders:
                try:
                    # 使用基础Agent系统处理占位符
                    logger.info(f"处理占位符: {placeholder['placeholder_text']}")
                    
                    # 准备Agent输入数据
                    agent_input = {
                        "placeholder_type": placeholder["placeholder_type"],
                        "description": placeholder["description"],
                        "data_source_id": data_source_id,
                        "context": task_config or {}
                    }
                    
                    # 通过orchestrator处理占位符
                    agent_result = await orchestrator.run(agent_input, task_config)
                    
                    if agent_result.success and agent_result.data:
                        # 从工作流结果中提取最终内容
                        workflow_result = agent_result.data
                        final_content, content_type, chart_path = self._extract_final_content(workflow_result)
                        
                        replacement = ReplacementResult(
                            success=True,
                            original_placeholder=placeholder["placeholder_text"],
                            replacement_content=final_content,
                            content_type=content_type,
                            chart_path=chart_path
                        )
                        successful_count += 1
                    else:
                        replacement = ReplacementResult(
                            success=False,
                            original_placeholder=placeholder["placeholder_text"],
                            replacement_content=placeholder["placeholder_text"],
                            content_type="error",
                            error_message=agent_result.error_message or "Agent处理失败",
                        )

                    replacements.append(replacement)

                except Exception as e:
                    logger.error(f"占位符处理失败: {e}")
                    replacement = ReplacementResult(
                        success=False,
                        original_placeholder=placeholder["placeholder_text"],
                        replacement_content=placeholder["placeholder_text"],
                        content_type="error",
                        error_message=str(e),
                    )
                    replacements.append(replacement)

            # 3. 应用替换到文档
            success = await self._apply_replacements_to_document(
                template_path, output_path, replacements
            )

            processing_time = (datetime.now() - start_time).total_seconds()

            return TemplateProcessingResult(
                success=success,
                processed_file_path=output_path if success else None,
                replacements=replacements,
                total_placeholders=len(placeholders),
                successful_replacements=successful_count,
                processing_time=processing_time,
            )

        except Exception as e:
            logger.error(f"传统模板处理失败: {e}")
            processing_time = (datetime.now() - start_time).total_seconds()

            return TemplateProcessingResult(
                success=False, error_message=str(e), processing_time=processing_time
            )

    async def _apply_replacements_to_document(
        self,
        template_path: str,
        output_path: str,
        replacements: List[ReplacementResult],
    ) -> bool:
        """应用替换到文档"""

        try:
            # 加载文档
            doc = docx.Document(template_path)

            # 创建替换映射
            replacement_map = {
                r.original_placeholder: r.replacement_content
                for r in replacements
                if r.success
            }

            # 替换段落中的占位符
            for paragraph in doc.paragraphs:
                for replacement in replacements:
                    placeholder = replacement.original_placeholder
                    if placeholder in paragraph.text:
                        if replacement.content_type == "chart_result" and hasattr(replacement, 'chart_path') and replacement.chart_path:
                            # 处理图表替换
                            self._replace_with_chart(paragraph, placeholder, replacement.chart_path)
                        else:
                            # 处理文本替换
                            paragraph.text = paragraph.text.replace(
                                placeholder, replacement.replacement_content
                            )

            # 替换表格中的占位符
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for replacement in replacements:
                            placeholder = replacement.original_placeholder
                            cell_text = ''.join([p.text for p in cell.paragraphs])
                            if placeholder in cell_text:
                                if replacement.content_type == "chart_result" and hasattr(replacement, 'chart_path') and replacement.chart_path:
                                    # 处理图表替换
                                    self._replace_with_chart_in_cell(cell, placeholder, replacement.chart_path)
                                else:
                                    # 处理文本替换
                                    for paragraph in cell.paragraphs:
                                        if placeholder in paragraph.text:
                                            paragraph.text = paragraph.text.replace(
                                                placeholder, replacement.replacement_content
                                            )

            # 保存文档
            doc.save(output_path)
            logger.info(f"文档已保存到: {output_path}")

            return True

        except Exception as e:
            logger.error(f"文档替换失败: {e}")
            return False

    async def generate_content_for_placeholder(
        self,
        placeholder_type: str,
        processed_data: Any,
        format_config: Optional[FormatConfig] = None,
    ) -> GeneratedContent:
        """为占位符生成内容"""

        if format_config is None:
            format_config = self.format_config

        content_generator = ContentGenerator()
        return await content_generator.generate_content(
            placeholder_type, processed_data, format_config
        )

    async def generate_chart_for_placeholder(
        self,
        chart_data: List[Dict[str, Any]],
        chart_config: Optional[ChartConfig] = None,
        output_format: str = "description",
    ) -> ChartResult:
        """为占位符生成图表"""

        if chart_config is None:
            chart_config = self.chart_config

        chart_generator = ChartGenerator()
        return await chart_generator.generate_chart(
            chart_data, chart_config, output_format
        )

    def insert_chart_into_document(
        self, doc: docx.Document, chart_result: ChartResult, placeholder_text: str
    ) -> bool:
        """将图表插入到文档中"""

        try:
            if chart_result.file_path:
                # 查找包含占位符的段落
                for paragraph in doc.paragraphs:
                    if placeholder_text in paragraph.text:
                        # 替换占位符文本
                        paragraph.text = paragraph.text.replace(placeholder_text, "")
                        # 插入图片
                        run = (
                            paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                        )
                        run.add_picture(chart_result.file_path, width=Inches(6))
                        return True

                # 如果在段落中没找到，检查表格
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if placeholder_text in paragraph.text:
                                    paragraph.text = paragraph.text.replace(
                                        placeholder_text, ""
                                    )
                                    run = (
                                        paragraph.runs[0]
                                        if paragraph.runs
                                        else paragraph.add_run()
                                    )
                                    run.add_picture(
                                        chart_result.file_path, width=Inches(4)
                                    )
                                    return True

            elif chart_result.description:
                # 使用图表描述替换
                for paragraph in doc.paragraphs:
                    if placeholder_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(
                            placeholder_text, chart_result.description
                        )
                        return True

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if placeholder_text in paragraph.text:
                                    paragraph.text = paragraph.text.replace(
                                        placeholder_text, chart_result.description
                                    )
                                    return True

            return False

        except Exception as e:
            logger.error(f"图表插入失败: {e}")
            return False

    def validate_template(self, file_path: str) -> Dict[str, Any]:
        """验证模板文件"""

        try:
            # 检查文件是否存在
            import os

            if not os.path.exists(file_path):
                return {"valid": False, "error": "模板文件不存在"}

            # 检查文件格式
            if not file_path.endswith(".docx"):
                return {"valid": False, "error": "仅支持.docx格式的模板文件"}

            # 尝试打开文档
            doc = docx.Document(file_path)

            # 解析占位符
            placeholder_data = self.parse_intelligent_placeholders(file_path)
            placeholders = placeholder_data.get("intelligent_placeholders", [])

            return {
                "valid": True,
                "placeholder_count": len(placeholders),
                "placeholders": placeholders,
                "file_size": os.path.getsize(file_path),
            }

        except Exception as e:
            return {"valid": False, "error": f"模板验证失败: {str(e)}"}

    def _extract_final_content(self, workflow_result) -> tuple[str, str, Optional[str]]:
        """
        从工作流结果中提取最终内容用于替换占位符
        
        Returns:
            tuple of (final_content, content_type, chart_path)
        """
        try:
            final_content = ""
            content_type = "text"
            chart_path = None
            
            if not hasattr(workflow_result, 'results'):
                return "数据处理完成", "text", None
            
            results = workflow_result.results
            
            # 优先级：内容生成 > 图表 > 分析结果 > 原始数据
            if "generate_content" in results and results["generate_content"].success:
                content_result = results["generate_content"].data
                if hasattr(content_result, 'content'):
                    final_content = content_result.content
                    content_type = "generated_content"
            
            elif "generate_report" in results and results["generate_report"].success:
                content_result = results["generate_report"].data
                if hasattr(content_result, 'content'):
                    final_content = content_result.content
                    content_type = "generated_content"
                    
            elif "generate_description" in results and results["generate_description"].success:
                content_result = results["generate_description"].data
                if hasattr(content_result, 'content'):
                    final_content = content_result.content
                    content_type = "generated_content"
            
            # 检查图表结果
            chart_steps = ["create_chart", "create_summary_chart", "create_trend_chart"]
            for step in chart_steps:
                if step in results and results[step].success:
                    chart_result = results[step].data
                    if hasattr(chart_result, 'file_path') and chart_result.file_path:
                        chart_path = chart_result.file_path
                        if not final_content and hasattr(chart_result, 'description'):
                            final_content = chart_result.description
                            content_type = "chart_result"
                        break
            
            # 如果没有内容生成，尝试从分析结果生成简单描述
            if not final_content:
                analysis_steps = ["analyze_data", "descriptive_analysis", "trend_analysis"]
                for step in analysis_steps:
                    if step in results and results[step].success:
                        analysis_result = results[step].data
                        if hasattr(analysis_result, 'summary'):
                            final_content = analysis_result.summary
                            content_type = "analysis_summary"
                            break
                        elif hasattr(analysis_result, 'results') and isinstance(analysis_result.results, dict):
                            # 生成简单的统计摘要
                            stats = []
                            for key, value in analysis_result.results.items():
                                if isinstance(value, dict) and "mean" in value:
                                    stats.append(f"{key}平均值: {value['mean']:.2f}")
                                elif isinstance(value, (int, float)):
                                    stats.append(f"{key}: {value}")
                            if stats:
                                final_content = "，".join(stats[:3])  # 取前3个统计项
                                content_type = "analysis_summary"
                            break
            
            # 最后的fallback
            if not final_content:
                if "fetch_data" in results and results["fetch_data"].success:
                    data_result = results["fetch_data"].data
                    if hasattr(data_result, 'row_count'):
                        final_content = f"数据查询完成，共{data_result.row_count}条记录"
                        content_type = "data_summary"
                    else:
                        final_content = "数据处理完成"
                        content_type = "text"
                else:
                    final_content = "处理完成"
                    content_type = "text"
            
            return final_content, content_type, chart_path
            
        except Exception as e:
            logger.error(f"提取最终内容失败: {e}")
            return "内容生成失败", "error", None

    def _replace_with_chart(self, paragraph, placeholder: str, chart_path: str):
        """在段落中用图表替换占位符"""
        try:
            # 先清除占位符文本
            paragraph.text = paragraph.text.replace(placeholder, "")
            
            # 添加图片
            if os.path.exists(chart_path):
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                
                # 根据段落位置确定图片大小
                from docx.shared import Inches
                if "表格" in paragraph.text or paragraph._element.getparent().tag.endswith('tc'):
                    # 在表格中，使用较小尺寸
                    run.add_picture(chart_path, width=Inches(4))
                else:
                    # 在段落中，使用标准尺寸  
                    run.add_picture(chart_path, width=Inches(6))
                    
                logger.info(f"图表已插入: {chart_path}")
            else:
                logger.warning(f"图表文件不存在: {chart_path}")
                # 如果图表文件不存在，插入描述文本
                paragraph.text = f"[图表: {placeholder}]"
                
        except Exception as e:
            logger.error(f"图表插入失败: {e}")
            paragraph.text = f"[图表插入失败: {placeholder}]"

    def _replace_with_chart_in_cell(self, cell, placeholder: str, chart_path: str):
        """在表格单元格中用图表替换占位符"""
        try:
            for paragraph in cell.paragraphs:
                if placeholder in paragraph.text:
                    self._replace_with_chart(paragraph, placeholder, chart_path)
                    break
        except Exception as e:
            logger.error(f"表格单元格图表替换失败: {e}")

    def _detect_template_type(self, template_path: str) -> str:
        """检测模板类型"""
        import os
        
        if not os.path.exists(template_path):
            return "unknown"
            
        file_extension = os.path.splitext(template_path)[1].lower()
        
        type_mapping = {
            '.docx': 'docx',
            '.doc': 'docx',
            '.xlsx': 'xlsx', 
            '.xls': 'xlsx',
            '.html': 'html',
            '.htm': 'html',
            '.pdf': 'pdf',
            '.txt': 'text'
        }
        
        return type_mapping.get(file_extension, 'text')


template_parser = TemplateParser()
