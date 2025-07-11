import docx
import io
import base64
import re
from typing import Dict, Any, List

class WordGeneratorService:
    TABLE_PLACEHOLDER_REGEX = re.compile(r"\[table:(\w+)\]")

    def generate_report(
        self, 
        template_path: str, 
        output_path: str, 
        data: Dict[str, Any]
    ):
        """
        Generates a .docx report by filling a template with provided data.
        - Replaces {{text_placeholder}} with string values.
        - Replaces [chart:chart_placeholder] with a base64 encoded image.
        - Fills tables marked with [table:table_placeholder] in the first cell.
        """
        doc = docx.Document(template_path)

        # --- 1. Fill Text and Chart Placeholders in Paragraphs ---
        for p in doc.paragraphs:
            self._replace_text_in_run(p, data)
            self._replace_chart_in_run(p, data)

        # --- 2. Fill Table Placeholders ---
        for table in doc.tables:
            # Check the first cell for a table placeholder
            first_cell_text = table.cell(0, 0).text
            match = self.TABLE_PLACEHOLDER_REGEX.search(first_cell_text)
            
            if match:
                key = match.group(1)
                table_data = data.get(key)
                
                if isinstance(table_data, list) and table_data:
                    # Clear the placeholder text from the first cell
                    table.cell(0, 0).text = table.cell(0, 0).text.replace(match.group(0), "").strip()

                    # Get headers from the first row of the template table
                    headers = [cell.text for cell in table.rows[0].cells]
                    
                    # Add data rows
                    for item in table_data:
                        row_cells = table.add_row().cells
                        for i, header in enumerate(headers):
                            # Allow for case-insensitive and flexible key matching
                            cell_value = self._find_value_for_header(header, item)
                            row_cells[i].text = str(cell_value)
        
        doc.save(output_path)

    def _replace_text_in_run(self, paragraph, data: Dict[str, Any]):
        # Simple text replacement for {{key}}
        for key, value in data.items():
            if isinstance(value, (str, int, float)) and f"{{{{{key}}}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))
    
    def _replace_chart_in_run(self, paragraph, data: Dict[str, Any]):
        # Replaces [chart:key] with a base64 image
        for key, value in data.items():
            if f"[chart:{key}]" in paragraph.text and self._is_base64(value):
                # Clear the placeholder text
                paragraph.text = paragraph.text.replace(f"[chart:{key}]", "")
                # Add image from base64 string
                try:
                    image_stream = io.BytesIO(base64.b64decode(value))
                    paragraph.add_run().add_picture(image_stream, width=docx.shared.Inches(5.0))
                except Exception as e:
                    print(f"Error adding picture for key {key}: {e}")

    def _find_value_for_header(self, header: str, item: Dict[str, Any]) -> Any:
        # Tries to find a value in the data item, ignoring case and looking for partial matches
        header_lower = header.lower().strip()
        for key, value in item.items():
            if key.lower().strip() == header_lower:
                return value
        return "" # Return empty string if no match found

    def _is_base64(self, s: Any) -> bool:
        if not isinstance(s, str):
            return False
        try:
            return base64.b64encode(base64.b64decode(s)).decode() == s
        except Exception:
            return False

word_generator_service = WordGeneratorService()
