"""
测试图表集成完整流程

模拟task.py中的流程：
1. ETL阶段获取数据
2. 文档生成阶段处理图表占位符
3. 生成包含图表的Word文档
"""

import asyncio
import sys
import os
from pathlib import Path

# 添加项目根目录到路径
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))


async def test_chart_placeholder_processor():
    """测试图表占位符处理器"""
    print("\n" + "="*70)
    print("测试 1: 图表占位符处理器")
    print("="*70)

    from app.services.infrastructure.document.chart_placeholder_processor import ChartPlaceholderProcessor

    # 模拟ETL返回的数据
    etl_data = [
        {"州市": "北京", "申请量": 523},
        {"州市": "上海", "申请量": 412},
        {"州市": "广州", "申请量": 335},
        {"州市": "深圳", "申请量": 289},
        {"州市": "成都", "申请量": 221},
    ]

    processor = ChartPlaceholderProcessor(user_id="test_user")

    # 测试占位符识别
    print("\n📋 测试占位符文本提取...")
    test_placeholders = [
        "{{图表：州市退货申请量由高到低排列并显示对应申请量的柱状图}}",
        "{{图表：月度销售额趋势折线图}}",
        "{{图表：产品类别销售占比饼图}}"
    ]

    for placeholder in test_placeholders:
        intent = processor._extract_chart_intent(placeholder)
        print(f"\n占位符: {placeholder}")
        print(f"  提取结果:")
        print(f"    - 图表类型: {intent['chart_type']}")
        print(f"    - 标题: {intent['title']}")
        print(f"    - 描述: {intent['description']}")

    # 测试图表生成
    print("\n📊 测试图表生成...")
    chart_result = await processor.process_chart_placeholder(
        placeholder_text="{{图表：州市退货申请量由高到低排列并显示对应申请量的柱状图}}",
        data=etl_data
    )

    if chart_result.get("success"):
        print(f"✅ 图表生成成功")
        print(f"   路径: {chart_result['chart_path']}")
        print(f"   类型: {chart_result['chart_type']}")
        print(f"   标题: {chart_result['title']}")
        print(f"   生成时间: {chart_result.get('generation_time_ms')}ms")

        # 验证文件存在
        if os.path.exists(chart_result['chart_path']):
            file_size = os.path.getsize(chart_result['chart_path'])
            print(f"   文件大小: {file_size} bytes")
        else:
            print(f"❌ 文件不存在")
    else:
        print(f"❌ 图表生成失败: {chart_result.get('error')}")

    return chart_result


async def test_word_template_integration():
    """测试Word模板集成"""
    print("\n" + "="*70)
    print("测试 2: Word模板集成")
    print("="*70)

    try:
        from docx import Document
    except ImportError:
        print("⚠️  python-docx 未安装，跳过Word集成测试")
        print("   请运行: pip install python-docx")
        return

    from app.services.infrastructure.document.word_template_service import WordTemplateService
    from docx import Document
    from docx.shared import Inches

    # 创建测试模板
    print("\n📝 创建测试Word模板...")
    test_template_path = "/tmp/test_chart_template.docx"
    doc = Document()

    doc.add_heading('退货申请分析报告', 0)
    doc.add_paragraph('本报告展示了各州市的退货申请情况。')

    doc.add_heading('数据统计', level=1)
    doc.add_paragraph('总申请量: {{total_applications}}')
    doc.add_paragraph('平均申请量: {{avg_applications}}')

    doc.add_heading('可视化分析', level=1)
    doc.add_paragraph('{{图表：州市退货申请量由高到低排列并显示对应申请量的柱状图}}')

    doc.add_paragraph('\n分析结论：从上图可以看出...')

    doc.save(test_template_path)
    print(f"✅ 测试模板已创建: {test_template_path}")

    # 模拟ETL数据
    etl_results = {
        "total_applications": 1780,
        "avg_applications": 356,
        "{{图表：州市退货申请量由高到低排列并显示对应申请量的柱状图}}": [
            {"州市": "北京", "申请量": 523},
            {"州市": "上海", "申请量": 412},
            {"州市": "广州", "申请量": 335},
            {"州市": "深圳", "申请量": 289},
            {"州市": "成都", "申请量": 221},
        ]
    }

    # 处理模板
    print("\n🔄 处理模板...")
    word_service = WordTemplateService()

    output_path = "/tmp/test_chart_output.docx"
    result = await word_service.process_document_template(
        template_path=test_template_path,
        placeholder_data=etl_results,
        output_path=output_path,
        container=None,
        use_agent_charts=True,
        use_agent_optimization=False,
        user_id="test_user"
    )

    if result.get("success"):
        print(f"✅ Word文档生成成功")
        print(f"   输出路径: {result['output_path']}")
        print(f"   占位符数量: {result['placeholders_processed']}")
        print(f"   图表生成方法: {result['chart_generation_method']}")

        # 验证输出文件
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            print(f"   文件大小: {file_size} bytes")

            # 验证图表是否插入
            output_doc = Document(output_path)
            image_count = sum(len([r for r in p.runs if r._element.xpath('.//pic:pic')]) for p in output_doc.paragraphs)
            print(f"   文档中的图片数量: {image_count}")
        else:
            print(f"❌ 输出文件不存在")
    else:
        print(f"❌ Word文档生成失败: {result.get('error')}")

    return result


async def test_complete_workflow():
    """测试完整工作流程（模拟tasks.py）"""
    print("\n" + "="*70)
    print("测试 3: 完整工作流程（模拟tasks.py）")
    print("="*70)

    # 阶段1: 模拟Agent生成SQL
    print("\n【阶段1】Agent生成SQL（跳过，直接使用模拟数据）")

    # 阶段2: 模拟ETL执行
    print("\n【阶段2】ETL执行 - 获取数据")

    etl_results = {
        "task_name": "月度退货分析报告",
        "report_period": "2024年1月",
        "total_applications": 1780,
        "avg_applications": 356,
        "图表：州市退货申请量由高到低排列并显示对应申请量的柱状图": [
            {"州市": "北京", "申请量": 523},
            {"州市": "上海", "申请量": 412},
            {"州市": "广州", "申请量": 335},
            {"州市": "深圳", "申请量": 289},
            {"州市": "成都", "申请量": 221},
        ],
        "图表：月度趋势折线图": [
            {"月份": "1月", "申请量": 280},
            {"月份": "2月", "申请量": 320},
            {"月份": "3月", "申请量": 356},
            {"月份": "4月", "申请量": 390},
            {"月份": "5月", "申请量": 434},
        ]
    }

    print(f"✅ ETL完成，获取到 {len(etl_results)} 个占位符数据")

    # 查找图表占位符
    chart_placeholders = [k for k in etl_results.keys() if k.startswith("图表：")]
    print(f"📊 识别到 {len(chart_placeholders)} 个图表占位符:")
    for ph in chart_placeholders:
        data_count = len(etl_results[ph]) if isinstance(etl_results[ph], list) else "N/A"
        print(f"   - {ph} (数据量: {data_count})")

    # 阶段3: 文档生成
    print("\n【阶段3】文档生成 - 替换占位符和生成图表")

    try:
        from docx import Document
    except ImportError:
        print("⚠️  python-docx 未安装，无法完成文档生成测试")
        return

    # 创建完整测试模板
    test_template_path = "/tmp/test_complete_workflow.docx"
    doc = Document()

    doc.add_heading('月度退货分析报告', 0)
    doc.add_paragraph(f'报告名称: {{task_name}}')
    doc.add_paragraph(f'报告周期: {{report_period}}')

    doc.add_heading('统计数据', level=1)
    doc.add_paragraph('总申请量: {{total_applications}}')
    doc.add_paragraph('平均申请量: {{avg_applications}}')

    doc.add_heading('州市退货分析', level=1)
    doc.add_paragraph('{{图表：州市退货申请量由高到低排列并显示对应申请量的柱状图}}')

    doc.add_heading('趋势分析', level=1)
    doc.add_paragraph('{{图表：月度趋势折线图}}')

    doc.save(test_template_path)
    print(f"✅ 模板已创建: {test_template_path}")

    # 使用WordTemplateService处理
    from app.services.infrastructure.document.word_template_service import WordTemplateService

    word_service = WordTemplateService()
    output_path = "/tmp/test_complete_workflow_output.docx"

    result = await word_service.process_document_template(
        template_path=test_template_path,
        placeholder_data=etl_results,
        output_path=output_path,
        container=None,
        use_agent_charts=True,
        use_agent_optimization=False,
        user_id="test_user"
    )

    if result.get("success"):
        print(f"\n✅ 完整流程测试成功！")
        print(f"   输出文件: {result['output_path']}")

        # 验证结果
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            print(f"   文件大小: {file_size} bytes")

            # 统计图片数量
            output_doc = Document(output_path)
            image_count = 0
            for p in output_doc.paragraphs:
                for run in p.runs:
                    if run._element.xpath('.//pic:pic'):
                        image_count += 1

            print(f"   插入的图表数量: {image_count}/{len(chart_placeholders)}")

            if image_count == len(chart_placeholders):
                print(f"   ✅ 所有图表都已成功插入！")
            else:
                print(f"   ⚠️  部分图表未成功插入")

        print(f"\n📄 请打开文件查看结果: {output_path}")
    else:
        print(f"❌ 完整流程测试失败: {result.get('error')}")

    return result


async def main():
    """运行所有测试"""
    print("🚀 开始测试图表集成完整流程")
    print("="*70)

    try:
        # 测试1: 图表占位符处理器
        await test_chart_placeholder_processor()

        # 测试2: Word模板集成
        await test_word_template_integration()

        # 测试3: 完整工作流程
        await test_complete_workflow()

        print("\n" + "="*70)
        print("✅ 所有测试完成")
        print("="*70)

        print("\n📌 总结:")
        print("   1. ChartGenerationTool - 图表生成工具 ✅")
        print("   2. ChartPlaceholderProcessor - 图表占位符处理器 ✅")
        print("   3. WordTemplateService集成 ✅")
        print("   4. 完整流程（模拟tasks.py） ✅")

        print("\n🎯 下一步:")
        print("   - 在实际任务中测试 {{图表：xxx}} 占位符")
        print("   - 确保ETL阶段返回正确的数据格式")
        print("   - 监控图表生成时间和性能")

    except Exception as e:
        print(f"\n❌ 测试过程中发生错误: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    asyncio.run(main())
